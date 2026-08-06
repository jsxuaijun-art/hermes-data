# -*- coding: utf-8 -*-
"""
高考英语词汇默写本 · 词根版v4（按天编排+学习计划索引）
- 78组 × 50词（28组词根 + 50组基础）
- 107日历天：每天学1组新词 + 复习到期的老组
- 前置学习计划索引（Day 1→107）
- 正文按天排列，翻到哪天做哪天
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import re, math, os, time, random

# ═══════════════════════════════════════════
# CONFIG
# ═══════════════════════════════════════════
SRC  = "/mnt/d/OneDrive/Desktop/tianyi_download/高考英语词汇表3817词.docx"
FREQ = "/home/administrator/.hermes/skills/productivity/vocab-memory-book/en_50k.txt"
OUT  = "/mnt/d/OneDrive/Desktop/tianyi_download/高考英语词汇默写本_词根版.docx"
WPG  = 50

# ═══════════════════════════════════════════
# MEMBEAN ROOTS (255+)
# ═══════════════════════════════════════════
RAW_ROOTS = """act|drive, do
ag|do, act, drive
alt|high
alter|another
anim|mind, spirit
ann|year
anthrop|human
aud|hear, listen to
auto|self, same
bell|war
bene|well
bio|life
brev|short
cad|fall
cap|take, seize
capit|head
carn|meat
ced|go, move, yield
ceed|go, move
celer|swift
centr|center
cern|sift, perceive
chron|time
circ|ring
cis|cut
civ|citizen
clam|shout
clar|clear
clin|lean
clud|shut, close
cogn|learn, know
cord|heart
corp|body
crat|rule
cred|believe
cresc|grow
crit|judge
cult|grow, cultivate
curr|run, course
cycl|circle, wheel
dem|people
dict|say
doc|teach
domin|master
duc|lead
dur|hard, harsh
dyn|power
err|wander, mistake
fac|make, do
fact|made, done
fend|ward off
fer|carry, bring, bear
fid|trust, faith
flect|bent
flict|struck
flor|flower
flu|flow
form|shape
fort|strong
fract|broken
frag|break
fug|flee
fus|pour
gen|born, produced, kind
gest|bear, carry
grad|step
graph|write, draw
grat|pleasing
grav|heavy, serious
greg|flock
gress|step
her|stick, cling
hibit|hold
jac|throw
ject|thrown
judic|judge
junct|joined
labor|work
lat|carry, wide
lect|gather, choose, read
leg|law, choose, read
liber|free
lig|tie, bind
limin|threshold
liter|letter
loc|place
log|word, study, reason
loqu|talk, speak
luc|light, brightness
lud|play
lumin|light
magn|large, great
mal|bad, evil
man|hand
mand|order
mania|madness
matr|mother
medi|middle
memor|remember
ment|mind
merc|trade
merg|plunge, sink
meter|measure, meter
migr|move, change place
min|lessen, make smaller
mir|wonder
miss|sent
mit|send, let go
mon|warn, advise
morph|shape, form
mort|death
mot|move
mov|move
mut|change, alter
nat|born
nav|ship
nect|tie, bind
neg|deny
neur|nerve
noc|harm
nom|name
nov|new
numer|number
nutr|nourish
oper|work
opt|choose, wish
ord|order, row
ori|rise, begin
orn|decorate
part|part, share
pass|suffer, endure
path|feeling, disease
patr|father
ped|foot
pel|push, drive
pend|hang, weigh
pens|hang, weigh
pet|seek, attack
phil|love
phon|sound, voice
pict|paint
plac|please, suit
plaud|applaud
plen|full, fill
plex|weave, braid
plic|fold
plor|weep
pod|foot
pon|put, place
popul|people
port|carry
pos|put, place
pot|power, ability
preci|worth, price
press|press
prim|first
priv|individual, separate
prob|prove, test
puls|push, drive
punct|point, prick
put|think, consider
quest|ask, seek
quiet|rest, still
radi|ray, spoke
rap|snatch, seize
ras|scrape, scratch
rect|right, straight
reg|rule, guide, direct
rid|laugh
riv|stream
rod|chew, eat away
rog|ask, request
rupt|burst, break
sacr|sacred, holy
sal|leap, jump
sal|health, safety
sanct|holy
sci|know
scrib|write
script|written
sect|cut
sed|sit, settle
sens|feel, sense, perceive
sent|feel, sense
sequ|follow
serv|keep, save, serve
sid|sit
sign|mark, seal, sign
simil|like, similar
sist|stand
sol|alone
solv|loosen, untie
somn|sleep
son|sound
soph|wise
spec|see, look
spect|see, look, watch over
spers|scatter
spir|breathe
spond|pledge, promise
stant|stand, stand still
stas|standing, standing still
stat|stand, station
string|draw tight, tighten
struct|build, arrange
surg|rise
tact|touched, touch
tain|have, hold
tang|touch, affect
techn|skill, art, craft
tend|stretch, extend
tens|stretch
ten|hold
tenu|thin
termin|boundary, end, limit
terr|earth, land
terr|frighten
test|witness, proof, evidence
therm|heat
tim|fear
tom|cut
tort|twisted
tox|poison
tract|drag, pull, draw
trop|turn, change
trud|thrust, push
turb|confusion, disorder
umbr|shadow, shade
und|wave
urb|city
vac|empty
vad|go
vag|wander, roam
val|be strong, be of value
van|empty, illusory
vari|different, diverse
veh|carry, bring
vel|curtain
ven|come
vent|come
verb|word
verg|move, turn
vers|turned, changed
vert|turn
ver|truth, true
vest|clothing, garments
vid|see
vig|be lively, active
vinc|conquer, win
vir|man, male
vis|see, look at
vit|life
viv|live, alive
voc|call, voice
vok|call, summon
vol|wish, want
volv|roll, wrap, turn
vor|eat, devour
vot|promise, pledge
cess|go, yield, move
cent|hundred
vit|vital, life
equ|equal, even
vol|fly
cura|care, attention
here|stick
gratia|favor, thanks
multi|many, much
omni|all, every
re|again, back"""

ROOTS = []
for line in RAW_ROOTS.strip().split('\n'):
    parts = line.split('|', 1)
    if len(parts) == 2: ROOTS.append((parts[0].strip(), parts[1].strip()))
ROOTS.sort(key=lambda x: -len(x[0]))

def detect_root(word):
    w = word.lower()
    for root, meaning in ROOTS:
        if root in w:
            if len(root) <= 2 and len(w) <= 4: continue
            return (root, meaning)
    return None

# ═══════════════════════════════════════════
# 1. LOAD VOCABULARY FROM DOCX
# ═══════════════════════════════════════════
print("Loading vocabulary from docx...")
from docx import Document as DocxReader
doc_src = DocxReader(SRC)

all_words_parsed = []
for table in doc_src.tables:
    for ri, row in enumerate(table.rows):
        if ri == 0: continue
        cells = [c.text.strip() for c in row.cells]
        if len(cells) >= 3:
            word = cells[0].strip().lower()
            phon = cells[1].strip()
            defn = cells[2].strip()[:90]
            if word:
                all_words_parsed.append({
                    'word': word, 'phonetic': phon, 'definition': defn,
                    'root': '', 'root_meaning': '',
                })
print(f"  {len(all_words_parsed)} words loaded")

# Detect roots
for w in all_words_parsed:
    ri = detect_root(w['word'])
    if ri:
        w['root'] = ri[0]
        w['root_meaning'] = ri[1]

# ═══════════════════════════════════════════
# 2. LOAD FREQUENCY DATA
# ═══════════════════════════════════════════
print("Loading frequency data...")
freq_data = {}
with open(FREQ, 'r') as f:
    for l in f:
        l = l.strip()
        if not l: continue
        parts = l.split(None, 1)
        if len(parts) == 2: freq_data[parts[0].lower()] = int(parts[1])
fv = sorted(freq_data.values(), reverse=True)
p80 = fv[max(0, len(fv)//5*4)] if fv else 0
p60 = fv[max(0, len(fv)//5*3)] if fv else 0
p40 = fv[max(0, len(fv)//5*2)] if fv else 0
p20 = fv[max(0, len(fv)//5*1)] if fv else 0
def stars(w):
    f = freq_data.get(w)
    if f is None: return 1
    if f >= p80: return 5
    if f >= p60: return 4
    if f >= p40: return 3
    if f >= p20: return 2
    return 1
for w in all_words_parsed:
    w['stars'] = stars(w['word'])

# ═══════════════════════════════════════════
# 3. SPLIT AND GROUP
# ═══════════════════════════════════════════
root_words = [w for w in all_words_parsed if w['root']]
basic_words = [w for w in all_words_parsed if not w['root']]
random.seed(20260704)
random.shuffle(root_words)
random.shuffle(basic_words)

root_groups = []
i = 0
while i < len(root_words):
    chunk = root_words[i:i+WPG]
    roots_in = list(set(w['root'] for w in chunk))
    root_groups.append({
        'label': '+'.join(roots_in[:3]) + ('...' if len(roots_in) > 3 else ''),
        'words': chunk,
        'part': 1,
        'group_num': len(root_groups) + 1,
    })
    i += WPG

basic_groups = []
i = 0
while i < len(basic_words):
    chunk = basic_words[i:i+WPG]
    basic_groups.append({
        'label': '基础词汇',
        'words': chunk,
        'part': 2,
        'group_num': len(basic_groups) + 1,
    })
    i += WPG

# Sequential index for all groups (Part 1 then Part 2)
all_groups = root_groups + basic_groups
for gidx, g in enumerate(all_groups):
    g['seq'] = gidx + 1  # 1-based sequential number

TOTAL_WORDS = len(all_words_parsed)
TOTAL_GROUPS = len(all_groups)
TOTAL_LEARN_DAYS = TOTAL_GROUPS
TOTAL_DAYS = TOTAL_LEARN_DAYS + 29  # last group's final review

print(f"  Root: {len(root_words)}w/{len(root_groups)}g | Basic: {len(basic_words)}w/{len(basic_groups)}g")
print(f"  Total: {TOTAL_GROUPS} groups, {TOTAL_DAYS} calendar days")

# ═══════════════════════════════════════════
# 4. BUILD DAY SCHEDULE
# ═══════════════════════════════════════════
# Each day: (day_num, new_group_info_or_None, list_of_review_tuples)
# review_tuple = (group_seq, part, group_num_in_part, label, round_number, round_name, mode)
# Round mapping:
#   Round 1 (学习日, 0 days): 英译汉 → cards shown (handled separately)
#   Round 2 (+1 day): 英→中 → mode='en2cn'
#   Round 3 (+3 days): 中→英 → mode='cn2en'
#   Round 4 (+6 days): 音标默写 → mode='phonetic'
#   Round 5 (+14 days): 高频词 中→英 → mode='highfreq' (cn2en for stars>=4 only)
#   Round 6 (+29 days): 混合终测 → mode='mixed'

ROUNDS = [
    (1,  '英→中',    'en2cn'),
    (3,  '中→英',    'cn2en'),
    (6,  '音标默写',  'phonetic'),
    (14, '高频词',    'highfreq'),
    (29, '混合终测',  'mixed'),
]

schedule = []
for day in range(1, TOTAL_DAYS + 1):
    new_group = None
    reviews = []

    # New group on this day?
    learn_idx = day - 1  # Group seq that starts learning on this day
    if learn_idx < TOTAL_GROUPS:
        g = all_groups[learn_idx]
        part_label = f"Part {g['part']}"
        group_label = f"第{g['group_num']:02d}组" if g['part'] == 1 else f"第{g['group_num']:02d}组"
        new_group = {
            'seq': g['seq'],
            'part': g['part'],
            'group_num': g['group_num'],
            'label': g['label'],
            'words': g['words'],
        }

    # Reviews due on this day?
    for check_gidx, g in enumerate(all_groups):
        learn_day = check_gidx + 1
        for offset, rname, mode in ROUNDS:
            if day == learn_day + offset:
                # Special case: highfreq only shows words with stars>=4
                sl = [w for w in g['words'] if w['stars'] >= 4] if mode == 'highfreq' else None
                reviews.append({
                    'seq': g['seq'],
                    'part': g['part'],
                    'group_num': g['group_num'],
                    'label': g['label'],
                    'round_offset': offset,
                    'round_name': rname,
                    'mode': mode,
                    'words': sl if mode == 'highfreq' else g['words'],
                    'highfreq_only': mode == 'highfreq',
                })

    schedule.append({
        'day': day,
        'new_group': new_group,
        'reviews': reviews,
    })

print(f"  Schedule built: {len(schedule)} days")

# ═══════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════
def set_shd(c, color):
    c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))
def set_bdr(c, **kw):
    tc = c._tc; p = tc.get_or_add_tcPr(); b = OxmlElement('w:tcBorders')
    for edge, val in kw.items():
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), val.get('val', 'single'))
        e.set(qn('w:sz'), val.get('sz', '4'))
        e.set(qn('w:color'), val.get('color', '000000'))
        e.set(qn('w:space'), '0')
        b.append(e); p.append(b)
def hex_rgb(h):
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)

def group_label_short(g):
    if g['part'] == 1:
        return f"Part 1 · 第{g['group_num']:02d}组"
    else:
        return f"Part 2 · 第{g['group_num']:02d}组"

# ═══ COVER ═══
def build_cover(doc, root_groups, basic_groups, all_words_parsed):
    bg = doc.add_table(rows=1, cols=1)
    bg.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = bg._tbl; tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    for c in list(tblPr):
        if c.tag == qn('w:tblW'): tblPr.remove(c)
    tw = OxmlElement('w:tblW')
    tw.set(qn('w:w'), '11906'); tw.set(qn('w:type'), 'dxa')
    tblPr.append(tw)
    cell = bg.cell(0, 0)
    set_shd(cell, '0A2463')
    set_bdr(cell, top={'val':'none','sz':'0','color':'auto'},
            bottom={'val':'none','sz':'0','color':'auto'},
            start={'val':'none','sz':'0','color':'auto'},
            end={'val':'none','sz':'0','color':'auto'})
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), '11906'); tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center'); tcPr.append(vAlign)
    def ctr(b, a):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(b)
        p.paragraph_format.space_after = Pt(a)
        p.paragraph_format.line_spacing = Pt(14)
        return p
    r = ctr(50, 10).add_run('✨    ⭐    ✨    ✨    ⭐    ✨')
    r.font.size = Pt(18); r.font.color.rgb = RGBColor(0xFF, 0xE7, 0x82)
    p = ctr(60, 10)
    r = p.add_run('高考英语词汇默写本')
    r.font.size = Pt(40); r.font.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p = ctr(8, 16)
    r = p.add_run('词根版 · 按天编排 · 艾宾浩斯记忆法')
    r.font.size = Pt(16); r.font.color.rgb = RGBColor(0xFF, 0xD7, 0x66)
    p = ctr(40, 12)
    r = p.add_run('🌊🌊  ⛵  ⛵⛵  🌊🌊🌊  ⛵  🌊🌊')
    r.font.size = Pt(24)
    p = ctr(36, 12)
    root_w = sum(len(g['words']) for g in root_groups)
    basic_w = sum(len(g['words']) for g in basic_groups)
    r = p.add_run(f'Part 1 词根词汇 {root_w}词 ({len(root_groups)}组) · '
                  f'Part 2 基础词汇 {basic_w}词 ({len(basic_groups)}组) · '
                  f'共 {TOTAL_WORDS}词 · {TOTAL_DAYS}天')
    r.font.size = Pt(11); r.font.color.rgb = RGBColor(0xBB, 0xDD, 0xFF)
    r.font.italic = True
    p = ctr(20, 8)
    r = p.add_run('「今天的每一分努力」\r都是明天看得见的风景')
    r.font.size = Pt(12); r.font.color.rgb = RGBColor(0xDD, 0xEE, 0xFF)
    r.font.italic = True
    ctr(10, 20).add_run('⭐  ✨  ⭐  ✨  ⭐').font.size = Pt(14)
    ctr(6, 30).add_run('⛵  🌅  ⛵  ⛵  🌅  ⛵').font.size = Pt(20)
    p = ctr(28, 4)
    r = p.add_run('苏州盈信企业管理有限公司')
    r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p = ctr(2, 2)
    r = p.add_run('公司注册 · 专注财税二十五年')
    r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xBB, 0xDD, 0xFF)
    p = ctr(2, 40)
    r = p.add_run('18912633863')
    r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xD7, 0x66)

# ═══ FOOTER ═══
def add_footer(doc):
    section = doc.sections[-1]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin')
    r1 = p.add_run(); r1.font.size = Pt(7); r1._r.append(fc1)
    it1 = OxmlElement('w:instrText'); it1.text = ' PAGE '
    r2 = p.add_run(); r2.font.size = Pt(7); r2._r.append(it1)
    fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'end')
    r3 = p.add_run(); r3.font.size = Pt(7); r3._r.append(fc2)
    r4 = p.add_run(' / '); r4.font.size = Pt(7)
    r4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'), 'begin')
    r5 = p.add_run(); r5.font.size = Pt(7); r5._r.append(fc3)
    it2 = OxmlElement('w:instrText'); it2.text = ' NUMPAGES '
    r6 = p.add_run(); r6.font.size = Pt(7); r6._r.append(it2)
    fc4 = OxmlElement('w:fldChar'); fc4.set(qn('w:fldCharType'), 'end')
    r7 = p.add_run(); r7.font.size = Pt(7); r7._r.append(fc4)

# ═══ LEARNING CARD ═══
def card_learn(doc, w):
    t = doc.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    sm = '★ ' if w['stars'] >= 4 else ''
    # Cell(0,0): word + phonetic
    c = t.cell(0, 0); c.text = ''
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{sm}{w['word']}")
    r.font.size = Pt(12); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E) if w['stars'] >= 4 else RGBColor(0x33, 0x33, 0x33)
    if w['phonetic']:
        r = p.add_run(f"  {w['phonetic']}")
        r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    set_shd(c, 'F0F4F8'); c.width = Cm(7.0)
    # Cell(0,1): definition only (no underlines, use full width for more text)
    c = t.cell(0, 1); c.text = ''
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(w['definition'])
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)  # Dark navy, readable
    c.width = Cm(7.5)
    # Cell(1,0): root info
    c = t.cell(1, 0); c.text = ''
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if w['root']:
        r = p.add_run(f"🔬 词根【{w['root']}】= {w['root_meaning']}")
        r.font.size = Pt(7.5); r.font.italic = True
        r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    else:
        r = p.add_run("📖 基础词汇")
        r.font.size = Pt(7.5); r.font.italic = True
        r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    # Cell(1,1): writing index
    c = t.cell(1, 1); c.text = ''
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    dots = '●' * w['stars'] + '○' * (5 - w['stars'])
    r = p.add_run(f"写作指数: {dots}  ")
    r.font.size = Pt(7); r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    if w['stars'] >= 4:
        r = p.add_run('⭐ 高频词')
        r.font.size = Pt(7); r.font.bold = True
        r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    for rowr in t.rows:
        for cell in rowr.cells:
            set_bdr(cell, top={'val':'single','sz':'4','color':'DDD'},
                    bottom={'val':'single','sz':'4','color':'DDD'},
                    start={'val':'single','sz':'4','color':'DDD'},
                    end={'val':'single','sz':'4','color':'DDD'})
    # No empty paragraph between cards — compact layout

# ═══ REVIEW TITLE ═══
def review_title(doc, label, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(label)
    r.font.size = Pt(13); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    r = p.add_run(f'  {subtitle}')
    r.font.size = Pt(7); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ═══ NAME/DATE HEADER ═══
def name_date(doc, total):
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (txt, w, clr) in enumerate([
        ('姓名: ___________', Cm(6.0), '2C3E50'),
        ('日期: ___________', Cm(5.0), '2C3E50'),
        (f'正确: ___ / {total}', Cm(4.0), 'C0392B'),
    ]):
        c = t.cell(0, i); c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = Pt(8)
        r = p.add_run(txt)
        r.font.size = Pt(8); r.font.bold = True
        r.font.color.rgb = RGBColor(*hex_rgb(clr))
        set_shd(c, 'F5F5F5')
        set_bdr(c, top={'val':'single','sz':'4','color':'CCC'},
                bottom={'val':'single','sz':'4','color':'CCC'},
                start={'val':'single','sz':'4','color':'CCC'},
                end={'val':'single','sz':'4','color':'CCC'})
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

# ═══ 6-COLUMN REVIEW TABLE ═══
CW = [Cm(0.2), Cm(3.56), Cm(5.24), Cm(0.2), Cm(3.56), Cm(5.24)]

def sixcol(doc, wl, mode):
    total = len(wl)
    if total == 0:
        return
    rn = math.ceil(total / 2)
    t = doc.add_table(rows=rn + 1, cols=6)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = t._tbl; tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    tw = OxmlElement('w:tblW')
    tw.set(qn('w:w'), '11000'); tw.set(qn('w:type'), 'dxa')
    tblPr.append(tw)
    h = {
        'en2cn':    ['#', '英文', '中文', '#', '英文', '中文'],
        'cn2en':    ['#', '中文释义', '英文', '#', '中文释义', '英文'],
        'phonetic': ['#', '🔊音标', '英+中', '#', '🔊音标', '英+中'],
        'mixed':    ['#', '原文', '填空', '#', '原文', '填空'],
        'highfreq': ['#', '中文释义', '英文', '#', '中文释义', '英文'],
    }.get(mode, ['#', '原文', '填空', '#', '原文', '填空'])

    for ci in range(6):
        c = t.cell(0, ci); c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h[ci])
        r.font.size = Pt(8); r.font.bold = True
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        set_shd(c, 'EAEAEA')
        set_bdr(c, top={'val':'single','sz':'4','color':'AAA'},
                bottom={'val':'single','sz':'6','color':'666'},
                start={'val':'single','sz':'4','color':'AAA'},
                end={'val':'single','sz':'4','color':'AAA'})
        c.width = CW[ci]
    trPr0 = t.rows[0]._tr.get_or_add_trPr()
    trH0 = OxmlElement('w:trHeight')
    trH0.set(qn('w:val'), '400'); trH0.set(qn('w:hRule'), 'atLeast')
    trPr0.append(trH0)

    for ri in range(rn):
        for col_offset, idx in [(0, ri), (3, ri + rn)]:
            if idx >= total: continue
            w = wl[idx]
            sm = '★ ' if w['stars'] >= 4 else ''
            # Column 1/4: #
            c = t.cell(ri + 1, col_offset); c.text = ''
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(idx + 1))
            r.font.size = Pt(7); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            v = OxmlElement('w:vAlign')
            v.set(qn('w:val'), 'center')
            c._tc.get_or_add_tcPr().append(v)
            # Column 2/5: prompt text
            c = t.cell(ri + 1, col_offset + 1); c.text = ''
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            if mode == 'en2cn':
                r = p.add_run(f"{sm}{w['word']}")
                r.font.size = Pt(10); r.font.bold = True
                r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E) if w['stars'] >= 4 else RGBColor(0x33, 0x33, 0x33)
            elif mode == 'cn2en' or mode == 'highfreq':
                r = p.add_run(w['definition'][:30])
                r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            elif mode == 'phonetic':
                r = p.add_run(f"🔊{w['phonetic']}")
                r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
            elif mode == 'mixed':
                if idx % 2 == 0:
                    r = p.add_run(f"{sm}{w['word']}")
                    r.font.size = Pt(10); r.font.bold = True
                    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E) if w['stars'] >= 4 else RGBColor(0x33, 0x33, 0x33)
                else:
                    r = p.add_run(w['definition'][:28])
                    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            v = OxmlElement('w:vAlign')
            v.set(qn('w:val'), 'center')
            c._tc.get_or_add_tcPr().append(v)
            # Column 3/6: blank area (no underlines, student writes freely)
            c = t.cell(ri + 1, col_offset + 2); c.text = ''
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            # Cell intentionally left blank for student to write
            v = OxmlElement('w:vAlign')
            v.set(qn('w:val'), 'center')
            c._tc.get_or_add_tcPr().append(v)
            # Borders
            for cix in range(3):
                set_bdr(t.cell(ri + 1, col_offset + cix),
                        top={'val':'single','sz':'4','color':'EEE'},
                        bottom={'val':'single','sz':'4','color':'EEE'},
                        start={'val':'single','sz':'4','color':'EEE'},
                        end={'val':'single','sz':'4','color':'EEE'})
        trPr = t.rows[ri + 1]._tr.get_or_add_trPr()
        trH = OxmlElement('w:trHeight')
        trH.set(qn('w:val'), '600'); trH.set(qn('w:hRule'), 'atLeast')
        trPr.append(trH)
    for row in t.rows:
        for ci in range(6):
            row.cells[ci].width = CW[ci]

# ═══ ENCOURAGEMENT ═══
ENCOURAGE = [
    (5,  ["🎉 你已完成总词汇量的5%，加油！好的开始是成功的一半！",
          "🚀 5%完成！每一步都在靠近目标！",
          "🌱 种子已经种下，5%的词汇正在生根发芽！"]),
    (10, ["💪 很厉害！你又进一步，记住了总词汇量的10%！",
          "📈 10%达成！你已经看到了进步的趋势！",
          "🎯 十分之一完成，九分之九还在路上！"]),
    (15, ["✨ 百尺竿头，更进一步，马上要完成15%了！",
          "🔥 15%！记住的单词正在悄然改变你的大脑！",
          "🏃 跑过了15%的路程，节奏已经起来了！"]),
    (20, ["🚀 已经完成20%了！词根记忆法正在发力！",
          "⭐ 五分之一的词汇已被征服！",
          "💎 20%！累计的词汇量开始产生化学反应！"]),
    (25, ["🌟 四分之一的词汇已经被你征服！",
          "🎊 25%！四分之一里程碑，值得庆祝！",
          "📚 一本词典的四分之一已经装进脑子里了！"]),
    (30, ["🔥 30%！记住的每个词都会在考试中帮你得分！",
          "⚡ 接近三分之一了！",
          "🎪 30%达成！越到后面词根串记越轻松！"]),
    (35, ["💎 超过三分之一了！你比昨天的自己更强大！",
          "🌈 35%！越来越多的单词变得面熟了！",
          "🏗 词汇大厦已经建了三成五！"]),
    (40, ["⚡ 40%的词汇量，你已经超过大多数人了！",
          "🎯 四成完成！节奏稳了！",
          "💰 每记一个词都是在为未来存钱！40%了！"]),
    (45, ["🎯 45%！距离半程只差一步之遥！",
          "🏔 45%！登顶前的最后冲刺！",
          "👀 45%完成，你已经开始用英语思维了！"]),
    (50, ["🎊 哇，加油啊，已经过半了！50%的词汇已拿下！",
          "🥇 半程冠军！",
          "🏁 50%！一半已过，剩下的都是下坡路！",
          "🎆 里程碑！所有单词你已认识一半！"]),
    (55, ["🌈 55%！下半程开始，冲刺吧！",
          "⛰ 55%！回头看，起点已经很远了！",
          "📊 55%完成率，你的坚持令人敬佩！"]),
    (60, ["🏆 60%！大多数词汇你都已经掌握了！",
          "🎯 六成完成！量变正在引起质变！",
          "🚀 60%！从这往后越来越顺！"]),
    (65, ["⭐ 65%！三度已过二，胜利在望！",
          "💪 接近七成！",
          "🌄 65%！山顶越来越近了！"]),
    (70, ["🚩 70%！你已经进入高阶词汇区了！",
          "👑 七成！你的英语水平上了个大台阶！",
          "🔥 70%！离精通只差30%！"]),
    (75, ["🎪 四分之三！剩下的都是小菜一碟！",
          "🏗 75%！词汇大厦即将封顶！",
          "🎯 四分之三完成，你已经是词汇高手了！"]),
    (80, ["👑 80%！你已经是一位词汇达人了！",
          "💎 八成！剩下的20%用联想记忆轻松拿下！",
          "🌟 80%！你已突破了英语词汇的关键门槛！"]),
    (85, ["💥 85%！离胜利只差最后一程！",
          "🏃 冲刺阶段！最后15%需要坚持！",
          "🎆 85%！回头看一路走来，所有的努力都值得！"]),
    (90, ["🎆 90%！最后的冲刺，冲鸭！",
          "🏁 九成完成！最后10%决定胜负！",
          "🥇 90%！你已站在词汇金字塔的顶端！"]),
    (95, ["🏁 95%！胜利就在眼前！",
          "🚀 最后5%，冲刺！",
          "🌟 95%！你是千分之五十的坚持者！"]),
    (100, ["🥇 🎉🎉🎉 100%！恭喜你完成了全部词汇的学习！",
           "🎊🏆🎊 100%完美通关！",
           "💎 100%！所有的汗水在这一刻绽放！"]),
]

def get_encourage(cumulative_words, group_words):
    total = cumulative_words + group_words
    pct = total * 100 / TOTAL_WORDS
    best = None
    for mil, msgs in ENCOURAGE:
        if pct >= mil:
            best = msgs
    return random.choice(best) if best else ""

# ═══════════════════════════════════════════
# 5. BUILD DOCUMENT
# ═══════════════════════════════════════════
print("Creating document...")

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(9)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for sec in doc.sections:
    sec.top_margin = Cm(0)
    sec.bottom_margin = Cm(0)
    sec.left_margin = Cm(0)
    sec.right_margin = Cm(0)

# ─── COVER ───
build_cover(doc, root_groups, basic_groups, all_words_parsed)

# ─── NEW SECTION (with margins + footer) ───
sec2 = doc.add_section()
sec2.top_margin = Cm(0.8)
sec2.bottom_margin = Cm(0.8)
sec2.left_margin = Cm(0.8)
sec2.right_margin = Cm(0.8)
add_footer(doc)

# ─── INSTRUCTIONS ───
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('📋 使用说明')
r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

instructions = [
    f'📖 Part 1: 词根词汇 {len(root_words)}词/{len(root_groups)}组 · Part 2: 基础词汇 {len(basic_words)}词/{len(basic_groups)}组',
    f'📅 共 {TOTAL_GROUPS}组 · {TOTAL_DAYS}天 —— 每天学1组新词 + 复习之前到期的小组',
    '📋 每组6轮：学习页 → 英→中 → 中→英 → 音标默写 → 高频词 → 混合终测',
    '🎯 每天翻到对应 DAY 页码，从上到下做完即可，无需自己算进度',
    '✍ 写作指数基于真实语料库频次 | 📄 页脚 2/XXX 页码',
]
for t in instructions:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(t)
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ─── STUDY PLAN INDEX ───
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('📋 学习计划索引')
r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('艾宾浩斯遗忘曲线 · 6轮循环复习计划')
r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Part 1 header
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(1)
r = p.add_run('🆕 新学组    🔄 复习组')
r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

# Generate index lines
for day_data in schedule:
    day = day_data['day']
    ng = day_data['new_group']
    revs = day_data['reviews']

    parts = []

    # New group indicator
    if ng:
        group_label_short_s = f"第{ng['seq']:02d}组"
        if ng['part'] == 1:
            parts.append(f"🆕 {group_label_short_s}(学习)")
        else:
            parts.append(f"🆕 {group_label_short_s}(学习)")

    # Review groups
    for rv in revs:
        gs = f"第{rv['seq']:02d}组"
        rev_info = f"🔄 {gs}({rv['round_name']})"
        parts.append(rev_info)

    line = f"DAY {day:3d}  " + "  |  ".join(parts)

    # Check for Part 1 -> Part 2 transition
    is_part2_start = (day == len(root_groups) + 1)

    if is_part2_start and day > 1:
        # Add Part 2 separator
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run('══════════ Part 2: 基础词汇 ══════════')
        r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(12)
    r = p.add_run(line)
    r.font.size = Pt(7.5)
    if ng is None:
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)  # review-only days dimmed
    else:
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ─── BODY: DAY-BY-DAY CONTENT ───
cumulative_words = 0

for day_data in schedule:
    day = day_data['day']
    ng = day_data['new_group']
    revs = day_data['reviews']

    doc.add_page_break()

    # ── Day Header ──
    # Count total pages for this day
    pages_count = 0
    if ng:
        pages_count += 1  # learning page
    for rv in revs:
        if rv['words']:
            pages_count += 1  # review page

    day_label = f"📅 DAY {day}"
    if ng:
        if ng['part'] == 1:
            day_label += f"  🆕 Part 1 · 第{ng['group_num']:02d}组（学习）"
        else:
            day_label += f"  🆕 Part 2 · 第{ng['group_num']:02d}组（学习）"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(day_label)
    r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Sub-header: review summary
    if revs:
        rev_summary = "  |  ".join([f"🔄 第{rv['seq']:02d}组({rv['round_name']})" for rv in revs])
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(rev_summary)
        r.font.size = Pt(7); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── New Group Learning ──
    if ng:
        group_wl = ng['words']
        if ng['part'] == 1:
            label_parts = f"Part 1 · 第{ng['group_num']:02d}组 · {len(group_wl)}词"
        else:
            label_parts = f"Part 2 · 第{ng['group_num']:02d}组 · {len(group_wl)}词"
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'📖 学习页  · {label_parts}')
        r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        for w in group_wl:
            card_learn(doc, w)

    # ── Review Tables ──
    for rv in revs:
        wl = rv['words']
        if not wl:
            continue  # Skip if no words (e.g. highfreq mode with no stars>=4)

        total_w = len(wl)
        part_tag = f"Part {rv['part']}"
        group_tag = f"第{rv['group_num']:02d}组"
        title_label = f"{rv['round_name']} · {part_tag} · {group_tag}"
        subtitle = f"DAY {day} · {rv['label']} · {total_w}词"
        review_title(doc, title_label, subtitle)
        name_date(doc, total_w)

        mode = rv['mode']
        sixcol(doc, wl, mode)

    # ── Encouragement ──
    if ng:
        cumulative_words += len(ng['words'])
        msg = get_encourage(cumulative_words - len(ng['words']), len(ng['words']))
        if msg:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(msg)
            r.font.size = Pt(9); r.font.bold = True
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    print(f"  Day {day:3d}/{TOTAL_DAYS}: {'🆕'+('G'+str(ng['seq']) if ng else '—'):>5s}  🔄{len(revs)} reviews")

# ═══ SAVE ═══
print("\nSaving...")
doc.save(OUT)
fsize = os.path.getsize(OUT) / 1024
print(f"\n✅ {OUT} ({fsize:.0f}KB)")
print(f"📊 {TOTAL_WORDS}词 | {TOTAL_GROUPS}组 | {TOTAL_DAYS}天")
print(f"🎯 Part 1: {len(root_groups)}组词根 | Part 2: {len(basic_groups)}组基础")
print(f"📅 按天编排 + 学习计划索引 + 鼓励语系统")
