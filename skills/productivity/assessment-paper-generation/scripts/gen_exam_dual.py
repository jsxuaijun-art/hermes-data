# -*- coding: utf-8 -*-
"""双版本试卷 docx 生成骨架（考生版无答案 + 教师版含答案）。
用同一份结构化试题数据渲染两份文档。填入你的题干/选项/答案/要点即可用。
运行：~/hermes-agent/venv/bin/python script.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, name='宋体', size=11, bold=False, color=None):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)  # 关键：中文字体
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_para(doc, text, size=11, bold=False, align=None, name='宋体',
             space_after=4, color=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, name=name, size=size, bold=bold, color=color)
    return p

# ====== 试题数据（换成你的内容）======
single = [  # (题干, [选项], 答案字母, 解析)
    ("1. 示例题：以下哪项正确？（    ）", ["A. ..", "B. ..", "C. ..", "D. .."], "B", "解析..."),
]
multiple = [  # (题干, [选项], [答案字母列表], 解析)
    ("2. 多选题（    ）", ["A. ..", "B. .."], ["A","B"], "解析..."),
]
truefalse = [  # (题干, 正确?, 解析)
    ("3. 判断题（    ）", True, "解析..."),
]
shortanswer = [  # (题干, [参考答案要点分点])
    ("4. 简答题", ["要点1", "要点2", "要点3"]),
]

def build(show_answer):
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
    title = "试卷标题"
    add_para(doc, title, size=18, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, name='黑体', space_after=4)
    add_para(doc, "满分100分　建议作答时间45分钟", size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "姓名：__________　岗位：__________　日期：__________",
             size=11, space_after=8)

    add_para(doc, "一、单选题（每题3分，共10题，30分）", size=12, bold=True, name='黑体')
    for q, opts, ans, ex in single:
        add_para(doc, q, size=11)
        for o in opts:
            add_para(doc, "　　" + o, size=11)
        if show_answer:
            add_para(doc, f"【答案】{ans}　{ex}", size=10, name='楷体',
                     color=RGBColor(0xC0, 0, 0))
        else:
            add_para(doc, "", size=6)

    add_para(doc, "二、多选题（每题4分，少选/多选/错选均不得分）", size=12, bold=True, name='黑体')
    for q, opts, ans, ex in multiple:
        add_para(doc, q, size=11)
        for o in opts:
            add_para(doc, "　　" + o, size=11)
        if show_answer:
            add_para(doc, f"【答案】{''.join(ans)}　{ex}", size=10, name='楷体',
                     color=RGBColor(0xC0, 0, 0))
        else:
            add_para(doc, "", size=6)

    add_para(doc, "三、判断题（每题2.5分，正确的打√，错误的打×）", size=12, bold=True, name='黑体')
    for q, ans, ex in truefalse:
        add_para(doc, q, size=11)
        if show_answer:
            add_para(doc, f"【答案】{'√' if ans else '×'}　{ex}", size=10, name='楷体',
                     color=RGBColor(0xC0, 0, 0))
        else:
            add_para(doc, "", size=6)

    add_para(doc, "四、简答题（每题7.5分，共30分）", size=12, bold=True, name='黑体')
    for q, points in shortanswer:
        add_para(doc, q, size=11)
        if show_answer:
            add_para(doc, "【参考答案要点】", size=10, name='楷体', bold=True,
                     color=RGBColor(0xC0, 0, 0))
            for k, pt in enumerate(points, 1):
                add_para(doc, f"　　{k}. {pt}", size=10, name='楷体',
                         color=RGBColor(0xC0, 0, 0))
        else:
            for _ in range(3):
                add_para(doc, "", size=11, space_after=10)
    return doc

teacher = build(show_answer=True)
teacher.save('/tmp/试卷_教师版含答案.docx')
student = build(show_answer=False)
student.save('/tmp/试卷_考生版.docx')
print("done: 教师版 + 考生版")
