"""
财务分析报告自动生成模板 — 通用版
基于苏州盈信企业管理16年实战经验

使用方法：
1. 修改 data = {} 中的财务数据（从科目余额表提取）
2. 修改 benchmark = {} 中的行业基准值（需搜索目标行业）
3. 修改公司名称、期间等基本信息
4. 运行: python3 /tmp/财务分析报告生成脚本.py
5. 输出：Word(.docx) + Markdown(.md) 到桌面
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os, json

# ═══════════════════════════════════════════
# ★★★ 用户修改区 START ★★★
# ═══════════════════════════════════════════

COMPANY = 'XX科技有限公司'           # 公司名称
PERIOD_1 = '2025年'                  # 基期
PERIOD_2 = '2026年'                  # 本期
MONTHS = 4                           # 本期涵盖月数（如1-4月=4）
COMPARE_LABEL = '2025年4月 vs 2026年4月'  # 对比标题

# 数据来源：从科目余额表提取的财务数据
data = {
    # ── 利润表（累计数据） ──
    'revenue_1': 0.0,       # 基期营业收入
    'revenue_2': 0.0,       # 本期营业收入
    'cost_1': 0.0,           # 基期营业成本
    'cost_2': 0.0,           # 本期营业成本
    'tax_surcharge_1': 0.0, # 基期税金及附加
    'tax_surcharge_2': 0.0,
    'sales_exp_1': 0.0,     # 基期销售费用
    'sales_exp_2': 0.0,
    'admin_exp_1': 0.0,     # 基期管理费用
    'admin_exp_2': 0.0,
    'rd_exp_1': 0.0,        # 基期研究费用（管理费用中）
    'rd_exp_2': 0.0,
    'finance_exp_1': 0.0,   # 基期财务费用
    'finance_exp_2': 0.0,
    'net_profit_1': 0.0,    # 基期净利润
    'net_profit_2': 0.0,

    # ── 资产负债表（期末数） ──
    'cash_1': 0.0,           # 基期货币资金
    'cash_2': 0.0,
    'ar_1': 0.0,             # 基期应收账款
    'ar_2': 0.0,
    'inventory_1': 0.0,      # 基期存货
    'inventory_2': 0.0,
    'current_assets_1': 0.0, # 基期流动资产合计
    'current_assets_2': 0.0,
    'fixed_assets_net_1': 0.0, # 基期固定资产净值
    'fixed_assets_net_2': 0.0,
    'total_assets_1': 0.0,   # 基期资产总计
    'total_assets_2': 0.0,
    'current_liab_1': 0.0,   # 基期流动负债合计
    'current_liab_2': 0.0,
    'total_liab_1': 0.0,     # 基期负债合计
    'total_liab_2': 0.0,
    'equity_1': 0.0,         # 基期所有者权益
    'equity_2': 0.0,
    'st_loan_1': 0.0,        # 基期短期借款
    'st_loan_2': 0.0,
    'ap_1': 0.0,             # 基期应付账款
    'ap_2': 0.0,
    'owner_loan_1': 0.0,     # 基期个人/股东借款（其他应付款）
    'owner_loan_2': 0.0,
}

# 行业基准值（需搜索目标行业）
benchmark = {
    'gross_margin': ('XX%–XX%', '请填写目标行业'),
    'net_margin': ('XX%–XX%', '请填写目标行业'),
    'rd_ratio': ('4%–8%', '高新技术企业认定≥5%（营收<5亿）'),
    'current_ratio': ('1.5–2.5', '请填写目标行业'),
    'debt_ratio': ('35%–55%', '请填写目标行业'),
    'ar_turnover': ('X–X次/年', '请填写目标行业'),
    'inventory_turnover': ('X–X次/年', '请填写目标行业'),
    'revenue_growth': ('X%–X%', '请填写目标行业'),
}

# 高风险自动发现规则阈值（可按行业调整）
THRESHOLDS = {
    'ar_concentration': 0.70,       # 单一客户应收占比>70% → 高风险
    'ar_concentration_medium': 0.40, # 单一客户应收占比>40% → 中风险
    'rd_ratio_high': 0.05,          # 研发费用占比<5%且收入<5000万 → 风险
    'rd_ratio_medium': 0.04,        # 研发费用占比<4%且收入>5000万 → 风险
    'inventory_growth_risk': 1.5,   # 存货增速/收入增速 > 1.5倍 → 积压风险
    'tax_rate_alert': 0.05,         # 实际税率<5% → 需确认
    'bs_diff_alert': 0.01,          # 资产负债表差额>0.01 → 不平
}

# ═══════════════════════════════════════════
# ★★★ 用户修改区 END ★★★
# ═══════════════════════════════════════════

# ── 派生值：收入拆分（如有明细数据可直接覆盖） ──
REVENUE_MAIN_1 = data['revenue_1']    # 主营业务收入（基期）
REVENUE_MAIN_2 = data['revenue_2']    # 主营业务收入（本期）
REVENUE_OTHER_1 = 0.0                 # 其他业务收入（基期）
REVENUE_OTHER_2 = 0.0                 # 其他业务收入（本期）

# ═══════════════════════════════════════════
# 自动发现规则引擎
# ═══════════════════════════════════════════

def run_auto_discovery(data):
    """基于数据的自动风险发现，返回风险列表"""
    risks = []
    
    # 规则1 & 2: 资产负债表检查
    bs_diff = data['total_assets_2'] - (data['total_liab_2'] + data['equity_2'])
    if abs(bs_diff) > max(THRESHOLDS['bs_diff_alert'], 1):
        risks.append(('🔴 高', '资产负债表不平',
            f'资产总计{data["total_assets_2"]:,.2f} ≠ 负债+权益{data["total_liab_2"]+data["equity_2"]:,.2f}，差额{bs_diff:,.2f}（占资产{abs(bs_diff)/data["total_assets_2"]*100:.1f}%）。需财务人员核实。'))
    
    # 规则3: 个人借款
    if data.get('owner_loan_1', 0) and data['owner_loan_1'] > 0:
        risks.append(('🟡 中', '个人/股东借款存在（基期）',
            f'基期个人借款{data["owner_loan_1"]:,.2f}元。如本期末已归还则风险解除，但需检查资金流水凭证。如长期挂账涉及个人所得税视同分红风险。'))
    
    if data.get('owner_loan_2', 0) and data['owner_loan_2'] > 0:
        risks.append(('🟡 中', '个人/股东借款未清（本期）',
            f'本期仍有个人借款{data["owner_loan_2"]:,.2f}元挂账。需关注个人所得税视同分红风险。'))
    
    # 规则4 & 5: 应收集中度（需配合科目余额表客户明细）
    # 在脚本外补充：单客户占比
    
    # 规则7: 零所得税风险
    net_profit = data.get('net_profit_2', 0)
    # 如果没有所得税数据，这里由分析人在撰写时补充
    
    # 规则9 & 10: 研发费用占比
    rd_ratio = data['rd_exp_2'] / data['revenue_2'] if data['revenue_2'] > 0 else 0
    revenue_annual = data['revenue_2'] / MONTHS * 12
    if revenue_annual < 50000000:  # 年营收<5,000万
        if rd_ratio < THRESHOLDS['rd_ratio_high'] and data['rd_exp_2'] > 0:
            risks.append(('🟡 中', '研发费用占比低于高企认定标准',
                f'本期研发费用占比{rd_ratio*100:.2f}%，低于高新技术企业认定≥5%的标准（年营收<5,000万）。如需申请或维护高企资质需补足研发投入。'))
    else:
        if rd_ratio < THRESHOLDS['rd_ratio_medium'] and data['rd_exp_2'] > 0:
            risks.append(('🟡 中', '研发费用占比值得关注',
                f'本期研发费用占比{rd_ratio*100:.2f}%，接近高新技术企业认定≥4%的标准（年营收>5,000万）。需确保达标。'))
    
    # 规则6: 存货增速对比
    if data.get('inventory_1', 0) and data['inventory_1'] > 0:
        inv_growth = (data['inventory_2'] - data['inventory_1']) / data['inventory_1']
        rev_growth_calc = (data['revenue_2'] - data['revenue_1']) / data['revenue_1'] if data['revenue_1'] > 0 else 0
        if inv_growth > rev_growth_calc * THRESHOLDS['inventory_growth_risk'] and rev_growth_calc > 0:
            risks.append(('🟡 中', '存货增长快于收入增长',
                f'存货增长{inv_growth*100:.1f}%（{data["inventory_1"]:,.0f}→{data["inventory_2"]:,.0f}），远超收入增长{rev_growth_calc*100:.1f}%。需关注是否存在积压风险。'))
    
    # 外汇风险：如财务费用中有大额汇兑损益
    finance_growth = data['finance_exp_2'] - data['finance_exp_1']
    if abs(finance_growth) > 50000 and data['revenue_2'] > 0:
        risks.append(('🟡 中', '财务费用大幅波动',
            f'财务费用从{data["finance_exp_1"]:,.2f}变为{data["finance_exp_2"]:,.2f}，变动额{finance_growth:,.2f}。如为出口型企业，需关注外汇风险。'))
    
    return risks

# ═══════════════════════════════════════════
# Word文档生成
# ═══════════════════════════════════════════

def build_doc():
    doc = Document()
    
    # 全局样式
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5
    
    def add_h(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        return h
    
    def add_tbl(headers, rows):
        table = doc.add_table(rows=1+len(rows), cols=len(headers))
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = table.rows[ri+1].cells[ci]
                cell.text = str(val)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.size = Pt(9)
        return table
    
    def kv(text, bold_part=None):
        p = doc.add_paragraph()
        if bold_part:
            r = p.add_run(bold_part); r.bold = True
            r.font.name = '微软雅黑'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r = p.add_run(text)
        r.font.name = '微软雅黑'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def note(text):
        p = doc.add_paragraph()
        r = p.add_run(f'⚠ {text}')
        r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
        r.font.name = '微软雅黑'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def red(text):
        p = doc.add_paragraph()
        r = p.add_run(f'🚨 {text}')
        r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00); r.bold = True
        r.font.name = '微软雅黑'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def green(text):
        r = doc.add_paragraph().add_run(text)
        r.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        r.font.name = '微软雅黑'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def fmt(v):
        if v is None or v == '': return '-'
        if isinstance(v, str): return v
        return f'{v:,.2f}' if abs(v) >= 10000 else f'{v:,.2f}'
    
    def pct(v): return f'{v*100:.1f}%'
    
    # ═══ 计算指标 ═══
    gm_1 = (data['revenue_1']-data['cost_1'])/data['revenue_1'] if data['revenue_1'] else 0
    gm_2 = (data['revenue_2']-data['cost_2'])/data['revenue_2'] if data['revenue_2'] else 0
    nm_1 = data['net_profit_1']/data['revenue_1'] if data['revenue_1'] else 0
    nm_2 = data['net_profit_2']/data['revenue_2'] if data['revenue_2'] else 0
    rd_1 = data['rd_exp_1']/data['revenue_1'] if data['revenue_1'] else 0
    rd_2 = data['rd_exp_2']/data['revenue_2'] if data['revenue_2'] else 0
    cr_1 = data['current_assets_1']/data['current_liab_1'] if data['current_liab_1'] else 0
    cr_2 = data['current_assets_2']/data['current_liab_2'] if data['current_liab_2'] else 0
    dr_1 = data['total_liab_1']/data['total_assets_1'] if data['total_assets_1'] else 0
    dr_2 = data['total_liab_2']/data['total_assets_2'] if data['total_assets_2'] else 0
    ar_t_1 = data['revenue_1']/data['ar_1']*(12/MONTHS) if data['ar_1'] else 0
    ar_t_2 = data['revenue_2']/data['ar_2']*(12/MONTHS) if data['ar_2'] else 0
    inv_t_1 = data['cost_1']/data['inventory_1']*(12/MONTHS) if data['inventory_1'] else 0
    inv_t_2 = data['cost_2']/data['inventory_2']*(12/MONTHS) if data['inventory_2'] else 0
    rev_g = (data['revenue_2']-data['revenue_1'])/data['revenue_1'] if data['revenue_1'] else 0
    prof_g = (data['net_profit_2']-data['net_profit_1'])/data['net_profit_1'] if data['net_profit_1'] else 0
    
    risks = run_auto_discovery(data)
    
    # ── 封面 ──
    title = doc.add_heading('', level=0)
    r = title.add_run(f'{COMPANY}\n财务分析报告')
    r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    r.font.name = '微软雅黑'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f'{COMPARE_LABEL}\n数据来源：苏州盈信企业管理有限公司')
    r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()
    
    # ── 一、执行概要 ──
    add_h('一、执行概要', 1)
    
    overall = '良好'
    if any('🔴' in r[0] for r in risks): overall = '需重点关注'
    elif any('🟡' in r[0] for r in risks): overall = '基本正常，部分需关注'
    
    kv(f'总体评价：{COMPANY}本期营收同比增长{pct(rev_g)}，净利润增长{pct(prof_g)}，毛利率{pct(gm_2)}。{overall}。')
    
    note_lines = [f'营业收入{fmt(data["revenue_2"])}，净利润{fmt(data["net_profit_2"])}']
    for level, title, desc in risks[:3]:
        note_lines.append(f'{level}: {title}')
    note('；'.join(note_lines[:3]))
    
    # 核心指标表
    add_h('核心指标对比', 2)
    add_tbl(
        ['核心指标', f'{PERIOD_1}', f'{PERIOD_2}', '变化', '行业基准', '评价'],
        [
            ['营业收入', fmt(data['revenue_1']), fmt(data['revenue_2']), pct(rev_g),
             benchmark['revenue_growth'][0], '🟢' if rev_g > 0.15 else '🟡'],
            ['毛利率', pct(gm_1), pct(gm_2), f'{pct(gm_2-gm_1)}',
             benchmark['gross_margin'][0], '🟢' if gm_2 > 0.35 else '🟡'],
            ['净利润率', pct(nm_1), pct(nm_2), f'{pct(nm_2-nm_1)}',
             benchmark['net_margin'][0], '🟢' if nm_2 > 0.10 else '🟡'],
            ['研发费用占比', pct(rd_1), pct(rd_2), f'{pct(rd_2-rd_1)}',
             benchmark['rd_ratio'][0], '🟡' if rd_2 < 0.05 else '🟢'],
            ['流动比率', f'{cr_1:.2f}', f'{cr_2:.2f}', f'{cr_2-cr_1:+.2f}',
             benchmark['current_ratio'][0], '🟢' if cr_2 > 1.5 else '🟡'],
            ['资产负债率', pct(dr_1), pct(dr_2), f'{pct(dr_2-dr_1)}',
             benchmark['debt_ratio'][0], '🟢' if dr_2 < 0.55 else '🟡'],
        ]
    )
    doc.add_page_break()
    
    # ── 二、经营业绩分析 ──
    add_h('二、经营业绩分析（利润表）', 1)
    
    add_h('2.1 收入情况', 2)
    add_tbl(
        ['项目', f'{PERIOD_1}', f'{PERIOD_2}', '变化额', '增幅'],
        [
            ['营业收入合计', fmt(data['revenue_1']), fmt(data['revenue_2']),
             fmt(data['revenue_2']-data['revenue_1']), pct(rev_g)],
        ]
    )
    
    add_h('2.2 毛利率分析', 2)
    add_tbl(
        ['项目', f'{PERIOD_1}', f'{PERIOD_2}', '变化'],
        [
            ['营业收入', fmt(data['revenue_1']), fmt(data['revenue_2']), f'+{fmt(data["revenue_2"]-data["revenue_1"])}'],
            ['营业成本', fmt(data['cost_1']), fmt(data['cost_2']), f'+{fmt(data["cost_2"]-data["cost_1"])}'],
            ['毛利额', fmt(data['revenue_1']-data['cost_1']), fmt(data['revenue_2']-data['cost_2']),
             f'+{fmt(data["revenue_2"]-data["cost_2"]-data["revenue_1"]+data["cost_1"])}'],
            ['毛利率', pct(gm_1), pct(gm_2), f'{pct(gm_2-gm_1)}'],
        ]
    )
    
    add_h('2.3 期间费用分析', 2)
    add_tbl(
        ['项目', f'{PERIOD_1}', '占比', f'{PERIOD_2}', '占比', '费用率变化'],
        [
            ['销售费用', fmt(data['sales_exp_1']), pct(data['sales_exp_1']/data['revenue_1']),
             fmt(data['sales_exp_2']), pct(data['sales_exp_2']/data['revenue_2']),
             pct(data['sales_exp_2']/data['revenue_2']-data['sales_exp_1']/data['revenue_1'])],
            ['管理费用', fmt(data['admin_exp_1']), pct(data['admin_exp_1']/data['revenue_1']),
             fmt(data['admin_exp_2']), pct(data['admin_exp_2']/data['revenue_2']),
             pct(data['admin_exp_2']/data['revenue_2']-data['admin_exp_1']/data['revenue_1'])],
            ['—其中：研究费用', fmt(data['rd_exp_1']), pct(rd_1),
             fmt(data['rd_exp_2']), pct(rd_2),
             pct(rd_2-rd_1)],
            ['财务费用', fmt(data['finance_exp_1']), pct(data['finance_exp_1']/data['revenue_1']),
             fmt(data['finance_exp_2']), pct(data['finance_exp_2']/data['revenue_2']),
             pct(data['finance_exp_2']/data['revenue_2']-data['finance_exp_1']/data['revenue_1'])],
        ]
    )
    
    if abs(data['finance_exp_2'] - data['finance_exp_1']) > 50000:
        red(f'财务费用显著变化：从{fmt(data["finance_exp_1"])}至{fmt(data["finance_exp_2"])}，变动额{fmt(data["finance_exp_2"]-data["finance_exp_1"])}。请结合外汇损益和利息费用具体分析。')
    
    add_h('2.4 净利润分析', 2)
    add_tbl(
        ['项目', f'{PERIOD_1}', f'{PERIOD_2}', '变化'],
        [
            ['净利润', fmt(data['net_profit_1']), fmt(data['net_profit_2']), f'+{fmt(data["net_profit_2"]-data["net_profit_1"])}'],
            ['净利率', pct(nm_1), pct(nm_2), f'+{pct(nm_2-nm_1)}'],
        ]
    )
    doc.add_page_break()
    
    # ── 三、资产结构与偿债能力 ──
    add_h('三、资产结构与偿债能力（资产负债表）', 1)
    
    add_h('3.1 资产结构', 2)
    add_tbl(
        ['项目', f'{PERIOD_1}', '占总资产%', f'{PERIOD_2}', '占总资产%', '变化'],
        [
            ['货币资金', fmt(data['cash_1']),
             f'{data["cash_1"]/data["total_assets_1"]*100:.1f}%' if data['total_assets_1'] else '-',
             fmt(data['cash_2']),
             f'{data["cash_2"]/data["total_assets_2"]*100:.1f}%' if data['total_assets_2'] else '-',
             fmt(data['cash_2']-data['cash_1'])],
            ['应收账款', fmt(data['ar_1']),
             f'{data["ar_1"]/data["total_assets_1"]*100:.1f}%',
             fmt(data['ar_2']),
             f'{data["ar_2"]/data["total_assets_2"]*100:.1f}%',
             fmt(data['ar_2']-data['ar_1'])],
            ['存货', fmt(data['inventory_1']),
             f'{data["inventory_1"]/data["total_assets_1"]*100:.1f}%',
             fmt(data['inventory_2']),
             f'{data["inventory_2"]/data["total_assets_2"]*100:.1f}%',
             fmt(data['inventory_2']-data['inventory_1'])],
            ['流动资产合计', fmt(data['current_assets_1']),
             f'{data["current_assets_1"]/data["total_assets_1"]*100:.1f}%',
             fmt(data['current_assets_2']),
             f'{data["current_assets_2"]/data["total_assets_2"]*100:.1f}%',
             fmt(data['current_assets_2']-data['current_assets_1'])],
            ['固定资产净值', fmt(data['fixed_assets_net_1']),
             f'{data["fixed_assets_net_1"]/data["total_assets_1"]*100:.1f}%',
             fmt(data['fixed_assets_net_2']),
             f'{data["fixed_assets_net_2"]/data["total_assets_2"]*100:.1f}%',
             fmt(data['fixed_assets_net_2']-data['fixed_assets_net_1'])],
            ['资产总计', fmt(data['total_assets_1']), '100%',
             fmt(data['total_assets_2']), '100%',
             fmt(data['total_assets_2']-data['total_assets_1'])],
        ]
    )
    
    add_h('3.2 负债结构', 2)
    add_tbl(
        ['项目', f'{PERIOD_1}', f'{PERIOD_2}', '变化'],
        [
            ['短期借款', fmt(data['st_loan_1']), fmt(data['st_loan_2']), fmt(data['st_loan_2']-data['st_loan_1'])],
            ['应付账款', fmt(data['ap_1']), fmt(data['ap_2']), fmt(data['ap_2']-data['ap_1'])],
            ['流动负债合计', fmt(data['current_liab_1']), fmt(data['current_liab_2']), fmt(data['current_liab_2']-data['current_liab_1'])],
        ]
    )
    
    add_h('3.3 偿债能力指标', 2)
    add_tbl(
        ['指标', f'{PERIOD_1}', f'{PERIOD_2}', '行业基准', '评价'],
        [
            ['流动比率', f'{cr_1:.2f}', f'{cr_2:.2f}', benchmark['current_ratio'][0],
             '🟢充裕' if cr_2 > 2 else ('🟡正常' if cr_2 > 1.2 else '🔴紧张')],
            ['资产负债率', pct(dr_1), pct(dr_2), benchmark['debt_ratio'][0],
             '🟢低杠杆' if dr_2 < 0.35 else ('🟡正常' if dr_2 < 0.55 else '🔴偏高')],
        ]
    )
    
    # 自动发现的资产负债表问题
    for level, title, desc in risks:
        if '资产负债表' in title or '个人借款' in title:
            if '🔴' in level: red(desc)
            else: note(desc)
    
    doc.add_page_break()
    
    # ── 四、运营效率 ──
    add_h('四、运营效率分析', 1)
    add_tbl(
        ['运营指标', f'{PERIOD_1}（年化）', f'{PERIOD_2}（年化）', '行业基准', '评价'],
        [
            ['应收账款周转率', f'{ar_t_1:.2f}次/年', f'{ar_t_2:.2f}次/年',
             benchmark['ar_turnover'][0], '🟢良好' if ar_t_2 > 4 else '🟡需关注'],
            ['应收账款周转天数', f'{365/ar_t_1:.0f}天' if ar_t_1 else '-',
             f'{365/ar_t_2:.0f}天' if ar_t_2 else '-', benchmark['ar_turnover'][0], '—'],
            ['存货周转率', f'{inv_t_1:.2f}次/年', f'{inv_t_2:.2f}次/年',
             benchmark['inventory_turnover'][0], '🟢正常'],
            ['存货周转天数', f'{365/inv_t_1:.0f}天' if inv_t_1 else '-',
             f'{365/inv_t_2:.0f}天' if inv_t_2 else '-', benchmark['inventory_turnover'][0], '—'],
        ]
    )
    doc.add_page_break()
    
    # ── 五、税务合规性提示 ──
    add_h('五、税务合规性提示', 1)
    kv('（需根据科目余额表中的应交税费明细补充分析）')
    kv('• 增值税：检查进销项匹配情况。出口企业关注出口退税申报。')
    kv('• 所得税：核查实际税率是否异常偏低。净利润>0但所得税=0需重点关注。')
    kv('• 其他税种：印花税、房产税、土地使用税是否足额申报。')
    kv('• 关联交易：个人借款、股东往来款涉及个人所得税视同分红风险。')
    
    doc.add_page_break()
    
    # ── 六、高企资质维护 ──
    add_h('六、高新技术企业资质维护分析', 1)
    add_tbl(
        ['高企认定条件', '要求标准', f'{COMPANY}现状', '是否达标'],
        [
            ['研发费用占比', '营收<5,000万：≥5%\n5,000万~2亿：≥4%\n>2亿：≥3%',
             f'本期：{pct(rd_2)}', '🟡需关注' if rd_2 < 0.05 else '🟢达标'],
            ['高新产品收入占比', '≥60%', '需确认', '❓待确认'],
            ['科技人员占比', '≥10%', '需确认', '❓待确认'],
            ['知识产权', '核心自主知识产权', '需确认', '❓待确认'],
        ]
    )
    
    if rd_2 < 0.05:
        red(f'研发费用占比{pct(rd_2)}，低于高企认定标准。建议2026年下半年增加研发投入。')
    
    doc.add_page_break()
    
    # ── 七、风险警示 ──
    add_h('七、风险警示', 1)
    if not risks:
        kv('未触发自动发现规则。但需结合客户明细、人员数据等补充分析。')
    else:
        for level, title, desc in risks:
            p = doc.add_paragraph()
            r = p.add_run(f'{level} {title}')
            r.bold = True; r.font.size = Pt(11)
            r.font.name = '微软雅黑'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            kv(desc)
    
    doc.add_page_break()
    
    # ── 八、管理建议 ──
    add_h('八、管理建议', 1)
    add_h('紧急（本周内）', 2)
    for level, title, desc in risks:
        if '🔴' in level:
            kv(f'✅ {title}：{desc[:80]}...')
    if not any('🔴' in r[0] for r in risks):
        kv('✅ 无紧急事项。')
    
    add_h('短期（1-3个月）', 2)
    kv('✅ 评估外汇风险管理方案（如适用）')
    kv('✅ 检查主要客户信用状况，制定催收计划')
    kv('✅ 确保研发费用占比达标')
    kv('✅ 与银行沟通短期借款续贷方案')
    kv('✅ 核实所得税申报准确性')
    
    add_h('中长期（3-12个月）', 2)
    kv('✅ 推动客户多元化，降低集中度')
    kv('✅ 申请/维护高新技术企业资质')
    kv('✅ 建立定期财务报告和经营分析会议制度')
    kv('✅ 评估数字化转型/ERP系统升级需求')
    kv('✅ 关注出口退税申报的及时性和准确性')
    
    # 结尾
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('— 报告完 —')
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    r.font.name = '微软雅黑'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('本报告由Hermes AI助手自动生成，数据来源为公司财务报表及科目余额表。\n建议客户结合实际情况核实后使用。')
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    return doc

# ═══════════════════════════════════════════
# Markdown版本生成
# ═══════════════════════════════════════════

def build_md():
    gm_1 = (data['revenue_1']-data['cost_1'])/data['revenue_1'] if data['revenue_1'] else 0
    gm_2 = (data['revenue_2']-data['cost_2'])/data['revenue_2'] if data['revenue_2'] else 0
    nm_1 = data['net_profit_1']/data['revenue_1'] if data['revenue_1'] else 0
    nm_2 = data['net_profit_2']/data['revenue_2'] if data['revenue_2'] else 0
    rd_2 = data['rd_exp_2']/data['revenue_2'] if data['revenue_2'] else 0
    rev_g = (data['revenue_2']-data['revenue_1'])/data['revenue_1'] if data['revenue_1'] else 0
    prof_g = (data['net_profit_2']-data['net_profit_1'])/data['net_profit_1'] if data['net_profit_1'] else 0
    
    md = f"""# {COMPANY} — 财务分析报告

**{COMPARE_LABEL}**
数据来源：苏州盈信企业管理有限公司

---

## 一、执行概要

**总体评价：** 营收同比增长{rev_g*100:.1f}%，净利润增长{prof_g*100:.1f}%，毛利率{gm_2*100:.1f}%。

| 核心指标 | {PERIOD_1} | {PERIOD_2} | 变化 | 行业基准 |
|---|---|---|---|---|
| 营业收入 | {data['revenue_1']:,.2f} | {data['revenue_2']:,.2f} | +{rev_g*100:.1f}% | {benchmark['revenue_growth'][0]} |
| 毛利率 | {gm_1*100:.1f}% | {gm_2*100:.1f}% | +{(gm_2-gm_1)*100:.1f}% | {benchmark['gross_margin'][0]} |
| 净利润率 | {nm_1*100:.1f}% | {nm_2*100:.1f}% | +{(nm_2-nm_1)*100:.1f}% | {benchmark['net_margin'][0]} |
| 研发费用占比 | {data['rd_exp_1']/data['revenue_1']*100:.1f}% | {rd_2*100:.1f}% | +{(rd_2-data['rd_exp_1']/data['revenue_1'])*100:.1f}% | {benchmark['rd_ratio'][0]} |

---

## 二、经营业绩分析（利润表）

| 项目 | {PERIOD_1} | {PERIOD_2} | 变化 |
|---|---|---|---|
| 营业收入 | {data['revenue_1']:,.2f} | {data['revenue_2']:,.2f} | +{data['revenue_2']-data['revenue_1']:,.2f} |
| 营业成本 | {data['cost_1']:,.2f} | {data['cost_2']:,.2f} | +{data['cost_2']-data['cost_1']:,.2f} |
| 毛利率 | {gm_1*100:.1f}% | {gm_2*100:.1f}% | +{(gm_2-gm_1)*100:.1f}% |
| 净利润 | {data['net_profit_1']:,.2f} | {data['net_profit_2']:,.2f} | +{data['net_profit_2']-data['net_profit_1']:,.2f} |
| 净利率 | {nm_1*100:.1f}% | {nm_2*100:.1f}% | +{(nm_2-nm_1)*100:.1f}% |

---

## 七、风险警示

{chr(10).join(f"- {'🔴 高' if '🔴' in r[0] else '🟡 中' if '🟡' in r[0] else '🟢 低'} **{r[1]}**: {r[2][:100]}..." for r in run_auto_discovery(data)) if run_auto_discovery(data) else '- 未触发自动发现规则，需结合客户明细补充分析。'}

---

*— 报告完 —*

*本报告由Hermes AI助手自动生成，数据来源为公司财务报表及科目余额表。*
"""
    return md

# ═══════════════════════════════════════════
# 生成与保存
# ═══════════════════════════════════════════

if __name__ == '__main__':
    # 生成Word
    doc = build_doc()
    base_dir = '/mnt/d/360MoveData/Users/Admin/Desktop/报表分析/'
    os.makedirs(base_dir, exist_ok=True)
    
    docx_path = os.path.join(base_dir, f'{COMPANY.split("有限公司")[0]}财务分析报告_{PERIOD_1[:4]}vs{PERIOD_2[:4]}.docx')
    doc.save(docx_path)
    print(f'✅ Word: {docx_path}')
    
    # 生成Markdown
    md_content = build_md()
    md_path = os.path.join(base_dir, f'{COMPANY.split("有限公司")[0]}财务分析报告_{PERIOD_1[:4]}vs{PERIOD_2[:4]}.md')
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_content)
    print(f'✅ Markdown: {md_path}')
    
    print(f'\n⚠ 别忘了修改 data = {{}} 中的财务数据！')
