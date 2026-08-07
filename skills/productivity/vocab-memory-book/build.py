# -*- coding: utf-8 -*-
"""
高考英语词汇默写本 · 50词 v9
- DAY X 标题行 + 六栏复习（序号缩窄，去✓□✗□）
- 写作指数来自真实英文语料库频次数据
- Day 31 → Day 30
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import random, math

DOCX_PATH = r"/mnt/c/Users/Admin/Desktop/高考英语词汇默写本_50词实验版.docx"

# ═══ EMBEDDED WORD DATA (definitions from 高考英语词汇表3817词.txt) ═══
WORD_DEFS = {
    "abstract":("ˈæbstrækt","抽象的（作品）；摘要"),"attract":("əˈtrækt","吸引，引起"),
    "attractive":("əˈtræktɪv","迷人的，有吸引力的"),"subtraction":("səbˈtrækʃən","减法"),
    "express":("ɪkˈspres","表达；表示；表情"),"expression":("ɪkˈspreʃən","表达；词句；表情"),
    "impress":("ɪmˈpres","留下极深的印象"),"impression":("ɪmˈpreʃən","印象，感觉"),
    "import":("ɪmˈpɔːt","进口，输入"),"portable":("ˈpɔːtəbəl","手提的，便携式的"),
    "report":("rɪˈpɔːt","报道，报告"),"support":("səˈpɔːt","支持，赞助"),
    "transport":("trænsˈpɔːt","运输"),"differ":("ˈdɪfə","相异，有区别"),
    "different":("ˈdɪfrənt","不同的，有差异的"),"offer":("ˈɒfə","提供；建议"),
    "prefer":("prɪˈfɜː","宁愿选择，更喜欢"),"refer":("rɪˈfɜː","谈到，提到，涉及"),
    "access":("ˈækses","通道，入径；存取"),"necessary":("ˈnesəsəri","必需的，必要的"),
    "process":("ˈprəʊses","过程，加工，处理"),"success":("səkˈses","成功"),
    "aspect":("ˈæspekt","方面，外观，外表"),"inspect":("ɪnˈspekt","检查；检验；审视"),
    "respect":("rɪˈspekt","尊敬，尊重"),"suspect":("səˈspekt","犯罪嫌疑人"),
    "construct":("kənˈstrʌkt","构筑；建造，建设"),"construction":("kənˈstrʌkʃən","建造，建设，建筑物"),
    "instruct":("ɪnˈstrʌkt","通知；指示；教"),"instruction":("ɪnˈstrʌkʃən","说明，须知；教导"),
    "abrupt":("əˈbrʌpt","突然的，意外的"),"corrupt":("kəˈrʌpt","贪污的，腐败的"),
    "erupt":("ɪˈrʌpt","（火山）爆发，喷发"),"interrupt":("ɪntəˈrʌpt","打扰，打断"),
    "admit":("ədˈmɪt","承认，准许进入"),"commit":("kəˈmɪt","犯（罪，错）；承诺"),
    "permit":("pəˈmɪt","许可，允许；执照"),"submit":("səbˈmɪt","提交，呈递"),
    "dismiss":("dɪsˈmɪs","让…离开；遣散；解雇"),"contradict":("kɒntrəˈdɪkt","反驳，驳斥，批驳"),
    "dictionary":("ˈdɪkʃənəri","词典，字典"),"predict":("prɪˈdɪkt","预言，预告，预报"),
    "assist":("əˈsɪst","帮助，协助"),"consist":("kənˈsɪst","包含，组成，构成"),
    "insist":("ɪnˈsɪst","坚持；坚决认为"),"resist":("rɪˈzɪst","抵抗；挡开"),
    "depend":("dɪˈpend","依靠，依赖，指望"),"independent":("ɪndɪˈpendənt","独立的，有主见的"),
    "expense":("ɪkˈspens","消费；支出"),"expensive":("ɪkˈspensɪv","昂贵的"),
}

# ═══ REAL FREQUENCY DATA (from English FrequencyWords corpus) ═══
# Values = frequency count; higher = more common
# Star ratings based on corpus frequency bands:
FREQ_STARS = {
    "abstract":1, "attract":2, "attractive":4, "subtraction":1,
    "express":4, "expression":3, "impress":3, "impression":4,
    "import":1, "portable":2, "report":5, "support":4, "transport":3,
    "differ":1, "different":5, "offer":5, "prefer":4, "refer":2,
    "access":4, "necessary":4, "process":4, "success":4,
    "aspect":2, "inspect":2, "respect":5, "suspect":4,
    "construct":1, "construction":3, "instruction":2,
    "abrupt":2, "corrupt":2, "erupt":2, "interrupt":3,
    "admit":4, "commit":4, "permit":3, "submit":2, "dismiss":2,
    "contradict":2, "dictionary":2, "predict":2,
    "assist":3, "consist":3, "insist":4, "resist":3,
    "depend":3, "independent":3, "expense":2, "expensive":4,
    "instruct":1, "instructed":1,
}

# ═══ 50 WORDS in 12 ROOT GROUPS ═══
ROOT_GROUPS = [
    {"root":"tract","meaning":"拉/抽","explain":"tract = to pull（拉）。拖拉机就是用来拉东西的",
     "words":[
        ("abstract","abs-(离开)","→从具体中抽离出来",True),
        ("attract","at-(朝向)","→把注意力拉过来",True),
        ("attractive","at-(朝向)+-ive(…的)","→迷人的",True),
        ("subtraction","sub-(向下)+-ion(名)","→往下抽→减法",False),
    ]},
    {"root":"press","meaning":"压","explain":"press = to press（压）。想法从心里压出来就是表达",
     "words":[
        ("express","ex-(向外)","→把想法压出来",True),
        ("expression","ex-(向外)+-ion(名)","→表达；表情",True),
        ("impress","im-(向内)","→压进脑子里",True),
        ("impression","im-(向内)+-ion(名)","→印象",True),
    ]},
    {"root":"port","meaning":"运/带","explain":"port = to carry（搬运）。港口是货物运输的地方",
     "words":[
        ("import","im-(向内)","→从外面运进来",True),
        ("portable","-able(能…的)","→能随身携带的",True),
        ("report","re-(回)","→把信息带回来",True),
        ("support","sup-(向上)","→从下面托着",True),
        ("transport","trans-(跨越)","→从一处运到另一处",True),
    ]},
    {"root":"fer","meaning":"带/拿","explain":"fer = to carry（携带）。从一个地方带到另一个地方",
     "words":[
        ("differ","dif-(分开)","→各自拿着不同的东西",True),
        ("different","dif-(分开)+-ent(…的)","→不同的",True),
        ("offer","of-(朝向)","→拿到对方面前",True),
        ("prefer","pre-(前/先)","→优先拿",True),
        ("refer","re-(回)","→拿回去参考",True),
    ]},
    {"root":"cess/cede","meaning":"走/去","explain":"cess/cede = to go（走）。从一个位置走到另一个",
     "words":[
        ("access","ac-(靠近)","→走过去靠近→通道",True),
        ("necessary","ne-(不)+-ary(…的)","→不能走开的→必需的",True),
        ("process","pro-(向前)","→向前走→过程/处理",True),
        ("success","suc-(紧接)","→紧跟着走下去",True),
    ]},
    {"root":"spect","meaning":"看","explain":"spect = to look（看）。核心是「看」的动作",
     "words":[
        ("aspect","a-(朝向)","→看过去的方面/外观",True),
        ("inspect","in-(向内)","→往里面看→检查",True),
        ("respect","re-(反复)","→反复看→因重视而尊敬",True),
        ("suspect","sus-(在下)","→从下面偷偷看→怀疑",True),
    ]},
    {"root":"struct","meaning":"建造/堆叠","explain":"struct = to build（建造）。把东西堆在一起就是建造",
     "words":[
        ("construct","con-(一起)","→堆在一起→建造",True),
        ("construction","con-(一起)+-ion(名)","→建造/建筑物",True),
        ("instruct","in-(向内)","→往脑子里构建知识→指导",True),
        ("instruction","in-(向内)+-ion(名)","→指导/说明",True),
    ]},
    {"root":"rupt","meaning":"断/破","explain":"rupt = to break（断裂）。东西突然崩断",
     "words":[
        ("abrupt","ab-(离开)","→突然断开→突然的",True),
        ("corrupt","cor-(共同)","→断了规矩→腐败的",True),
        ("erupt","e-(向外)","→从内部崩出来→喷发",True),
        ("interrupt","inter-(中间)","→从中间打断",True),
    ]},
    {"root":"mit/miss","meaning":"送/发","explain":"mit/miss = to send（发送）。把东西送出去",
     "words":[
        ("admit","ad-(朝向)","→送进来→准许进入/承认",True),
        ("commit","com-(共同)","→把想法投入行动→犯/承诺",True),
        ("permit","per-(通过)","→送过去让其通过→允许",True),
        ("submit","sub-(在下)","→从下往上送→提交",True),
        ("dismiss","dis-(离开)","→把人送走→解散/解雇",True),
    ]},
    {"root":"dict","meaning":"说","explain":"dict = to say（说）。字典 dictionary 就是说话的工具书",
     "words":[
        ("contradict","contra-(相反)","→反着说→反驳",True),
        ("dictionary","-(ion)ary","→用来查说法的书→词典",True),
        ("predict","pre-(提前)","→提前说→预测",True),
    ]},
    {"root":"sist","meaning":"站立","explain":"sist = to stand（站立）。核心是站在某处/坚持立场",
     "words":[
        ("assist","as-(靠近)","→站在你旁边→帮助",True),
        ("consist","con-(一起)","→站在一起→组成/构成",True),
        ("insist","in-(在上)","→站在…上不动→坚持",True),
        ("resist","re-(反)","→站在对面→抵抗",True),
    ]},
    {"root":"pend/pens","meaning":"悬挂/花费","explain":"pend/pens = to hang/weigh。挂着→称重→花钱",
     "words":[
        ("depend","de-(向下)","→挂在…下面→依靠",True),
        ("independent","in-(不)+de-(向下)+-ent","→不挂在别人下面→独立",True),
        ("expense","ex-(向外)","→把钱花出去→消费",True),
        ("expensive","ex-(向外)+-ive(…的)","→要花很多钱的→昂贵",True),
    ]},
]

ALL_WORDS = []
for g in ROOT_GROUPS:
    for w, affix, desc, star in g["words"]:
        info = WORD_DEFS.get(w, ('', ''))
        wf_stars = FREQ_STARS.get(w, 2)
        ALL_WORDS.append({
            "word": w, "phonetic": info[0], "definition": info[1],
            "affix_hint": affix, "root_desc": desc,
            "group_root": g["root"], "group_meaning": g["meaning"],
            "star": star, "wf": wf_stars,
        })
star_w = [w for w in ALL_WORDS if w['star']]
print(f"Total: {len(ALL_WORDS)} words, {len(ROOT_GROUPS)} root groups")

# ═══════════════════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════════════════
def set_shd(c, color):
    c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))
def set_bdr(c, **kw):
    tc=c._tc; p=tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for edge,val in kw.items():
        e=OxmlElement(f'w:{edge}'); e.set(qn('w:val'),val.get('val','single'))
        e.set(qn('w:sz'),val.get('sz','4')); e.set(qn('w:color'),val.get('color','000000'))
        e.set(qn('w:space'),'0'); b.append(e)
    p.append(b)
def bdr_all(c, c2='DDD'):
    set_bdr(c,top={'val':'single','sz':'4','color':c2},bottom={'val':'single','sz':'4','color':c2},
            start={'val':'single','sz':'4','color':c2},end={'val':'single','sz':'4','color':c2})
def hex_rgb(h):
    return int(h[0:2],16),int(h[2:4],16),int(h[4:6],16)

# ═══ COVER ═══
def build_cover(doc):
    # Remove all default paragraph spacing at start
    bg = doc.add_table(rows=1, cols=1); bg.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Set table to exact page width (A4 = 21cm = 11906 dxa)
    tbl = bg._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # Remove any existing tblW
    for child in list(tblPr):
        if child.tag == qn('w:tblW'):
            tblPr.remove(child)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '11906')  # A4 full width at 0 margins
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    
    cell = bg.cell(0,0); set_shd(cell, '0A2463')
    set_bdr(cell,top={'val':'none','sz':'0','color':'auto'},bottom={'val':'none','sz':'0','color':'auto'},
            start={'val':'none','sz':'0','color':'auto'},end={'val':'none','sz':'0','color':'auto'})
    # Set cell to fill full height
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), '11906'); tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)
    vAlign = OxmlElement('w:vAlign'); vAlign.set(qn('w:val'),'center')
    tcPr.append(vAlign)
    
    def ctr(before=0, after=0): 
        p = cell.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = Pt(14)
        return p
    
    p = ctr(50,10); r = p.add_run('✨    ⭐    ✨    ✨    ⭐    ✨'); r.font.size = Pt(18); r.font.color.rgb = RGBColor(0xFF,0xE7,0x82)
    p = ctr(60,10); r = p.add_run('高考英语词汇默写本'); r.font.size = Pt(40); r.font.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p = ctr(8,16); r = p.add_run('词根词缀 · 艾宾浩斯记忆法'); r.font.size = Pt(18); r.font.color.rgb = RGBColor(0xFF,0xD7,0x66)
    p = ctr(40,12); r = p.add_run('🌊🌊  ⛵  ⛵⛵  🌊🌊🌊  ⛵  🌊🌊'); r.font.size = Pt(24)
    p = ctr(36,16); r = p.add_run('学海无涯，扬帆起航'); r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p = ctr(6,10); r = p.add_run('每一个词根，都是你驶向未来的帆'); r.font.size = Pt(14); r.font.color.rgb = RGBColor(0xBB,0xDD,0xFF); r.font.italic = True
    p = ctr(24,8); r = p.add_run('～🌊～🌊～🌊～🌊～🌊～🌊～🌊～'); r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x4F,0xA8,0xEA)
    p = ctr(24,10); r = p.add_run('「今天的每一分努力」\r都是明天看得见的风景'); r.font.size = Pt(12); r.font.color.rgb = RGBColor(0xDD,0xEE,0xFF); r.font.italic = True
    p = ctr(14,20); r = p.add_run('⭐  ✨  ⭐  ✨  ⭐'); r.font.size = Pt(14); r.font.color.rgb = RGBColor(0xFF,0xE7,0x82)
    p = ctr(6,30); r = p.add_run('⛵  🌅  ⛵  ⛵  🌅  ⛵'); r.font.size = Pt(20); r.font.color.rgb = RGBColor(0xFF,0xA0,0x50)
    p = ctr(28,4); r = p.add_run('苏州盈信企业管理有限公司'); r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p = ctr(2,2); r = p.add_run('公司注册 · 专注财税二十五年'); r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xBB,0xDD,0xFF)
    p = ctr(2,40); r = p.add_run('18912633863'); r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = RGBColor(0xFF,0xD7,0x66)

# ═══════════════════════════════════════════════════════════════
# LEARNING CARD (Day 1) — only 2 rows
# ═══════════════════════════════════════════════════════════════
def card_learn(doc, w):
    t = doc.add_table(rows=2, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    sm = '★ ' if w['star'] else '  '
    
    # Row 0: Word | Definition (LEFT-aligned against left border)
    c = t.cell(0,0); c.text = ''
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{sm}{w['word']}")
    r.font.size = Pt(12); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E) if w['star'] else RGBColor(0x33,0x33,0x33)
    if w['phonetic']:
        r = p.add_run(f"  {w['phonetic']}"); r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x66,0x66,0x66)
    set_shd(c, 'F0F4F8'); c.width = Cm(7.0)

    c = t.cell(0,1); c.text = ''
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # LEFT-aligned, flush to left border
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{w['definition']}   ")
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xBB,0xBB,0xBB); r.font.italic = True
    r = p.add_run(f"({'＿' * 12})")
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xCC,0xCC,0xCC)
    c.width = Cm(7.5)
    
    # Row 1: Root (left) + blank | Writing index + star (right)
    c = t.cell(1,0); c.text = ''
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f"🔬 {w['word']} = {w['affix_hint']} + {w['group_root']} {w['root_desc']}")
    r.font.size = Pt(7.5); r.font.italic = True; r.font.color.rgb = RGBColor(0x2C,0x3E,0x50)
    
    c = t.cell(1,1); c.text = ''
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    dots = '●'*w['wf']+'○'*(5-w['wf'])
    r = p.add_run(f"写作指数: {dots}  ")
    r.font.size = Pt(7); r.font.color.rgb = RGBColor(0x88,0x88,0x88)
    if w['star']:
        r = p.add_run('⭐ 高频词'); r.font.size = Pt(7); r.font.bold = True; r.font.color.rgb = RGBColor(0xC0,0x39,0x2B)

    for rowr in t.rows:
        for cell in rowr.cells: bdr_all(cell)
    sp = doc.add_paragraph(); sp.paragraph_format.space_before = Pt(0); sp.paragraph_format.space_after = Pt(0)

def group_header(doc, g):
    p = doc.add_paragraph()
    r = p.add_run(f'▸ 【{g["root"]}】= {g["meaning"]}')
    r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    p = doc.add_paragraph()
    r = p.add_run(f'  {g["explain"]}'); r.font.size = Pt(7.5); r.font.italic = True; r.font.color.rgb = RGBColor(0x66,0x66,0x66)
    sp = doc.add_paragraph(); sp.paragraph_format.space_before = Pt(0); sp.paragraph_format.space_after = Pt(2)
    pPr = sp._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    bt = OxmlElement('w:bottom'); bt.set(qn('w:val'),'single'); bt.set(qn('w:sz'),'4')
    bt.set(qn('w:color'),'CCC'); bt.set(qn('w:space'),'1'); pBdr.append(bt); pPr.append(pBdr)

# ═══════════════════════════════════════════════════════════════
# 六栏复习排版 (序号缩窄，去✓□✗□)
# ═══════════════════════════════════════════════════════════════
def review_header(doc, day_label, title, subtitle, total):
    """精简标题：DAY X + 标题在同一行，居中"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(f'{day_label}  {title}')
    r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    r = p.add_run(f'  {subtitle}'); r.font.size = Pt(7); r.font.color.rgb = RGBColor(0x99,0x99,0x99)
    
    t = doc.add_table(rows=1, cols=3); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,(txt,w,clr) in enumerate([('姓名: ___________',Cm(6.0),'2C3E50'),
        ('日期: ___________',Cm(5.0),'2C3E50'),(f'正确: ___ / {total}',Cm(4.0),'C0392B')]):
        c = t.cell(0,i); c.text = ''; p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = Pt(8)
        r = p.add_run(txt); r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = RGBColor(*hex_rgb(clr))
        set_shd(c, 'F5F5F5')
        set_bdr(c, top={'val':'single','sz':'4','color':'CCC'},bottom={'val':'single','sz':'4','color':'CCC'},
                start={'val':'single','sz':'4','color':'CCC'},end={'val':'single','sz':'4','color':'CCC'})
    # Compact spacer
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)

def sixcol_review(doc, word_list, mode):
    """六栏复习：序号极限窄，行高填满页面"""
    total = len(word_list)
    rows_needed = math.ceil(total / 2)
    
    t = doc.add_table(rows=rows_needed + 1, cols=6)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Full page width
    tbl = t._tbl; tblPr = tbl.tblPr
    if tblPr is None: tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'),'11000'); tblW.set(qn('w:type'),'dxa')
    tblPr.append(tblW)
    # 序号列极限窄0.15cm，填空列最宽
    col_widths = [Cm(0.15), Cm(3.5), Cm(5.35), Cm(0.15), Cm(3.5), Cm(5.35)]
    
    if mode == 'en2cn':
        h = ['序号','英文','中文（填空）','序号','英文','中文（填空）']
    elif mode == 'cn2en':
        h = ['序号','中文释义','英文（填空）','序号','中文释义','英文（填空）']
    elif mode == 'phonetic':
        h = ['序号','🔊音标','英+中（填空）','序号','🔊音标','英+中（填空）']
    else:
        h = ['序号','原文','填空','序号','原文','填空']
    
    for ci in range(6):
        c = t.cell(0, ci); c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h[ci]); r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = RGBColor(0x33,0x33,0x33)
        set_shd(c, 'EAEAEA')
        set_bdr(c, top={'val':'single','sz':'4','color':'AAA'},bottom={'val':'single','sz':'6','color':'666'},
                start={'val':'single','sz':'4','color':'AAA'},end={'val':'single','sz':'4','color':'AAA'})
        c.width = col_widths[ci]
    # Set header row height
    t.cell(0,0)._tc.get_or_add_tcPr()
    trPr0 = t.rows[0]._tr.get_or_add_trPr()
    trH0 = OxmlElement('w:trHeight'); trH0.set(qn('w:val'),'400'); trH0.set(qn('w:hRule'),'atLeast')
    trPr0.append(trH0)
    
    # Calculate row height to fill page: ~26.4cm for 25 rows = 1.05cm ≈ 2977 dxa
    row_h = '600'
    for ri in range(rows_needed):
        for col_offset, idx in [(0, ri), (3, ri + rows_needed)]:
            col_idx = col_offset  # 0 or 3
            if idx >= total: continue
            w = word_list[idx]
            sm = '★ ' if w['star'] else ''
            
            # Number cell — ultra narrow
            c = t.cell(ri+1, col_offset); c.text = ''
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(idx+1)); r.font.size = Pt(7); r.font.color.rgb = RGBColor(0x99,0x99,0x99)
            # Vertical center
            tcPr = c._tc.get_or_add_tcPr()
            vAlign = OxmlElement('w:vAlign'); vAlign.set(qn('w:val'),'center'); tcPr.append(vAlign)
            
            # Content cell
            c = t.cell(ri+1, col_offset+1); c.text = ''
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
            if mode == 'en2cn':
                r = p.add_run(f"{sm}{w['word']}")
                r.font.size = Pt(10); r.font.bold = True
                r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E) if w['star'] else RGBColor(0x33,0x33,0x33)
            elif mode == 'cn2en':
                r = p.add_run(w['definition'][:30])
                r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x44,0x44,0x44)
            elif mode == 'phonetic':
                r = p.add_run(f"🔊{w['phonetic']}")
                r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x2C,0x3E,0x50)
            tcPr = c._tc.get_or_add_tcPr()
            vAlign = OxmlElement('w:vAlign'); vAlign.set(qn('w:val'),'center'); tcPr.append(vAlign)
            
            # Blank cell — extend to right
            c = t.cell(ri+1, col_offset+2); c.text = ''
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
            if mode == 'en2cn':
                r = p.add_run(f"({'＿'*14})")
            elif mode == 'cn2en':
                r = p.add_run(f"({'＿'*14})")
            elif mode == 'phonetic':
                r = p.add_run(f"英[{'＿'*8}]中[{'＿'*8}]")
            r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xCC,0xCC,0xCC)
            tcPr = c._tc.get_or_add_tcPr()
            vAlign = OxmlElement('w:vAlign'); vAlign.set(qn('w:val'),'center'); tcPr.append(vAlign)
            
            # Borders
            for cix in range(3):
                set_bdr(t.cell(ri+1, col_offset+cix),
                    top={'val':'single','sz':'4','color':'EEE'},bottom={'val':'single','sz':'4','color':'EEE'},
                    start={'val':'single','sz':'4','color':'EEE'},end={'val':'single','sz':'4','color':'EEE'})
        
        # Set row height to fill available space
        trPr = t.rows[ri+1]._tr.get_or_add_trPr()
        trH = OxmlElement('w:trHeight'); trH.set(qn('w:val'), row_h); trH.set(qn('w:hRule'), 'atLeast')
        trPr.append(trH)
    
    for row in t.rows:
        for ci in range(6): row.cells[ci].width = col_widths[ci]

# ═══════════════════════════════════════════════════════════════
# BUILD
# ═══════════════════════════════════════════════════════════════
doc = Document()
style = doc.styles['Normal']; style.font.name = '微软雅黑'; style.font.size = Pt(9)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
for sec in doc.sections:
    sec.top_margin = Cm(0); sec.bottom_margin = Cm(0)
    sec.left_margin = Cm(0); sec.right_margin = Cm(0)

build_cover(doc)

# New section for content pages with normal margins
new_sec = doc.add_section()
new_sec.top_margin = Cm(0.8); new_sec.bottom_margin = Cm(0.8)
new_sec.left_margin = Cm(0.8); new_sec.right_margin = Cm(0.8)

# ═══ INSTRUCTIONS ═══
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('📋 使用说明 & 复习日程'); r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
for t in ['① 按词根分组 — 同根词批量记忆，一个词根解锁一串单词',
          '② 写作指数来自真实英文语料库频次数据（1-5星）',
          '③ 第一轮学习为单栏卡片，第二轮起为六栏紧凑排版',
          '④ 六栏复习：序号列缩窄，填空列加宽，无✓□✗□标记']:
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(t); r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x44,0x44,0x44)
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('🗓 复习日程表'); r.font.size = Pt(12); r.font.bold = True; r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
sched = [['DAY 1','📖 学习','英→中','单栏卡片（词根拆解）'],
         ['DAY 2','🔄 复习①','英→中','六栏：英文→写中文'],
         ['DAY 4','🔄 复习②','中→英','六栏：中文→写英文'],
         ['DAY 7','🔄 复习③','音标→英+中','六栏：音标→写单词+释义'],
         ['DAY 15','🔄 复习④','中→英 ★','六栏：仅⭐高频词'],
         ['DAY 30','🔄 复习⑤','混合','六栏：英→中/中→英各半']]
t = doc.add_table(rows=len(sched)+1, cols=4); t.style='Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i,h in enumerate(['天数','类型','方向','版式']):
    c = t.cell(0,i); c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(h); r.font.size=Pt(8.5); r.font.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    set_shd(c,'2C3E50')
for ri,row in enumerate(sched):
    for ci,val in enumerate(row):
        c=t.cell(ri+1,ci); c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
        r=p.add_run(val); r.font.size=Pt(8)
        if ci==2 and '中→英' in val: r.font.bold=True; r.font.color.rgb=RGBColor(0xC0,0x39,0x2B)
        if ci==2 and '音标' in val: r.font.bold=True; r.font.color.rgb=RGBColor(0x2C,0x3E,0x50)
        if ri%2: set_shd(c,'F5F5F5')

# ═══ DAY 1: 学习页 ═══
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('DAY 1  学习页'); r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('按词根分组  50词汇'); r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x99,0x99,0x99)

for g in ROOT_GROUPS:
    group_header(doc, g)
    for wd, affix, desc, star in g["words"]:
        w_info = WORD_DEFS.get(wd, ('', ''))
        card_learn(doc, {"word":wd,"phonetic":w_info[0],"definition":w_info[1],
                          "affix_hint":affix,"root_desc":desc,
                          "group_root":g["root"],"group_meaning":g["meaning"],
                          "star":star,"wf":FREQ_STARS.get(wd,2)})

# ═══ DAY 2: 英→中 六栏 ═══
doc.add_page_break()
order = list(range(50))
random.seed(42); random.shuffle(order)
w_en = [ALL_WORDS[i] for i in order]
review_header(doc, 'DAY 2', '英→中', '复习① · 左列1-25 / 右列26-50', 50)
sixcol_review(doc, w_en, 'en2cn')

# ═══ DAY 4: 中→英 六栏 ═══
doc.add_page_break()
random.seed(77); random.shuffle(order)
w_cn = [ALL_WORDS[i] for i in order]
review_header(doc, 'DAY 4', '中→英', '复习② · 左列1-25 / 右列26-50', 50)
sixcol_review(doc, w_cn, 'cn2en')

# ═══ DAY 7: 音标→英+中 六栏 ═══
doc.add_page_break()
random.seed(123); random.shuffle(order)
w_ph = [ALL_WORDS[i] for i in order]
review_header(doc, 'DAY 7', '🔊 音标→英文+中文', '复习③ · 左列1-25 / 右列26-50', 50)
sixcol_review(doc, w_ph, 'phonetic')

# ═══ DAY 15: ★ 中→英 六栏 ═══
doc.add_page_break()
star_list = sorted(star_w, key=lambda x: -x['wf'])
review_header(doc, 'DAY 15', f'★ 中→英  高频词拼写', f'复习④ · ⭐高频词（{len(star_list)}词）', len(star_list))
sixcol_review(doc, star_list, 'cn2en')

# ═══ DAY 30: 混合终测 六栏 ═══
doc.add_page_break()
random.seed(999); random.shuffle(order)
w_mix = [ALL_WORDS[i] for i in order]
review_header(doc, 'DAY 30', '🎲 混合终测', '复习⑤ · 英→中/中→英各半', 50)

total = 50; rows_needed = math.ceil(total/2)
t = doc.add_table(rows=rows_needed+1, cols=6)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl = t._tbl; tblPr = tbl.tblPr
if tblPr is None: tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'),'11000'); tblW.set(qn('w:type'),'dxa')
tblPr.append(tblW)
col_widths = [Cm(0.15), Cm(3.5), Cm(5.35), Cm(0.15), Cm(3.5), Cm(5.35)]
for ci in range(6):
    c = t.cell(0,ci); c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(['序号','原文','填空','序号','原文','填空'][ci])
    r.font.size=Pt(8); r.font.bold=True; r.font.color.rgb=RGBColor(0x33,0x33,0x33)
    set_shd(c,'EAEAEA')
    set_bdr(c,top={'val':'single','sz':'4','color':'AAA'},bottom={'val':'single','sz':'6','color':'666'},
            start={'val':'single','sz':'4','color':'AAA'},end={'val':'single','sz':'4','color':'AAA'})
    c.width=col_widths[ci]
# Header row height
trPr0 = t.rows[0]._tr.get_or_add_trPr()
trH0 = OxmlElement('w:trHeight'); trH0.set(qn('w:val'),'400'); trH0.set(qn('w:hRule'),'atLeast')
trPr0.append(trH0)

for ri in range(rows_needed):
    for col_offset, idx in [(0, ri), (3, ri+rows_needed)]:
        if idx >= total: continue
        w = w_mix[idx]; is_en = (idx % 2 == 0); sm = '★ ' if w['star'] else ''
        
        c = t.cell(ri+1, col_offset); c.text=''
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
        r=p.add_run(str(idx+1)); r.font.size=Pt(7); r.font.color.rgb=RGBColor(0x99,0x99,0x99)
        tcPr=c._tc.get_or_add_tcPr()
        vAlign=OxmlElement('w:vAlign'); vAlign.set(qn('w:val'),'center'); tcPr.append(vAlign)
        
        c = t.cell(ri+1, col_offset+1); c.text=''
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
        if is_en:
            r=p.add_run(f"{sm}{w['word']}"); r.font.size=Pt(10); r.font.bold=True
            r.font.color.rgb=RGBColor(0x1A,0x1A,0x2E) if w['star'] else RGBColor(0x33,0x33,0x33)
        else:
            r=p.add_run(w['definition'][:28]); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x44,0x44,0x44)
        tcPr=c._tc.get_or_add_tcPr()
        vAlign=OxmlElement('w:vAlign'); vAlign.set(qn('w:val'),'center'); tcPr.append(vAlign)
        
        c = t.cell(ri+1, col_offset+2); c.text=''
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
        if is_en: r=p.add_run(f"中文[{'＿'*10}]")
        else: r=p.add_run(f"英文[{'＿'*10}]")
        r.font.size=Pt(9); r.font.color.rgb=RGBColor(0xCC,0xCC,0xCC)
        tcPr=c._tc.get_or_add_tcPr()
        vAlign=OxmlElement('w:vAlign'); vAlign.set(qn('w:val'),'center'); tcPr.append(vAlign)
        
        for cix in range(3):
            set_bdr(t.cell(ri+1, col_offset+cix),
                top={'val':'single','sz':'4','color':'EEE'},bottom={'val':'single','sz':'4','color':'EEE'},
                start={'val':'single','sz':'4','color':'EEE'},end={'val':'single','sz':'4','color':'EEE'})
    
    # Fill height
    trPr = t.rows[ri+1]._tr.get_or_add_trPr()
    trH = OxmlElement('w:trHeight'); trH.set(qn('w:val'),'600'); trH.set(qn('w:hRule'),'atLeast')
    trPr.append(trH)

for row in t.rows:
    for ci in range(6): row.cells[ci].width = col_widths[ci]

# ═══ 高频词速查 ═══
doc.add_page_break()
p = doc.add_paragraph()
r = p.add_run('✍ 高频写作词速查（按词根分组）'); r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
p = doc.add_paragraph()
r = p.add_run('⭐ 写作指数基于真实英文语料库频次数据（1星=低频 · 5星=最高频）')
r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x99,0x99,0x99)

for g in ROOT_GROUPS:
    gs = [w for w in star_w if w['group_root']==g['root']]
    if not gs: continue
    p = doc.add_paragraph()
    r = p.add_run(f'▸ {g["root"]} ({g["meaning"]})')
    r.font.size = Pt(9.5); r.font.bold = True; r.font.color.rgb = RGBColor(0x2C,0x3E,0x50)
    for w in sorted(gs, key=lambda x: -x['wf']):
        dots = '●'*w['wf']+'○'*(5-w['wf'])
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(f'{w["word"]} — {w["definition"][:30]}  {dots}')
        r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x44,0x44,0x44)

# ═══ BACK ═══
doc.add_page_break()
for _ in range(4): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('50词 × 12个词根 = 批量记忆的起点'); r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('写作指数基于真实语料库频次数据'); r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x7F,0x8C,0x8D)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('满意了？说一声 → 启动完整版 3875词'); r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x55,0x55,0x55)
doc.add_paragraph(); doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('苏州盈信企业管理有限公司'); r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xAA,0xAA,0xAA)

doc.save(DOCX_PATH)
print(f"\n✅ {DOCX_PATH}")
print(f"📊 50词 / {len(ROOT_GROUPS)}词根组")
print(f"✍ 写作指数基于真实英文语料库频次数据")
print(f"📐 学习=2行卡片 | 复习=六栏(序号缩窄·无✓□✗□·Day30)")
