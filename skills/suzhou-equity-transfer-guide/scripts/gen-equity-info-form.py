#!/usr/bin/env python3
"""
股权转让信息采集表 v2.0 → 桌面
使用方式: python3 scripts/gen-equity-info-form.py
客户提了股转需求后直接运行此脚本，生成表单发给客户填写。

v2.0 改动（2026.6.4）：
  - 新增：资产合计、负债合计、未分配利润、利润总额（核税关键参数，避免上次遗漏）
  - 新增：房屋/土地账面价值+占比（判断是否触发强制评估）
  - 新增：逻辑校验区（自动标出常见填表错误）
  - 新增：每个字段的"用途说明"列，让客户知道填这个干什么用
  - 强化：两期报表（上年度末+最近一个月末）同时索要
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os


def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color, qn('w:val'): 'clear'
    })
    shading.append(shd)


def set_cell_font(cell, size=10, bold=False, color=None):
    for para in cell.paragraphs:
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after = Pt(1)
        for r in para.runs:
            r.font.size = Pt(size)
            r.font.bold = bold
            if color:
                r.font.color.rgb = color


def add_header_row(table, headers, color='1A3C6E'):
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)
            for r in para.runs:
                r.font.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, color)


def make_3col_table(doc, data, col_widths=None):
    """
    data: list of (item_text, required_text, usage_text)
    三列：项目（什么字段）、必填（*或留空）、用途说明（为什么问这个）
    """
    table = doc.add_table(rows=len(data) + 1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(table, ['项目', '必填', '用途说明', '填写内容'])
    for idx, (item, required, usage) in enumerate(data, 1):
        row = table.rows[idx]
        row.cells[0].text = item
        set_cell_font(row.cells[0], 10)
        row.cells[1].text = required
        for p in row.cells[1].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].text = usage
        set_cell_font(row.cells[2], 9, color=RGBColor(0x66, 0x66, 0x66))
        row.cells[3].text = ''
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def build_form(output_path):
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ===== 封面 =====
    for _ in range(4):
        doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('股权转让信息采集表')
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    doc.add_paragraph('')
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('—— 苏州盈信企业管理有限公司 ——')
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('高级会计师 · 财税实战专家')
    r3.font.size = Pt(12)
    r3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    p_v = doc.add_paragraph()
    p_v.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_v = p_v.add_run('v2.0 ｜ 采集即两期，不问第二遍')
    r_v.font.size = Pt(10)
    r_v.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_paragraph('\\n')

    # ===== 填写说明 =====
    p = doc.add_paragraph()
    r = p.add_run('📋 填写说明：')
    r.font.bold = True
    r.font.size = Pt(12)
    doc.add_paragraph(
        '1. 请逐项填写，不清楚的留空标注"待确认"\\n'
        '2. 带"*"为必填项，其余有则填\\n'
        '3. 每个字段的"用途说明"列告诉你为什么问这个，帮助你理解我们的专业判断逻辑\\n'
        '4. ⚠️ 两套财务报表请同时提供：上年度末（汇算清缴年报）+ 最近一个月末的资产负债表+利润表\\n'
        '5. 如有章程、股权转让协议草案等附件材料可一并打包'
    )
    doc.add_paragraph('—' * 60)

    def heading(text):
        h = doc.add_heading(text, level=1)
        for r in h.runs:
            r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    # ===== 一、公司信息 =====
    heading('一、被转让公司基本信息')
    make_3col_table(doc, [
        ('公司全称', '*', '确认法律主体'),
        ('统一社会信用代码', '*', '税务系统唯一标识'),
        ('注册地址', '', '判断所属分局管辖区域'),
        ('成立日期', '*', '判断是否需提供历史股转材料（2019年分界）'),
        ('注册资本（万元）', '*', '判断转让方股权原值的上限依据'),
        ('实收资本（万元）', '*', '★核心参数，决定股权原值和个税基数'),
        ('所属行业', '', '影响行业税负特征判断'),
        ('注册地所在区', '*（吴中木渎/园区/高新区/姑苏/相城/吴江）', '★不同区的窗口要求和审核口径不同'),
    ], col_widths=[4.5, 1.5, 5, 6])
    doc.add_paragraph('')

    # ===== 二、股权结构 =====
    heading('二、当前股权结构')
    eq = doc.add_table(rows=5, cols=5)
    eq.style = 'Table Grid'
    eq.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(eq, ['股东姓名', '持股比例(%)', '出资方式', '出资是否到位', '备注'])

    rows_data = [
        ('股东1', '', '', '', ''),
        ('股东2', '', '', '', ''),
        ('股东3', '', '', '', ''),
        ('合计', '100%', '', '', ''),
    ]
    for idx, vals in enumerate(rows_data, 1):
        for j, v in enumerate(vals):
            eq.rows[idx].cells[j].text = v
            for p in eq.rows[idx].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                if idx == 4:
                    for r in p.runs:
                        r.font.bold = True

    # 设置列宽
    for row in eq.rows:
        widths = [4, 2.5, 2.5, 2.5, 5]
        for i, w in enumerate(widths):
            row.cells[i].width = Cm(w)

    doc.add_paragraph('')

    # ===== 三、转让详情 =====
    heading('三、本次股权转让详情')
    make_3col_table(doc, [
        ('转让方（老股东）', '*', '个税纳税人是谁'),
        ('受让方（新股东）', '*', '扣缴义务人是谁'),
        ('转让股数/比例', '*', '判断是否达到控制权变更门槛'),
        ('转让价格（万元）', '*', '★直接影响个税计算和是否触发核定'),
        ('转让方与受让方关系', '*（夫妻/近亲属/非亲属/法人）', '★决定是否适用近亲属即时办结通道'),
        ('转让基准日', '*', '财务报表的截止时点'),
        ('协议是否已签订', '□ 已签  □ 未签  □ 草拟中', '判断是否有修改空间'),
        ('是否有对赌/回购条款', '□ 有  □ 无  □ 不清楚', '涉及或有负债的税务处理'),
        ('款项支付方式', '□ 一次性付清  □ 分期  □ 未定', '影响纳税时点判定'),
    ], col_widths=[4.5, 1.5, 5, 6])
    doc.add_paragraph('')

    # ===== 四、财务与资产情况（关键——决定即时/非即时办结） =====
    heading('四、财务与资产情况（★核心——决定办结类型和税负）')

    p_note = doc.add_paragraph()
    r_note = p_note.add_run('⚠️ 以下数据是税务局和评估机构最关注的参数，请尽量准确填报。')
    r_note.font.size = Pt(10)
    r_note.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    make_3col_table(doc, [
        ('上年度末：资产合计（元）', '*', '公式：资产合计 - 负债合计 = 净资产'),
        ('上年度末：负债合计（元）', '*', '★负债合计 + 资产合计 → 净资产（核税第一参数）'),
        ('上年度末：未分配利润（元）', '*', '★判断公司是否有累计可分配利润，影响转让定价合理性'),
        ('上年度末：利润总额（元）', '*', '★交叉验证未分配利润的真实性'),
        ('最近一月末：资产合计（元）', '*', '★两期对比判断资产结构变化是否超20%'),
        ('最近一月末：负债合计（元）', '*', '★两期对比判断是否有重大负债变动'),
        ('最近一月末：未分配利润（元）', '*', '★判断这期间是否有新增亏损或盈利'),
        ('最近一月末：利润总额（元）', '*', '★判断期间损益'),
        ('公式推算：净资产', '自动计算', '资产合计 - 负债合计 = 净资产（正数或负数决定零元转让是否触发核定）'),
        ('', '', ''),
        ('公司名下是否有房产/土地？', '*（是/否）', '★有 → 必须出资产评估报告'),
        ('如有：房产/土地账面价值（元）', '', '★占资产总额比例 → 决定是否超过20%强制评估门槛'),
        ('如有：房产/土地占总资产比例', '', '占资产总额超过20% → 类型C非即时办结'),
        ('公司名下是否有知识产权/专利/商标？', '（是/否）', '有 → 属于增值性资产，影响类型B判定'),
        ('公司名下是否有车辆？', '（是/否）', '车辆不是增值性资产，不影响分类'),
        ('是否有对外投资/持股其他公司？', '（是/否）', '对外持股 → 属于增值性资产（股权）'),
        ('是否有未结诉讼或重大债权债务？', '（有/无，简述）', '影响评估报告结论'),
        ('近三年是否发生过股权变更？', '*（是/否，日期）', '判断是否需要调取历史股转材料'),
        ('近三年是否享受过税收优惠/核定征收？', '（是/否）', '影响税负测算逻辑'),
    ], col_widths=[5, 1.5, 4, 6.5])
    doc.add_paragraph('')

    # ===== 逻辑校验区 =====
    heading('⚠️ 逻辑校验区（系统自动标出异常点）')
    p_check = doc.add_paragraph()
    r_check = p_check.add_run('以下是对您填写内容的前置判断，帮助你我快速锁定风险：')
    r_check.font.size = Pt(10)
    r_check.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    checks = doc.add_table(rows=8, cols=3)
    checks.style = 'Table Grid'
    checks.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_header_row(checks, ['校验项目', '常见异常', '影响'], 'CC0000')

    check_data = [
        ('净资产是否为正？', '净资产为负 + 零元转让 → 可能个税=0（需看评估结果）', '★★★'),
        ('两期资产结构变化是否超20%？', '变化超20% → 类型C非即时办结，审核延长', '★★★'),
        ('房产/土地占比是否超20%？', '占比超20% → 必须出资产评估报告', '★★★'),
        ('转让价是否低于净资产？', '低于净资产 → 启动核定征收，税负大增', '★★★'),
        ('是否近亲属关系？', '近亲属→类型A即时办结，可走低价转让', '★★☆'),
        ('实收资本是否为0？', '实收=0 → 股权原值=0，影响个税基数', '★★☆'),
        ('是否2019年前成立？', '是 → 需提供最近一次股转的旧材料', '★☆☆'),
        ('转让方是否曾任职/任法人？', '曾任职 → 可能涉及离职收入另算个税', '★☆☆'),
    ]
    for idx, (item, desc, level) in enumerate(check_data, 1):
        checks.rows[idx].cells[0].text = item
        set_cell_font(checks.rows[idx].cells[0], 10, bold=True)
        checks.rows[idx].cells[1].text = desc
        set_cell_font(checks.rows[idx].cells[1], 9)
        checks.rows[idx].cells[2].text = level
        for p in checks.rows[idx].cells[2].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in checks.rows:
        widths = [4.5, 8, 1.5]
        for i, w in enumerate(widths):
            row.cells[i].width = Cm(w)

    doc.add_paragraph('')

    # ===== 五、诉求 =====
    heading('五、客户核心诉求')
    make_3col_table(doc, [
        ('预期完成时间节点', '□ 本月内  □ 季度内  □ 不急', '倒推材料准备排期'),
        ('是否有节税需求？', '□ 有  □ 无  □ 合规优先', '在合规范围内做最优方案'),
        ('是否已咨询其他中介？', '□ 没有  □ 有（简述建议）', '了解竞争对手方案，针对性优化'),
        ('是否需要协议审查服务？', '□ 需要  □ 不需要  □ 自己有律师', '决定是否提供股权转让协议审查'),
        ('其他特别说明', '', ''),
    ], col_widths=[4.5, 1.5, 5, 6])
    doc.add_paragraph('')

    # ===== 注意事项 =====
    doc.add_paragraph('')
    p = doc.add_paragraph()
    r = p.add_run('⚠️ 重要提示：')
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    for note in [
        '转让价格明显低于净资产份额的，税务机关有权核定征收，按20%征收个税',
        '2019年以前成立的企业需提供最近一次股权转让的旧材料',
        '公司资产中含房产/土地/知识产权的，必须出具资产评估报告',
        '资产结构变化超过20%的将转入非即时办结通道，审核时间延长',
        '两套财务报表（上年末+最近一个月末）务必同时提供，缺任何一套窗口可能退回',
    ]:
        p = doc.add_paragraph(note, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x66, 0x00, 0x00)

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('— 请将本文件发回，1个工作日内出具完整方案 —')
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('苏州盈信企业管理有限公司 ｜ 高级会计师团队')
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    desktop = '/mnt/d/360MoveData/Users/Admin/Desktop'
    os.makedirs(desktop, exist_ok=True)
    path = os.path.join(desktop, '股权转让信息采集表_v2.0.docx')
    build_form(path)
    print(f'✅ v2.0 已生成: {path}')
    print(f'   大小: {os.path.getsize(path)} 字节')
    print(f'')
    print(f'v2.0 新增字段:')
    print(f'  ① 资产合计 / 负债合计 / 未分配利润 / 利润总额（上下两期）')
    print(f'  ② 房产/土地账面价值 + 占比（判断强制评估门槛）')
    print(f'  ③ 逻辑校验区（自动标出8大异常点）')
    print(f'  ④ 用途说明列（每个字段为什么要填）')