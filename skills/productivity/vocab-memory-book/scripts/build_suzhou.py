# -*- coding: utf-8 -*-
"""苏州中考英语词汇默写本 - 基于频次筛选的考纲词汇"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import re, math, os, time, random

SRC = r"/tmp/suzhou_vocab.txt"
FREQ = r"/home/dmin/.hermes/skills/productivity/vocab-memory-book/en_50k.txt"
OUT = r"/mnt/c/Users/Admin/Desktop/苏州中考英语词汇默写本.docx"
WPG = 30

def set_shd(c, color):
    c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))
def set_bdr(c, **kw):
    tc=c._tc; p=tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for edge,val in kw.items():
        e=OxmlElement(f'w:{edge}'); e.set(qn('w:val'),val.get('val','single'))
        e.set(qn('w:sz'),val.get('sz','4')); e.set(qn('w:color'),val.get('color','000000'))
        e.set(qn('w:space'),'0'); b.append(e)
    p.append(b)
def hex_rgb(h): return int(h[0:2],16),int(h[2:4],16),int(h[4:6],16)

# Parse vocab
all_words = []
with open(SRC, 'r', errors='replace') as f:
    lines = f.read().strip().split('\n')
for l in lines[3:]:
    l=l.strip()
    if not l: continue
    m=re.match(r'([a-zA-Z\-\(\)]+)\s',l)
    if not m: continue
    word=m.group(1).lower().rstrip(',)')
    ph=re.search(r'\[([^\]]+)\]',l)
    phon=ph.group(1) if ph else ''
    rest=l[m.end():].strip()
    if phon: rest=rest.replace(f'[{phon}]','',1).strip()
    all_words.append({'word':word,'phonetic':phon,'definition':rest[:90]})

# Frequency
freq_data={}
with open(FREQ,'r') as f:
    for l in f:
        l=l.strip()
        if not l: continue
        parts=l.split(None,1)
        if len(parts)==2: freq_data[parts[0].lower()]=int(parts[1])
fv=sorted(freq_data.values(),reverse=True)
p80=fv[max(0,len(fv)//5*4)] if fv else 0
p60=fv[max(0,len(fv)//5*3)] if fv else 0
p40=fv[max(0,len(fv)//5*2)] if fv else 0
p20=fv[max(0,len(fv)//5*1)] if fv else 0
def stars(w):
    f=freq_data.get(w)
    if f is None: return 1
    if f>=p80: return 5
    if f>=p60: return 4
    if f>=p40: return 3
    if f>=p20: return 2
    return 1
for w in all_words: w['stars']=stars(w['word'])

# Group
random.seed(20260704); random.shuffle(all_words)
groups=[all_words[i:i+WPG] for i in range(0,len(all_words),WPG)]

# Build doc
doc=Document()
style=doc.styles['Normal']; style.font.name='微软雅黑'; style.font.size=Pt(9)
style.element.rPr.rFonts.set(qn('w:eastAsia'),'微软雅黑')
for sec in doc.sections:
    sec.top_margin=Cm(0); sec.bottom_margin=Cm(0)
    sec.left_margin=Cm(0); sec.right_margin=Cm(0)

# Cover
bg=doc.add_table(rows=1,cols=1); bg.alignment=WD_TABLE_ALIGNMENT.CENTER
tbl=bg._tbl; tblPr=tbl.tblPr
if tblPr is None: tblPr=OxmlElement('w:tblPr'); tbl.insert(0,tblPr)
for child in list(tblPr):
    if child.tag==qn('w:tblW'): tblPr.remove(child)
tblW=OxmlElement('w:tblW'); tblW.set(qn('w:w'),'11906'); tblW.set(qn('w:type'),'dxa'); tblPr.append(tblW)
cell=bg.cell(0,0); set_shd(cell,'0A2463')
set_bdr(cell,top={'val':'none','sz':'0','color':'auto'},bottom={'val':'none','sz':'0','color':'auto'},
        start={'val':'none','sz':'0','color':'auto'},end={'val':'none','sz':'0','color':'auto'})
tc=cell._tc; tcPr=tc.get_or_add_tcPr()
tcW=OxmlElement('w:tcW'); tcW.set(qn('w:w'),'11906'); tcW.set(qn('w:type'),'dxa'); tcPr.append(tcW)
vAlign=OxmlElement('w:vAlign'); vAlign.set(qn('w:val'),'center'); tcPr.append(vAlign)

def ctr(before=0,after=0):
    p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after)
    p.paragraph_format.line_spacing=Pt(14); return p

ctr(50,10).add_run('✨    ⭐    ✨    ✨    ⭐    ✨').font.size=Pt(18)
r=ctr(60,10).add_run('苏州中考英语词汇默写本'); r.font.size=Pt(36); r.font.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
r=ctr(8,16).add_run('全国通用中考考纲 · 艾宾浩斯记忆法'); r.font.size=Pt(16); r.font.color.rgb=RGBColor(0xFF,0xD7,0x66)
r=ctr(40,12).add_run('🌊🌊  ⛵  ⛵⛵  🌊🌊🌊  ⛵  🌊🌊').font.size=Pt(24)
r=ctr(36,12).add_run(f'{len(all_words)}词 · {len(groups)}组 · 每组{WPG}词 · 艾宾浩斯六轮复习')
r.font.size=Pt(11); r.font.color.rgb=RGBColor(0xBB,0xDD,0xFF); r.font.italic=True
r=ctr(20,8).add_run('「今天的每一分努力」\r都是明天看得见的风景')
r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xDD,0xEE,0xFF); r.font.italic=True
r=ctr(28,4).add_run('苏州盈信企业管理有限公司'); r.font.size=Pt(13); r.font.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
r=ctr(2,2).add_run('公司注册 · 专注财税二十五年'); r.font.size=Pt(10); r.font.color.rgb=RGBColor(0xBB,0xDD,0xFF)
r=ctr(2,40).add_run('18912633863'); r.font.size=Pt(13); r.font.bold=True; r.font.color.rgb=RGBColor(0xFF,0xD7,0x66)

# Content section
sec2=doc.add_section()
sec2.top_margin=Cm(0.8); sec2.bottom_margin=Cm(0.8); sec2.left_margin=Cm(0.8); sec2.right_margin=Cm(0.8)
# Footer
footer=sec2.footer; footer.is_linked_to_previous=False
fp=footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; fp.paragraph_format.space_before=Pt(2); fp.paragraph_format.space_after=Pt(2)
fc1=OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'),'begin')
r1=fp.add_run(); r1.font.size=Pt(7); r1._r.append(fc1)
it1=OxmlElement('w:instrText'); it1.text=' PAGE '
r2=fp.add_run(); r2.font.size=Pt(7); r2._r.append(it1)
fc2=OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'),'end')
r3=fp.add_run(); r3.font.size=Pt(7); r3._r.append(fc2)
r4=fp.add_run(' / '); r4.font.size=Pt(7); r4.font.color.rgb=RGBColor(0x99,0x99,0x99)
fc3=OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'),'begin')
r5=fp.add_run(); r5.font.size=Pt(7); r5._r.append(fc3)
it2=OxmlElement('w:instrText'); it2.text=' NUMPAGES '
r6=fp.add_run(); r6.font.size=Pt(7); r6._r.append(it2)
fc4=OxmlElement('w:fldChar'); fc4.set(qn('w:fldCharType'),'end')
r7=fp.add_run(); r7.font.size=Pt(7); r7._r.append(fc4)

CW=[Cm(0.2),Cm(3.56),Cm(5.24),Cm(0.2),Cm(3.56),Cm(5.24)]

def card(w):
    t=doc.add_table(rows=2,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    sm='★ ' if w['stars']>=4 else ''
    c=t.cell(0,0); c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(f"{sm}{w['word']}"); r.font.size=Pt(12); r.font.bold=True
    r.font.color.rgb=RGBColor(0x1A,0x1A,0x2E) if w['stars']>=4 else RGBColor(0x33,0x33,0x33)
    if w['phonetic']: r=p.add_run(f"  {w['phonetic']}"); r.font.size=Pt(8); r.font.color.rgb=RGBColor(0x66,0x66,0x66)
    set_shd(c,'F0F4F8')
    c=t.cell(0,1); c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(f"{w['definition']}   "); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0xBB,0xBB,0xBB); r.font.italic=True
    r=p.add_run(f"({'＿'*12})"); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0xCC,0xCC,0xCC)
    c=t.cell(1,0); c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
    r=p.add_run("📖 中考考纲词汇"); r.font.size=Pt(7.5); r.font.italic=True; r.font.color.rgb=RGBColor(0x2C,0x3E,0x50)
    c=t.cell(1,1); c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
    dots='●'*w['stars']+'○'*(5-w['stars'])
    r=p.add_run(f"写作指数: {dots}  "); r.font.size=Pt(7); r.font.color.rgb=RGBColor(0x88,0x88,0x88)
    if w['stars']>=4: r=p.add_run('⭐ 高频词'); r.font.size=Pt(7); r.font.bold=True; r.font.color.rgb=RGBColor(0xC0,0x39,0x2B)
    for rowr in t.rows:
        for cell in rowr.cells:
            set_bdr(cell,top={'val':'single','sz':'4','color':'DDD'},bottom={'val':'single','sz':'4','color':'DDD'},
                    start={'val':'single','sz':'4','color':'DDD'},end={'val':'single','sz':'4','color':'DDD'})
    sp=doc.add_paragraph(); sp.paragraph_format.space_before=Pt(0); sp.paragraph_format.space_after=Pt(0)

def rtitle(label,sub):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(1)
    r=p.add_run(label); r.font.size=Pt(13); r.font.bold=True; r.font.color.rgb=RGBColor(0x1A,0x1A,0x2E)
    r=p.add_run(f'  {sub}'); r.font.size=Pt(7); r.font.color.rgb=RGBColor(0x99,0x99,0x99)

def ndtable(total):
    t=doc.add_table(rows=1,cols=3); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(txt,w,clr) in enumerate([('姓名: ___________',Cm(6.0),'2C3E50'),('日期: ___________',Cm(5.0),'2C3E50'),(f'正确: ___ / {total}',Cm(4.0),'C0392B')]):
        c=t.cell(0,i); c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1); p.paragraph_format.line_spacing=Pt(8)
        r=p.add_run(txt); r.font.size=Pt(8); r.font.bold=True; r.font.color.rgb=RGBColor(*hex_rgb(clr))
        set_shd(c,'F5F5F5')
        set_bdr(c,top={'val':'single','sz':'4','color':'CCC'},bottom={'val':'single','sz':'4','color':'CCC'},
                start={'val':'single','sz':'4','color':'CCC'},end={'val':'single','sz':'4','color':'CCC'})
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)

def sixcol(wl,mode):
    total=len(wl); rn=math.ceil(total/2)
    t=doc.add_table(rows=rn+1,cols=6); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    tbl=t._tbl; tblPr=tbl.tblPr
    if tblPr is None: tblPr=OxmlElement('w:tblPr'); tbl.insert(0,tblPr)
    tblW=OxmlElement('w:tblW'); tblW.set(qn('w:w'),'11000'); tblW.set(qn('w:type'),'dxa'); tblPr.append(tblW)
    h={'en2cn':['#','英文','中文（填空）','#','英文','中文（填空）'],
       'cn2en':['#','中文释义','英文（填空）','#','中文释义','英文（填空）'],
       'phonetic':['#','🔊音标','英+中（填空）','#','🔊音标','英+中（填空）'],
       'mixed':['#','原文','填空','#','原文','填空']}.get(mode,['#','原文','填空','#','原文','填空'])
    for ci in range(6):
        c=t.cell(0,ci); c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
        r=p.add_run(h[ci]); r.font.size=Pt(8); r.font.bold=True; r.font.color.rgb=RGBColor(0x33,0x33,0x33)
        set_shd(c,'EAEAEA')
        set_bdr(c,top={'val':'single','sz':'4','color':'AAA'},bottom={'val':'single','sz':'6','color':'666'},
                start={'val':'single','sz':'4','color':'AAA'},end={'val':'single','sz':'4','color':'AAA'})
        c.width=CW[ci]
    trPr0=t.rows[0]._tr.get_or_add_trPr()
    trH0=OxmlElement('w:trHeight'); trH0.set(qn('w:val'),'400'); trH0.set(qn('w:hRule'),'atLeast'); trPr0.append(trH0)
    for ri in range(rn):
        for col_offset, idx in [(0,ri),(3,ri+rn)]:
            if idx>=total: continue
            w=wl[idx]; sm='★ ' if w['stars']>=4 else ''
            c=t.cell(ri+1,col_offset); c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
            r=p.add_run(str(idx+1)); r.font.size=Pt(7); r.font.color.rgb=RGBColor(0x99,0x99,0x99)
            tcPr=c._tc.get_or_add_tcPr(); vAlign=OxmlElement('w:vAlign'); vAlign.set(qn('w:val'),'center'); tcPr.append(vAlign)
            c=t.cell(ri+1,col_offset+1); c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
            if mode=='en2cn':
                r=p.add_run(f"{sm}{w['word']}"); r.font.size=Pt(10); r.font.bold=True
                r.font.color.rgb=RGBColor(0x1A,0x1A,0x2E) if w['stars']>=4 else RGBColor(0x33,0x33,0x33)
            elif mode=='cn2en': r=p.add_run(w['definition'][:30]); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x44,0x44,0x44)
            elif mode=='phonetic': r=p.add_run(f"🔊{w['phonetic']}"); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x2C,0x3E,0x50)
            elif mode=='mixed':
                if idx%2==0:
                    r=p.add_run(f"{sm}{w['word']}"); r.font.size=Pt(10); r.font.bold=True
                    r.font.color.rgb=RGBColor(0x1A,0x1A,0x2E) if w['stars']>=4 else RGBColor(0x33,0x33,0x33)
                else: r=p.add_run(w['definition'][:28]); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x44,0x44,0x44)
            tcPr=c._tc.get_or_add_tcPr(); vAlign=OxmlElement('w:vAlign'); vAlign.set(qn('w:val'),'center'); tcPr.append(vAlign)
            c=t.cell(ri+1,col_offset+2); c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
            if mode=='en2cn': r=p.add_run(f"({'＿'*14})")
            elif mode=='cn2en': r=p.add_run(f"({'＿'*14})")
            elif mode=='phonetic':
                r=p.add_run(f"英[{'＿'*8}]"); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0xCC,0xCC,0xCC)
                r.add_break(); r=p.add_run(f"中[{'＿'*8}]"); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0xCC,0xCC,0xCC)
            elif mode=='mixed':
                if idx%2==0: r=p.add_run(f"中文[{'＿'*10}]")
                else: r=p.add_run(f"英文[{'＿'*10}]")
            r.font.size=Pt(9); r.font.color.rgb=RGBColor(0xCC,0xCC,0xCC)
            tcPr=c._tc.get_or_add_tcPr(); vAlign=OxmlElement('w:vAlign'); vAlign.set(qn('w:val'),'center'); tcPr.append(vAlign)
            for cix in range(3):
                set_bdr(t.cell(ri+1,col_offset+cix),top={'val':'single','sz':'4','color':'EEE'},bottom={'val':'single','sz':'4','color':'EEE'},
                        start={'val':'single','sz':'4','color':'EEE'},end={'val':'single','sz':'4','color':'EEE'})
        trPr=t.rows[ri+1]._tr.get_or_add_trPr()
        trH=OxmlElement('w:trHeight'); trH.set(qn('w:val'),'600'); trH.set(qn('w:hRule'),'atLeast'); trPr.append(trH)
    for row in t.rows:
        for ci in range(6): row.cells[ci].width=CW[ci]

ENCOURAGE=[(5,['🎉 已学5%！加油，开个好头！','🚀 5%完成！每一步都算数！']),(10,['💪 10%！很不错，继续保持！','📈 十分之一完成！']),(15,['✨ 15%！节奏已经起来了！','🔥 15%！坚持就是胜利！']),(20,['⭐ 20%！五分之一的词汇已拿下！','🚀 20%完成！']),(25,['🌟 四分之一！继续前进！','🎊 25%！里程碑！']),(30,['🔥 30%！越学越顺！','⚡ 接近三分之一了！']),(35,['💎 35%！比你昨天的自己更强！','🌈 35%完成！']),(40,['⚡ 40%！超过大多数人了！','🎯 四成完成！']),(45,['🎯 45%！距离半程一步之遥！','🏔 45%！']),(50,['🎊 过半了！50%拿下！','🥇 半程冠军！','🎆 50%！所有词汇已认识一半！']),(55,['🌈 55%！下半程冲刺！','📊 55%完成率！']),(60,['🏆 60%！量变引起质变！','🚀 六成完成！']),(65,['⭐ 65%！胜利在望！','🌄 65%！']),(70,['🚩 70%！高阶词汇区已进入！','👑 七成！']),(75,['🎪 四分之三！','🏗 75%！词汇大厦即将封顶！']),(80,['👑 80%！词汇达人！','💎 八成！']),(85,['💥 85%！最后冲刺！','🎆 85%！']),(90,['🎆 90%！终点在望！','🏁 九成！']),(95,['🏁 95%！胜利就在眼前！','🚀 最后5%！']),(100,['🥇🎉 100%！全部通关！太棒了！','🎊🏆 完美通关！英语词汇大师！'])]
GENERIC=['📖 第{}组完成！加油！','💪 又一组拿下！','🌟 第{}组记住了！','🏃 第{}组拿下！','🚀 第{}组完成！离目标又近了！']

print(f"Generating {len(groups)} groups...")
cumulative=0
for gi,gwl in enumerate(groups):
    tw=len(gwl); t0=time.time()
    print(f"  G{gi+1}/{len(groups)} ({tw}w)...",end='')
    doc.add_page_break()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f'DAY 1  学习页  · 第{gi+1}/{len(groups)}组'); r.font.size=Pt(14); r.font.bold=True; r.font.color.rgb=RGBColor(0x1A,0x1A,0x2E)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f'苏州中考考纲词汇 · {tw}词'); r.font.size=Pt(8); r.font.color.rgb=RGBColor(0x99,0x99,0x99)
    for w in gwl: card(w)
    doc.add_page_break(); rtitle(f'DAY 2 · 第{gi+1}组','英→中'); ndtable(tw); sixcol(gwl,'en2cn')
    doc.add_page_break(); rtitle(f'DAY 4 · 第{gi+1}组','中→英'); ndtable(tw); sixcol(gwl,'cn2en')
    doc.add_page_break(); rtitle(f'DAY 7 · 第{gi+1}组','🔊 音标'); ndtable(tw); sixcol(gwl,'phonetic')
    sl=[w for w in gwl if w['stars']>=4]
    if sl: doc.add_page_break(); rtitle(f'DAY 15 · 第{gi+1}组',f'★ 高频词 ⭐{len(sl)}'); ndtable(len(sl)); sixcol(sl,'cn2en')
    doc.add_page_break(); rtitle(f'DAY 30 · 第{gi+1}组','🎲 混合终测'); ndtable(tw); sixcol(gwl,'mixed')
    cumulative+=tw
    pct=cumulative*100/1186
    msg=None
    for m,ms in ENCOURAGE:
        if abs(pct-m)<2.0 or (m==100 and pct>=100): msg=eval('random.choice(ms)'); break
    if not msg: msg=eval('random.choice(GENERIC)').format(gi+1)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(6)
    r=p.add_run(msg); r.font.size=Pt(10); r.font.bold=True; r.font.color.rgb=RGBColor(0xC0,0x39,0x2B)
    print(f" {time.time()-t0:.0f}s")

doc.save(OUT)
fsize=os.path.getsize(OUT)/1024
print(f"\n✅ {OUT} ({fsize:.0f}KB) {len(groups)}组")
