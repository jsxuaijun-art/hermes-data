#!/usr/bin/env python3
"""生成初次洽谈SOP Word文档"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ===== 样式设置 =====
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ===== 封面 =====
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('苏州盈信·初次洽谈全流程 SOP 与话术')
run.font.size = Pt(22)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('电话获客 → 面谈诊断 → 跟进成交')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('适用：盈信财税团队  |  版本 v1.0  |  2026.6.6')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ===== 辅助函数 =====
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, italic=False, size=None, color=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    return p

def add_quote(text):
    """添加引用块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.italic = True

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # 数据行
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

def add_divider():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 50)
    run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    run.font.size = Pt(8)

# ===== 正文 =====

# === 核心武器 ===
add_heading('一、核心武器：高级会计师锚点', level=1)
add_para('全流程中必须出现的 3 个锚点：', bold=True)
add_para('① 介绍公司时 → 「我们老板 2018 年就评上高级会计师了，全江苏八九百人里唯一一家代账公司通过的」', indent=True)
add_para('② 聊专业时 → 「不是挂靠大公司的，是以我们小代账公司的名义，自己凭本事评的，含金量比挂靠高多了」', indent=True)
add_para('③ 逼单时 → 「您让同行去找一个高会带队、又做代账的公司比比看」', indent=True)

# === 第1阶段 ===
add_heading('二、第1阶段：电话获客（3-5分钟）', level=1)

add_heading('2.1 目标', level=2)
add_para('✅ 客户接了电话没挂')
add_para('✅ 拿到一个约见面/视频的理由')
add_para('✅ 建立「这人有点东西」的第一印象')

add_heading('2.2 话术模板', level=2)

add_para('模板A：打新公司（刚注册/刚拿执照）', bold=True)
add_quote('「喂，（客户称呼）对吧？我是苏州盈信的（你的名字）。您公司刚下来，通知您一声——执照下来30天内要做税务登记，超期有罚款。您现在税务这块有人弄了没有？」')
add_para('关键：用政策信息开场，不是推销。客户回答后 ↓')
add_quote('「那我跟您说一下流程：税务登记→票种核定→三方协议→每月申报。如果没人弄，这些全部要您自己跑。您现在全职还是兼职在做？」')

add_para('模板B：打老公司（挖同行墙角）', bold=True)
add_quote('「（客户称呼），我是苏州盈信的（你的名字）。今天打电话没别的事，就想问一句——您现在的会计，是什么样的资质？助理会计师还是中级？」')
add_quote('「您知道吗，现在代理记账行业有一个很稀缺的资质——高级会计师。全江苏省八九百人里，唯一以代账公司名义评上高会的，就在我们公司。做账报税这种事情，初级200一个月也能做，高会带队做也是这个价——一样的钱为啥不找专业对口的？」')

add_para('模板C：转介绍/被动咨询', bold=True)
add_quote('「（客户称呼），感谢您给我们这个机会。我先花30秒介绍一下我们是谁——苏州盈信，做了16年财税，老板是高级会计师，2018年全省唯一以代账公司评上的高会。您公司什么情况，我需要了解几个问题，帮您判断匹配不匹配。」')

add_heading('2.3 电话必问6题（需求挖掘）', level=2)
for i, q in enumerate([
    '「您公司什么行业？开了多久了？」',
    '「现在账谁在做？（自己/兼职/代账/全职会计）」',
    '「每个月开票量大吗？一般纳税人还是小规模？」',
    '「最近有没有遇到过税务问题？稽查、预警、补税之类的？」',
    '「您现在每个月付多少服务费？」',
    '「您公司对财税最看重什么？省心？合规？省钱？还是能帮您筹划？」'
], 1):
    add_para(f'{i}. {q}', indent=True)

add_heading('2.4 电话目标转化', level=2)
add_table(
    ['情况', '目标动作'],
    [
        ['客户有兴趣', '→ 「我整理一份您行业的对比方案，明天上午还是下午方便，我到您公司坐15分钟？」'],
        ['客户犹豫', '→ 「那我先把我们公司的资料发您微信，您看看。不管做不做，有问题随时问我。」'],
        ['客户直接拒绝', '→ 「没关系。那我加您个微信，以后税务方面的问题您随时问，免费解答。」'],
    ]
)

# === 第2阶段 ===
add_heading('三、第2阶段：面谈/视频诊断（30-45分钟）', level=1)

add_heading('3.1 开场（3分钟）· 破冰+亮身份', level=2)
add_quote('「（称呼），感谢您抽时间。我先简单介绍一下我自己——盈信财税，（你的名字）。我们公司做了16年，我们老板37岁就评上了高级会计师——不是挂靠大公司的，是以代账公司的名义自己评的。2018年全江苏八九百人，就她一个代账公司出来的高会。」')
add_para('停顿，等对方反应——一般会追问「那挺厉害的，高会难考吗？」')
add_quote('「难。会计从初级到高会，正常十几年。而且高会不是考过就行的，要评审——把你这十几年的业绩、案例、资质全部摆给专家团打分。大部分通过的都是国企财务总监、大学教授、上市公司CFO。代理记账公司能过的，全国都没几个。」')

add_heading('3.2 企业诊断（15分钟）· SPIN提问法', level=2)

add_para('S — 背景问题', bold=True)
add_para('「公司目前主要做什么业务？」', indent=True)
add_para('「团队多少人？有专职财务吗？」', indent=True)
add_para('「现在用什么财务软件？还是手工做账？」', indent=True)

add_para('P — 难点问题', bold=True)
add_para('「现在的账务有没有让你操心的点——比如报税逾期、发票出问题、有些数据对不上？」', indent=True)
add_para('「有没有收到过税务预警？或者被专管员找过？」', indent=True)
add_para('「每个月报表出来你看得懂吗？有没有分析过成本、利润、税负率？」', indent=True)

add_para('I — 暗示问题（放大后果🔥）', bold=True)
add_para('「如果漏报一个税种，金税系统自动预警，轻则补税罚款，重则进异常名录，经营受影响——这一点您了解吗？」', indent=True)
add_para('「报表看不懂，等于您做了半年一年的生意，到底赚没赚钱、哪儿赚钱、哪儿漏钱——完全不清楚。」', indent=True)
add_para('「现在金税四期上线以后，数据都是自动比对的。发票、社保、银行流水都在一张网上，以前能混过去的现在混不过去了。」', indent=True)

add_para('N — 需求-效益问题（让客户自己说价值）', bold=True)
add_para('「如果有一套专业团队帮您盯紧税务、出月报分析、提前预警风险——每个月多花两三百块钱，值得吗？」', indent=True)
add_para('「您最希望解决哪一块？」', indent=True)

add_heading('3.3 方案呈现（10分钟）· FABG法则', level=2)
add_table(
    ['层级', '内容', '话术'],
    [
        ['F（Feature）', '做账+报税+合规+预警+分析', '—'],
        ['A（Advantage）', '双会计交叉审核，高会带队把关', '「别人家一个会计管70家，我们一个会计只管20家，每份账还有高会二审」'],
        ['B（Benefit）', '你不操心、不被罚、看得懂', '「合作之后您就一件事——专心搞业务。税务那边我们全包」'],
        ['G（Grabber）', '高级会计师全流程把关', '「全省代账公司里，有高会带队把关的，你数不出5家」'],
    ]
)

add_heading('3.4 报价（5分钟）· 三档报价法', level=2)
add_para('「我们按照您的需求，有三种方案：', bold=True)
add_para('基础方案——每月xx元。标准代账+报税，适合业务简单的客户。', indent=True)
add_para('标准方案——每月xx元。基础服务+合规审核+季度报表分析+风险预警。大多数客户选这个。', indent=True)
add_para('高端方案——每月xx元。标准全包+高级会计师一对一托管+税务筹划建议+不定期上门。适合业务复杂、对财税重视的客户。', indent=True)
add_para('您先看看哪个方案听着比较顺，咱们细聊。')

add_heading('3.5 常见异议现场处理', level=2)
add_table(
    ['异议', '话术'],
    [
        ['「太贵了」', '「一年平均一天一杯奶茶钱。请一个全职会计至少五六万一年。出一次风险，补税罚款都超过这个数了。」'],
        ['「同行比你们便宜」', '「他是不是高会带队？出问题赔不赔？我们做错账，全额赔付。一分钱一分货。」'],
        ['「先签一个月试试」', '「所有合作一年起签。先签一年，不满意全额退款。不是死板——而是财务服务是按年度规划的。」'],
        ['「和股东商量一下」', '「我把方案整理发您。或者安排个时间我当面给股东讲一遍，现场解答。」'],
        ['「现在零申报用不着」', '「零申报每个季度还是要报。忘了——罚款500起步。金税不认『没业务』，只认有没有按时申报。」'],
    ]
)

# === 第3阶段 ===
add_heading('四、第3阶段：跟进与成交', level=1)

add_heading('4.1 跟进节奏表', level=2)
add_table(
    ['时间节点', '动作', '话术/内容'],
    [
        ['面谈后2小时', '微信发送名片+需求确认', '「xx总，今天聊得很愉快。我整理了一下您的需求……您看看对不对？」'],
        ['次日（第1天）', '电话回访+发补充资料', '「昨天聊完之后我又查了您行业的最新政策，发现……您跟股东聊得怎么样了？」'],
        ['第3天', '政策节点提醒', '「提醒您——本月的申报截止日期是××号。另外最近出了个新优惠……」'],
        ['第7天', '案例驱动跟进', '「最近签了一个同行业的客户。给您看看他之前遇到的问题——跟您情况很像。」'],
        ['第14天', '最后真诚跟进', '「不管您选不选我们，之前发的政策信息您留着用。以后有任何财税问题，随时问我，免费解答。」'],
    ]
)

add_heading('4.2 逼单7法（按场景选用）', level=2)
add_table(
    ['方法', '适用场景', '话术'],
    [
        ['二选一法', '客户犹豫选方案', '「您觉得是基础方案够用，还是标准方案更放心？」'],
        ['稀缺性法', '一直拖着', '「每个会计服务名额有限，目前对得上您的会计只剩最后一个名额了。」'],
        ['从众法', '不相信小公司', '「我们合作了上千家客户，光您这个行业就一百多家。」'],
        ['案例法', '担心风险/质量', '「上个月一个做贸易的客户，之前代账少报了十几万收入，金税预警了。我们帮他一周解决。」'],
        ['损失厌恶法', '觉得不急', '「这个价格是季度优惠价，下周一调回原价。」'],
        ['高会锚定法', '还在比价', '「您问任何一家同行——问他有没有高会带队，一问就知道了。」'],
        ['假定成交法', '聊得差不多了', '「我先帮您把资料建好，明天安排会计跟您对接。」'],
    ]
)

add_heading('4.3 成交后的动作', level=2)
add_para('「xx总，合作愉快。明天我安排您的专属会计加您微信，先做一次全面的账务体检。之后每个月的申报情况、每季度报表分析，我都会同步发给您。」')
add_para('成交后24小时内必做：', bold=True)
for i, item in enumerate([
    '拉专属服务群（客户+你+会计+主管）',
    '发《服务流程说明》',
    '做首次账务交接/资料收集清单',
    '发朋友圈（脱敏版）：「欢迎xx总加入盈信大家庭」'
], 1):
    add_para(f'{i}. {item}', indent=True)

# === 辅助工具 ===
add_heading('五、辅助工具', level=1)

add_heading('5.1 材料准备清单（面谈前检查）', level=2)
add_table(
    ['类型', '材料', '用途'],
    [
        ['公司介绍', '公司简介视频/图文', '破冰+信任'],
        ['高会锚点', '高级会计师证书照片/文章', '核心差异化'],
        ['案例', '2-3个同行业客户案例', '从众+证明'],
        ['方案', '三档报价单', '清晰报价'],
        ['证书', 'TSC五级/营业执照', '资质证明'],
        ['合同', '标准服务合同', '当场签约'],
    ]
)

add_heading('5.2 核心数据速记（脱口就能说）', level=2)
for item in [
    '「16年——我们做了16年」',
    '「高级会计师——2018年全省唯一代账公司评上的」',
    '「1000+客户——90%是老客户介绍来的」',
    '「TSC五级——税务局给的最高信用评级，全行业没几家」',
    '「双城——苏州上海都有公司」',
]:
    add_para(item, indent=True)

# === 避坑指南 ===
add_heading('六、避坑指南', level=1)
add_table(
    ['常见错误', '正确做法'],
    [
        ['电话里报全价', '电话只透露大致范围，详细报价见面/视频聊'],
        ['过早让步降价', '先价值重塑再谈价格，不要一上来就送'],
        ['客户说贵就解释', '不要解释——反问「跟什么比觉得贵？」'],
        ['比价格只比价格', '对比专业、资质、服务配置'],
        ['客户说再考虑就真等', '定好下次联系时间，不给晾着'],
        ['只聊产品不聊痛点', '80%时间挖痛点，20%时间给方案'],
        ['忘了亮高会身份', '开场必须亮，中间回扣，逼单再用'],
    ]
)

# ===== 保存 =====
output_path = '/mnt/c/Users/Admin/Desktop/初次洽谈SOP与话术_完整版.docx'
doc.save(output_path)
print(f'✅ Saved to {output_path}')
