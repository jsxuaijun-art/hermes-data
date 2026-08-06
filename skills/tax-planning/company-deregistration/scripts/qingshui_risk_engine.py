#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
注销清税智能体 v2.0 —— 科目余额表 / 财务报表 解析引擎
=====================================================
苏州盈信企业管理有限公司 | 高级会计师江敏 16年实务

功能：
  1. 自动读取「科目余额表.xlsx」—— 识别科目编码/名称/期末余额
  2. 自动读取「资产负债表 / 利润表」—— 识别关键科目
  3. 对照注销雷区规则 → 输出「注销雷区诊断报告」+ 处理建议 + 引流文案钩子

用法：
  python 注销清税解析引擎_v2.py --科目余额表 路径.xlsx
  python 注销清税解析引擎_v2.py --资产负债表 路径.xlsx --利润表 路径.xlsx
  python 注销清税解析引擎_v2.py --科目余额表 a.xlsx --自动诊断   # 一键诊断
"""
import sys, os, re
from collections import OrderedDict
try:
    import openpyxl
except ImportError:
    sys.exit("需要 openpyxl: ~/.venv-hermes/bin/pip install openpyxl")

# ═══════════════════════════════════════════
# 雷区规则库 v2.0（科目关键词 → 风险+处理+红线）
# ═══════════════════════════════════════════
RISK_RULES = OrderedDict()

def _rule(key, keywords, risk, handle, redline, level="🔴", trigger_min=0.0, ratio_ref=None):
    RISK_RULES[key] = {
        "keywords": keywords, "risk": risk, "handle": handle,
        "redline": redline, "level": level,
        "trigger_min": trigger_min, "ratio_ref": ratio_ref,
    }

# ① 应收账款过大
_rule("应收账款", ["应收账款", "应收帐款"],
      "收不回仍挂账 → 虚增清算所得，影响清税证明",
      "坏账核销须备齐25号公告证据(破产裁定/生效判决/超3年账龄+催收记录)",
      "无证据硬核销 → 税局不认，视同未分配利润", "🔴")

# ② 存货过大
_rule("存货", ["存货", "库存商品", "原材料", "在产品", "生产成本", "产成品"],
      "账实不符/长期不卖 → 注销重点审查，最高罚5倍",
      "三选一:变现(交增值税)/分配股东(视同销售按组成计税价)/报废(非正常损失进项转出)",
      "管理不善致存货损失 → 进项税额必须转出", "🔴")

# ③ 未分配利润高
_rule("未分配利润", ["未分配利润"],
      "清算分配 → 自然人股东涉20%个税",
      "清算所得先缴企税，剩余财产分配对股东补个税",
      "不能靠注销规避个税，税局按清算分配核定", "🔴")

# ④ 股东/关联借款（记在其他应收款借方，或公司向股东借款记其他应付款贷方）
_rule("股东借款", ["股东借款", "法人借款", "老板", "股东", "往来款"],
      "股东借款不还 → 视同分红",
      "注销前结清或作分配处理",
      "12月底前未还的股东借款视同利润分配", "🟡")

# ④b 其他应收款（股东/个人借款、保证金藏匿地）
_rule("其他应收款", ["其他应收款"],
      "股东/关联方个人借款挂账 → 注销时视同利润分配补个税",
      "核查明细：个人借款须注销前归还或作分配；保证金/押金应退回应退",
      "其他应收款挂个人大额借款 → 税局按股息红利核定20%个税", "🔴")

# ⑤ 固定资产
_rule("固定资产", ["固定资产"],
      "变卖/分配未缴增值税",
      "对外处置交增值税，清算确认所得",
      "固定资产净值≠账面残值，处置价偏低会被核定", "🟡")

# ⑥ 货币资金大（现金+银行存款）
_rule("货币资金", ["货币资金", "银行存款", "库存现金", "现金", "其他货币资金"],
      "余额较大 → 是清算财产，须分配；现金余额异常大易被疑账外经营/私户",
      "清算时作为剩余财产分配，正常不额外涉税（分配环节股东补个税）",
      "现金巨量且无业务支撑 → 疑私账收款未入账，先自查", "🟡")

# ⑦ 应付账款有余额（无法清偿）
_rule("应付账款", ["应付账款", "应付帐款", "其他应付款"],
      "无法清偿的应付 → 须转入营业外收入缴企业所得税",
      "注销前核实能否清偿；确实无需支付的应付转营业外收入（25%企税）",
      "长期挂账无法清偿的应付，税局可能按营业外收入核定补税", "🟡")

# ⑧ 预收账款
_rule("预收账款", ["预收账款", "预收帐款", "合同负债"],
      "预收未履约 → 注销时须确认收入或退款",
      "已履约的确认收入缴税；未履约的退款冲销",
      "长期挂预收不处理 → 税局疑隐匿收入", "🟡")

# ⑨ 长期待摊/在建/无形资产
_rule("长期资产", ["长期待摊费用", "在建工程", "无形资产", "商誉", "长期股权投资", "长期应收款"],
      "未摊销/未处置 → 清算需作价或核销",
      "可处置变现或作废处理，长期股权投资需清理子公司",
      "长期股权投资不清理 → 注销受阻(有子公司须先注销子公司)", "🟡")

# ⑩ 应交税费科目（正常结转 vs 有余额）
_rule("应交税费", ["应交税费", "应交税金"],
      "贷方余额=未缴税，借方余额=留抵/多缴",
      "贷方须结清税款；借方留抵可申请退税或结转",
      "贷方余额未清 → 无法取得清税证明", "🔴")

# ⑪ 应付职工薪酬
_rule("应付职工薪酬", ["应付职工薪酬", "应付工资", "应付福利费"],
      "欠付员工工资/社保 → 清算优先清偿",
      "注销前结清职工工资、社保",
      "欠薪注销 → 员工维权+清算组责任", "🟡")

# ═══════════════════════════════════════════
# 工具函数
# ═══════════════════════════════════════════
def _clean(v):
    """清洗单元格数字"""
    if v is None:
        return 0.0
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).replace(",", "").replace("，", "").strip()
    if not s:
        return 0.0
    try:
        return float(s)
    except ValueError:
        return 0.0

def _find_header_row(ws, max_col):
    """定位表头行：优先找含'科目名称/科目编码'且其后出现'借方/贷方'子行的行"""
    # 候选：含"科目编码"或"科目名称"
    candidates = []
    for r in range(1, min(ws.max_row, 10) + 1):
        row_vals = [str(ws.cell(r, c).value or "") for c in range(1, max_col + 1)]
        joined = "".join(row_vals)
        has_name = ("科目名称" in joined or "科目编码" in joined)
        has_fx = ("借方" in joined or "贷方" in joined)
        if has_name:
            candidates.append((r, has_fx))
    if candidates:
        # 优先选同时含借贷方向词的表头行
        for r, has_fx in candidates:
            if has_fx:
                return r
        # 否则看下一行是否含"借方/贷方"子行
        for r, _ in candidates:
            nxt = [str(ws.cell(r+1, c).value or "") for c in range(1, max_col + 1)]
            if "借方" in "".join(nxt) or "贷方" in "".join(nxt):
                return r
        return candidates[0][0]
    return None

def parse_subject_balance(path):
    """解析科目余额表(xlsx/xls)，返回 {科目名称: 期末余额(借贷净值)} + 明细"""
    if path.lower().endswith(".xls"):
        return parse_subject_balance_xls(path)
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.active
    max_col = ws.max_column
    header_row = _find_header_row(ws, max_col)
    if not header_row:
        header_row = 1
    cols = {"name": None, "code": None, "debit_end": None, "credit_end": None}
    for c in range(1, max_col + 1):
        h = str(ws.cell(header_row, c).value or "").strip()
        if "科目名称" in h or h == "科目":
            cols["name"] = c
        elif "科目编码" in h or "科目编号" in h:
            cols["code"] = c
        elif "期末余额" in h:
            if "借方" in h: cols["debit_end"] = c
            elif "贷方" in h: cols["credit_end"] = c
        elif "借方余额" in h or ("借方" in h and "期末" in h):
            cols["debit_end"] = c
        elif "贷方余额" in h or ("贷方" in h and "期末" in h):
            cols["credit_end"] = c
    if not cols["name"]:
        raise ValueError("未找到科目名称列")
    # 若表头行下方有子行(如"借方|贷方"第二行)，数据从再下一行开始，且借贷列以子行为准
    data_start = header_row + 1
    subvals = [str(ws.cell(data_start, c).value or "").strip() for c in range(1, max_col + 1)]
    if "借方" in "".join(subvals) or "贷方" in "".join(subvals):
        # 借/贷列从子行识别（表头行是合并表头）
        cols["debit_end"] = cols["credit_end"] = None
        for c in range(1, max_col + 1):
            sv = subvals[c-1]
            if sv == "借方" and cols["debit_end"] is None:
                cols["debit_end"] = c
            elif sv == "贷方" and cols["credit_end"] is None:
                cols["credit_end"] = c
        data_start = header_row + 2
    subjects = OrderedDict()
    detail = []
    for r in range(data_start, ws.max_row + 1):
        name = str(ws.cell(r, cols["name"]).value or "").strip()
        if not name or name.startswith("合计") or name.startswith("总计"):
            continue
        debit = _clean(ws.cell(r, cols["debit_end"]).value) if cols["debit_end"] else 0.0
        credit = _clean(ws.cell(r, cols["credit_end"]).value) if cols["credit_end"] else 0.0
        code = str(ws.cell(r, cols["code"]).value or "").strip() if cols["code"] else ""
        is_top = (len(code) == 4) if code else True
        net = debit - credit
        if is_top and name not in subjects:
            subjects[name] = net
        detail.append({"code": code, "name": name, "debit_end": debit, "credit_end": credit, "net": net})
    return subjects, detail, cols

def parse_subject_balance_xls(path):
    """解析 xls 老格式科目余额表（表头合并行 + 借方/贷方子行）"""
    import xlrd
    wb = xlrd.open_workbook(path)
    ws = wb.sheet_by_index(0)
    header_row = None
    for r in range(ws.nrows):
        rowvals = "".join(str(ws.cell_value(r, c)) for c in range(ws.ncols))
        if "科目编码" in rowvals or "科目名称" in rowvals:
            header_row = r
            break
    if header_row is None:
        raise ValueError("xls科目余额表未找到表头")
    data_start = header_row + 1
    nxt = "".join(str(ws.cell_value(data_start, c)) for c in range(ws.ncols))
    if "借方" in nxt or "贷方" in nxt:
        data_start = header_row + 2
    cols = {"name": None, "code": None, "debit_end": None, "credit_end": None}
    for c in range(ws.ncols):
        h = str(ws.cell_value(header_row, c)).strip()
        if "科目名称" in h or h == "科目":
            cols["name"] = c
        elif "科目编码" in h or "科目编号" in h:
            cols["code"] = c
        elif "借方" in h:
            cols["debit_end"] = c
        elif "贷方" in h:
            cols["credit_end"] = c
    if not cols["name"]:
        cols["name"] = 1; cols["code"] = 0
        cols["debit_end"] = ws.ncols - 2; cols["credit_end"] = ws.ncols - 1
    subjects = OrderedDict()
    detail = []
    for r in range(data_start, ws.nrows):
        name = str(ws.cell_value(r, cols["name"])).strip()
        if not name or name.startswith("合计") or name.startswith("总计"):
            continue
        debit = float(ws.cell_value(r, cols["debit_end"]) or 0)
        credit = float(ws.cell_value(r, cols["credit_end"]) or 0)
        code = str(ws.cell_value(r, cols["code"])).strip() if cols["code"] else ""
        is_top = (len(code) == 4) if code else True
        net = debit - credit
        if is_top and name not in subjects:
            subjects[name] = net
        detail.append({"code": code, "name": name, "debit_end": debit, "credit_end": credit, "net": net})
    return subjects, detail, cols

def parse_bs_and_pl(bs_path=None, pl_path=None):
    """解析资产负债表/利润表（xlsx/xls，两表合一）"""
    result = {}
    for label, path in [("资产负债表", bs_path), ("利润表", pl_path)]:
        if not path:
            continue
        if path.lower().endswith(".xls"):
            import xlrd
            wb = xlrd.open_workbook(path)
            for ws in wb.sheets():
                # 只处理资产负债表和利润表，排除现金流量表（行名会误匹配雷区关键词）
                if "资产负债表" not in ws.name and "利润表" not in ws.name:
                    continue
                # 资产负债表为左右双栏：A名称|B行次|C期末|D年初|E负债名|F行次|G期末|H年初
                # 左栏取C(期末)，右栏取G(期末)；利润表取"本年累计金额"
                is_bs = "资产负债表" in ws.name
                for r in range(ws.nrows):
                    # 左栏
                    name = str(ws.cell_value(r, 0)).strip()
                    if name and name not in ("资产", "流动资产：", "非流动资产：", "资产负债表"):
                        if is_bs and ws.ncols >= 4 and isinstance(ws.cell_value(r, 2), (int, float)):
                            result[name] = float(ws.cell_value(r, 2))
                        elif not is_bs and ws.ncols >= 3 and isinstance(ws.cell_value(r, 2), (int, float)):
                            result[name] = float(ws.cell_value(r, 2))
                    # 右栏（负债和所有者权益）
                    if is_bs and ws.ncols >= 7:
                        rname = str(ws.cell_value(r, 4)).strip()
                        if rname and rname not in ("负债和所有者权益", "流动负债：", "非流动负债："):
                            if isinstance(ws.cell_value(r, 6), (int, float)):
                                result[rname] = float(ws.cell_value(r, 6))
            continue
        wb = openpyxl.load_workbook(path, data_only=True)
        for ws in wb.worksheets:
            for r in range(1, ws.max_row + 1):
                name = str(ws.cell(r, 1).value or "").strip()
                vals = [_clean(ws.cell(r, c).value) for c in range(2, ws.max_column + 1)]
                non_zero = [v for v in vals if abs(v) > 0.01]
                if non_zero:
                    result[name] = non_zero[-1]
    return result

# ═══════════════════════════════════════════
# 雷区诊断引擎
# ═══════════════════════════════════════════
def diagnose(subjects, bs_pl=None, thresholds=None):
    """匹配雷区规则，返回诊断结果列表"""
    if bs_pl is None:
        bs_pl = {}
    if thresholds is None:
        thresholds = {"应收": 50000, "存货": 50000, "货币资金": 100000,
                      "应付": 50000, "未分配利润": 0, "股东借款": 10000}
    findings = []
    # 合并科目余额表 + 报表数据（科目余额表更细、更权威，覆盖报表）
    pool = {}
    pool.update(bs_pl)      # 先用报表
    pool.update(subjects)   # 科目余额表覆盖（应付/应交/往来款以科目表为准）
    # 总资产（用于占比）——报表若无，从科目表资产合计取
    total_assets = 0
    for k, v in pool.items():
        if "资产总计" in k or "资产合计" in k:
            total_assets = abs(v)

    for key, rule in RISK_RULES.items():
        matched = []
        for kw in rule["keywords"]:
            for name, val in pool.items():
                if kw in name and abs(val) > 0.01:
                    matched.append((name, val))
        if not matched:
            continue
        # 取绝对值最大的
        name, val = max(matched, key=lambda x: abs(x[1]))
        absv = abs(val)
        # 判断是否触发（按默认阈值，简单规则）
        trigger = True
        note = f"{name} = {val:,.2f} 元"
        findings.append({
            "rule": key, "level": rule["level"],
            "subject": name, "value": val, "note": note,
            "risk": rule["risk"], "handle": rule["handle"], "redline": rule["redline"],
        })
    return findings, total_assets

# ═══════════════════════════════════════════
# 清算所得税测算（财税〔2009〕60号）
# ═══════════════════════════════════════════
def calc_qingshui_tax(pool, discount_rate=1.0, liq_cost=None, corp_tax_rate=0.25, div_tax_rate=0.20, bs=None):
    """
    依据财税〔2009〕60号测算注销清算涉税额。
    公式：①清算所得=可变现价值−计税基础−清算费用−相关税费+债务清偿损益
          ②清算所得税=清算所得×25%
          ③剩余财产=可变现价值−清算费用−工资社保−清算所得税−欠税−清偿债务
          ④股息所得=(未分配利润+盈余公积)→20%个税
          ⑤转让所得=剩余财产−股息−投资成本→20%个税
    bs: 资产负债表权威数据（期末总额口径），优先于此；pool为兜底。
    """
    def get(*keys, base=None):
        src = base if base is not None else pool
        for k in keys:
            for name, v in src.items():
                if k in name and abs(v) > 0.01:
                    return v
        return 0.0

    # 资产端（注意科目余额表里贷方科目净额为负；这里统一取绝对值逻辑）
    # 总资产：优先用资产负债表权威口径（bs），pool兜底
    total_assets = get("资产总计", "资产合计", base=bs) or get("资产总计", "资产合计")
    if not total_assets:
        # 无总资产时，从科目余额表累加资产类（编码1开头）
        total_assets = sum(v for k, v in pool.items() if k[:1] in "12" or "货币资金" in k or "应收账款" in k or "存货" in k or "固定资产" in k or "库存" in k)

    net_assets_book = abs(total_assets)                 # 计税基础≈账面净值（简化）
    realizable = net_assets_book * discount_rate        # 可变现价值（可按折扣调整）

    # 债务清偿损益：无法清偿的负债（应付/预收/其他应付款）→ 视同清偿收益
    # 用报表期末总额（资产负债表权威口径，负债为正数贷方；科目余额表为负）
    unavoidable_liab = 0.0
    for key in ["应付账款", "预收账款", "其他应付款", "应付帐款", "预收帐款"]:
        v = get(key, base=bs)
        if v:
            # bs报表为正数贷方，科目表为负数；统一取绝对值作为无法清偿负债
            if bs is not None and key in [k for k in bs if key in k]:
                unavoidable_liab += abs(v)
            elif v < 0:
                unavoidable_liab += abs(v)

    # 清算费用：默认按可变现价值的2%，可覆盖
    if liq_cost is None:
        liq_cost = realizable * 0.02
    # 相关税费（简化：增值税等按存货/固定资产处置估算，此处留0占位，可由用户补）
    related_tax = 0.0

    # ① 清算所得
    qing_suo_de = realizable - net_assets_book - liq_cost - related_tax + unavoidable_liab
    # ② 清算所得税
    qing_suo_tax = max(0.0, qing_suo_de) * corp_tax_rate

    # ③ 剩余财产
    salary_welfare = abs(get("应付职工薪酬", "应付工资", base=bs) or get("应付职工薪酬", "应付工资"))   # 工资社保优先清偿
    unpaid_tax = abs(get("应交税费", "应交税金", base=bs) or get("应交税费", "应交税金"))            # 欠税（贷方）
    normal_liab = abs(get("负债合计", base=bs)) if get("负债合计", base=bs) else (unavoidable_liab + salary_welfare + unpaid_tax)
    surplus = realizable - liq_cost - salary_welfare - qing_suo_tax - unpaid_tax - (normal_liab - unavoidable_liab)

    # ④ 股息所得 + 个税
    retained = abs(get("未分配利润", "利润分配", base=bs) or get("未分配利润", "利润分配")) + abs(get("盈余公积", "盈余公积"))
    div_tax_base = retained
    div_tax = div_tax_base * div_tax_rate
    # ⑤ 转让所得 + 个税
    invest_cost = abs(get("实收资本", "实收资本", "股本", base=bs) or get("实收资本", "实收资本", "股本"))
    transfer_income = max(0.0, surplus - div_tax_base - invest_cost)
    transfer_tax = transfer_income * div_tax_rate

    total_shareholder_tax = div_tax + transfer_tax

    return {
        "realizable": realizable, "net_book": net_assets_book, "liq_cost": liq_cost,
        "unavoidable_liab": unavoidable_liab, "qing_suo_de": qing_suo_de,
        "qing_suo_tax": qing_suo_tax, "salary_welfare": salary_welfare,
        "unpaid_tax": unpaid_tax, "surplus": surplus,
        "div_tax_base": div_tax_base, "div_tax": div_tax,
        "transfer_income": transfer_income, "transfer_tax": transfer_tax,
        "total_shareholder_tax": total_shareholder_tax,
        "corp_tax_rate": corp_tax_rate, "div_tax_rate": div_tax_rate,
        "discount_rate": discount_rate,
    }

# ═══════════════════════════════════════════
# 报告输出
# ═══════════════════════════════════════════
def gen_docx(findings, total_assets, out_path, tax=None):
    """生成 Word 注销雷区诊断报告"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = '微软雅黑'; st._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    st.font.size = Pt(10.5)
    def set_ea(run, name='微软雅黑'):
        run.font.name = name
        rpr = run._element.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = rpr.makeelement(qn('w:rFonts'), {})
            rpr.append(rf)
        rf.set(qn('w:eastAsia'), name)
    def H(t, s=16, c=(0x1F,0x49,0x7D), center=False):
        p = doc.add_paragraph(); 
        if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(t); r.bold = True; r.font.size = Pt(s); r.font.color.rgb = RGBColor(*c)
        set_ea(r)
    def P(t, s=10.5, c=None, b=False):
        p = doc.add_paragraph(); r = p.add_run(t); r.font.size = Pt(s); r.bold = b
        if c: r.font.color.rgb = RGBColor(*c)
        set_ea(r)
    H("公司注销 · 税务清税 雷区诊断报告", 20, (0x1F,0x49,0x7D), True)
    P("苏州盈信企业管理有限公司 · 高级会计师江敏", 10.5, (0x99,0x99,0x99), True)
    P("基于科目余额表/财务报表自动解析生成", 9, (0x99,0x99,0x99))
    P("")
    H("一、诊断结论")
    if not findings:
        P("未发现明显注销雷区科目。", c=(0x00,0x80,0x00), b=True)
    else:
        for i, f in enumerate(findings, 1):
            P(f"[{i}] {f['level']} {f['rule']}（{f['note']}）", 10.5, (0xCC,0x00,0x00) if "🔴" in f["level"] else (0xCC,0x66,0x00), True)
    if total_assets:
        P(f"报表总资产约 {total_assets:,.2f} 元", 9, (0x99,0x99,0x99))
    P("")
    H("二、清算所得税测算（财税〔2009〕60号）")
    if tax:
        def money(v):
            return f"{v:,.0f}" if abs(v) >= 1 else f"{v:,.2f}"
        P(f"① 全部资产可变现价值（按账面 {tax['discount_rate']:.0%}）: {money(tax['realizable'])} 元")
        P(f"② 清算费用（估算 {tax['liq_cost']:.0f} 元）")
        P(f"③ 无法清偿负债（应付/预收/其他应付）视同清偿收益: {money(tax['unavoidable_liab'])} 元")
        P(f"④ 清算所得 = 可变现价值 − 计税基础 − 清算费用 + 债务清偿损益 = {money(tax['qing_suo_de'])} 元", b=True)
        P(f"⑤ 清算所得税 = 清算所得 × {tax['corp_tax_rate']:.0%} = {money(tax['qing_suo_tax'])} 元 🔴", c=(0xCC,0x00,0x00), b=True)
        P(f"⑥ 剩余财产（清偿工资社保/欠税/债务后）: {money(tax['surplus'])} 元")
        P(f"⑦ 股息所得（未分配利润+盈余公积）: {money(tax['div_tax_base'])} 元 → 股东个税 {money(tax['div_tax'])} 元 🔴", c=(0xCC,0x00,0x00), b=True)
        P(f"⑧ 转让所得: {money(tax['transfer_income'])} 元 → 转让个税 {money(tax['transfer_tax'])} 元")
        P(f"★ 注销合计涉税（清算所得税+股东个税）: {money(tax['qing_suo_tax']+tax['total_shareholder_tax'])} 元 🔴", c=(0xCC,0x00,0x00), b=True)
        P("注：可变现价值按账面净值简化估算，相关税费/增值税未计；实际以处置价格和税局核定为准。", 8, (0x99,0x99,0x99))
        P("")
    H("三、各雷区处理建议")
    for f in findings:
        H(f"{f['level']} {f['rule']}", 12, (0x1F,0x49,0x7D))
        P(f"风险：{f['risk']}")
        P(f"处理：{f['handle']}")
        P(f"红线：{f['redline']}", 9, (0xCC,0x00,0x00))
        P("")
    H("四、引流文案钩子（基于本报告发现）")
    for f in findings[:3]:
        P(f"▸ 你账上的「{f['subject']}」余额 {abs(f['value']):,.0f} 元，注销前不处理，可能要多缴一笔冤枉税。私信我帮你做清算体检。")
    P("")
    P("⚡ 本报告基于报表自动解析生成，实际执行前请结合原始凭证进一步核实。", 9, (0x99,0x99,0x99))
    doc.save(out_path)
    return out_path

# ═══════════════════════════════════════════
# main
# ═══════════════════════════════════════════
def main():
    args = sys.argv[1:]
    if not args or args[0] in ("-h", "--help"):
        print(__doc__)
        return
    subjects = OrderedDict()
    bs_pl = {}
    auto = "--自动诊断" in args
    files = [a for a in args if a.lower().endswith((".xlsx", ".xls"))]
    mode = None
    for a in args:
        if a.startswith("--") and a not in ("--自动诊断",):
            mode = a[2:]
    # 模式判定：若未显式指定科目余额表模式，则按文件名自动分流
    #   文件名含"科目余额表" → 科目余额表；否则 → 财务报表
    for f in files:
        base = os.path.basename(f).lower()
        is_kemu = (mode == "科目余额表") or ("科目余额表" in base)
        if is_kemu:
            try:
                s, d, cols = parse_subject_balance(f)
                subjects.update(s)
                print(f"[解析] 科目余额表 {f} → {len(s)}个顶级科目")
            except Exception as e:
                print(f"[错误] 解析 {f}: {e}")
        else:
            bs_pl.update(parse_bs_and_pl(f, f))
    findings, total_assets = diagnose(subjects, bs_pl)
    # 合并数据用于清算所得税测算（bs优先用报表权威口径，pool兜底）
    pool = {}
    pool.update(bs_pl)
    pool.update(subjects)
    tax = calc_qingshui_tax(pool, bs=bs_pl)
    print("=" * 50)
    print("🔍 注销雷区诊断结果")
    print("=" * 50)
    if not findings:
        print("✅ 未发现明显雷区")
    for f in findings:
        print(f"{f['level']} {f['rule']} | {f['note']}")
        print(f"    处理: {f['handle']}")
    print("=" * 50)
    print(f"💰 清算所得税测算（财税60号）:")
    print(f"    清算所得 = {tax['qing_suo_de']:,.2f} 元 → 清算所得税 = {tax['qing_suo_tax']:,.2f} 元")
    print(f"    股东个税（股息+转让） = {tax['total_shareholder_tax']:,.2f} 元")
    print(f"    ★ 注销合计涉税 = {tax['qing_suo_tax']+tax['total_shareholder_tax']:,.2f} 元")
    if findings or True:
        out = "/mnt/c/Users/Admin/Desktop/注销雷区诊断报告.docx"
        try:
            gen_docx(findings, total_assets, out, tax)
            print(f"\n📄 报告已生成: {out}")
        except Exception as e:
            print(f"\n[docx生成失败] {e}")

if __name__ == "__main__":
    main()
