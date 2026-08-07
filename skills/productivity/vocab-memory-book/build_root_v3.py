# -*- coding: utf-8 -*-
"""
高考英语词汇默写本 · 词根版v3
- "#"列 0.35cm，余宽均分给其他列
- Part 1: 词根词汇 50词一组（不同词根可混合），完整6轮
- Part 2: 基础词汇 50词一组，4轮复习(无学习/无终测)
- 页码: "2/365" 格式
- 各组最后一句鼓励语
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import re, math, os, time, random

SRC = r"/mnt/c/Users/Admin/Desktop/高考英语词汇表3817词.txt"
FREQ = r"/home/dmin/.hermes/skills/productivity/vocab-memory-book/en_50k.txt"
OUT = r"/mnt/c/Users/Admin/Desktop/高考英语词汇默写本_词根版.docx"
WPG = 50  # Words Per Group

# ═══ ROOTS ═══
RAW_ROOTS = open('/tmp/build_root_v2.py','r').read()
start = RAW_ROOTS.index('RAW_ROOTS = """') + len('RAW_ROOTS = """')
end = RAW_ROOTS.index('"""', start)
root_text = RAW_ROOTS[start:end]
ROOTS = []
for line in root_text.strip().split('\n'):
    parts = line.split('|', 1)
    if len(parts) == 2: ROOTS.append((parts[0].strip(), parts[1].strip()))
ROOTS.sort(key=lambda x: -len(x[0]))

def detect_root(word):
    w = word.lower()
    for root, meaning in ROOTS:
        if root in w:
            if len(root) <= 2 and len(w) <= 4: continue
            return (root, meaning)
    return None

# ═══ PARSE ═══
print("Loading vocab...")
all_words = []
with open(SRC, 'r', errors='replace') as f:
    lines = f.read().strip().split('\n')
for l in lines[2:]:
    l = l.strip()
    if not l: continue
    m = re.match(r'([a-zA-Z\-\(\)]+)\s', l)
    if not m: continue
    word = m.group(1).lower().rstrip(',)')
    ph = re.search(r'\[([^\]]+)\]', l)
    phon = ph.group(1) if ph else ''
    rest = l[m.end():].strip()
    if phon: rest = rest.replace(f'[{phon}]','',1).strip()
    ri = detect_root(word)
    all_words.append({
        'word': word, 'phonetic': phon, 'definition': rest[:90],
        'root': ri[0] if ri else '', 'root_meaning': ri[1] if ri else '',
    })
print(f"Total: {len(all_words)} words")

# ═══ FREQ ═══
print("Loading frequency...")
freq_data = {}
with open(FREQ, 'r') as f:
    for l in f:
        l = l.strip()
        if not l: continue
        parts = l.split(None, 1)
        if len(parts) == 2: freq_data[parts[0].lower()] = int(parts[1])
fv = sorted(freq_data.values(), reverse=True)
p80 = fv[max(0,len(fv)//5*4)] if fv else 0
p60 = fv[max(0,len(fv)//5*3)] if fv else 0
p40 = fv[max(0,len(fv)//5*2)] if fv else 0
p20 = fv[max(0,len(fv)//5*1)] if fv else 0
def stars(w):
    f = freq_data.get(w)
    if f is None: return 1
    if f >= p80: return 5
    if f >= p60: return 4
    if f >= p40: return 3
    if f >= p20: return 2
    return 1
for w in all_words: w['stars'] = stars(w['word'])

# ═══ GROUPING ═══
root_words = [w for w in all_words if w['root']]
basic_words = [w for w in all_words if not w['root']]
random.seed(20260704)
random.shuffle(root_words)
random.shuffle(basic_words)

# Group root words into chunks of WPG (keep root-tagged)
root_groups = []
i = 0
while i < len(root_words):
    chunk = root_words[i:i+WPG]
    # Get the dominant root(s) for display
    roots_in_chunk = list(set(w['root'] for w in chunk))
    root_groups.append({'label': '+'.join(roots_in_chunk[:3]) + ('...' if len(roots_in_chunk)>3 else ''),
                        'roots': roots_in_chunk, 'words': chunk})
    i += WPG

# Group basic words into chunks of WPG
basic_groups = []
i = 0
while i < len(basic_words):
    chunk = basic_words[i:i+WPG]
    basic_groups.append({'words': chunk})
    i += WPG

print(f"Root: {len(root_words)}w → {len(root_groups)} groups")
print(f"Basic: {len(basic_words)}w → {len(basic_groups)} groups")
total_groups = len(root_groups) + len(basic_groups)
total_root_words = len(root_words)
total_basic_words = len(basic_words)
TOTAL_WORDS = len(all_words)

# ═══════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════
def set_shd(c, color):
    c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))
def set_bdr(c, **kw):
    tc=c._tc; p=tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for edge,val in kw.items():
        e=OxmlElement(f'w:{edge}'); e.set(qn('w:val'),val.get('val','single'))
        e.set(qn('w:sz'),val.get('sz','4')); e.set(qn('w:color'),val.get('color','000000'))
        e.set(qn('w:space'),'0'); b.append(e)
    p.append(b)
def hex_rgb(h): return int(h[0:2],16),int(h[2:4],16),int(h[4:6],16)

# ═══ COVER ═══
def build_cover(doc, rt, bt, rr, br):
    bg = doc.add_table(rows=1, cols=1); bg.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = bg._tbl; tblPr = tbl.tblPr
    if tblPr is None: tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    for child in list(tblPr):
        if child.tag == qn('w:tblW'): tblPr.remove(child)
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'),'11906'); tblW.set(qn('w:type'),'dxa'); tblPr.append(tblW)
    cell = bg.cell(0,0); set_shd(cell, '0A2463')
    set_bdr(cell,top={'val':'none','sz':'0','color':'auto'},bottom={'val':'none','sz':'0','color':'auto'},
            start={'val':'none','sz':'0','color':'auto'},end={'val':'none','sz':'0','color':'auto'})
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'),'11906'); tcW.set(qn('w:type'),'dxa'); tcPr.append(tcW)
    vAlign = OxmlElement('w:vAlign'); vAlign.set(qn('w:val'),'center'); tcPr.append(vAlign)
    def ctr(before=0, after=0):
        p = cell.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = Pt(14)
        return p
    ctr(50,10).add_run('✨    ⭐    ✨    ✨    ⭐    ✨').font.size = Pt(18)
    run = ctr(60,10).add_run('高考英语词汇默写本'); run.font.size = Pt(40); run.font.bold = True; run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    run = ctr(8,16).add_run('词根版 · 乱序编排 · 艾宾浩斯记忆法'); run.font.size = Pt(16); run.font.color.rgb = RGBColor(0xFF,0xD7,0x66)
    run = ctr(40,12).add_run('🌊🌊  ⛵  ⛵⛵  🌊🌊🌊  ⛵  🌊🌊'); run.font.size = Pt(24)
    run = ctr(36,12).add_run(f'词根词汇 {rt}词 ({rr}组) · 基础词汇 {bt}词 ({br}组) · 共 {TOTAL_WORDS}词'); run.font.size = Pt(11); run.font.color.rgb = RGBColor(0xBB,0xDD,0xFF); run.font.italic = True
    run = ctr(20,8).add_run('「今天的每一分努力」\r都是明天看得见的风景'); run.font.size = Pt(12); run.font.color.rgb = RGBColor(0xDD,0xEE,0xFF); run.font.italic = True
    ctr(10,20).add_run('⭐  ✨  ⭐  ✨  ⭐').font.size = Pt(14)
    ctr(6,30).add_run('⛵  🌅  ⛵  ⛵  🌅  ⛵').font.size = Pt(20)
    run = ctr(28,4).add_run('苏州盈信企业管理有限公司'); run.font.size = Pt(13); run.font.bold = True; run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    run = ctr(2,2).add_run('公司注册 · 专注财税二十五年'); run.font.size = Pt(10); run.font.color.rgb = RGBColor(0xBB,0xDD,0xFF)
    run = ctr(2,40).add_run('18912633863'); run.font.size = Pt(13); run.font.bold = True; run.font.color.rgb = RGBColor(0xFF,0xD7,0x66)

# ═══ PAGE FOOTER ═══
def add_page_number_footer(doc):
    """Add '2/365' style page number to the content section footer."""
    section = doc.sections[-1]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    
    # PAGE field
    fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin')
    r1 = p.add_run(); r1.font.size = Pt(7); r1._r.append(fc1)
    it1 = OxmlElement('w:instrText'); it1.text = ' PAGE '
    r2 = p.add_run(); r2.font.size = Pt(7); r2._r.append(it1)
    fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'end')
    r3 = p.add_run(); r3.font.size = Pt(7); r3._r.append(fc2)
    
    r4 = p.add_run(' / '); r4.font.size = Pt(7); r4.font.color.rgb = RGBColor(0x99,0x99,0x99)
    
    # NUMPAGES field
    fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'), 'begin')
    r5 = p.add_run(); r5.font.size = Pt(7); r5._r.append(fc3)
    it2 = OxmlElement('w:instrText'); it2.text = ' NUMPAGES '
    r6 = p.add_run(); r6.font.size = Pt(7); r6._r.append(it2)
    fc4 = OxmlElement('w:fldChar'); fc4.set(qn('w:fldCharType'), 'end')
    r7 = p.add_run(); r7.font.size = Pt(7); r7._r.append(fc4)

# ═══ CARD ═══
def card_learn(doc, w):
    t = doc.add_table(rows=2, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    sm = '★ ' if w['stars'] >= 4 else ''
    c = t.cell(0,0); c.text = ''
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{sm}{w['word']}")
    r.font.size = Pt(12); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E) if w['stars']>=4 else RGBColor(0x33,0x33,0x33)
    if w['phonetic']: r = p.add_run(f"  {w['phonetic']}"); r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x66,0x66,0x66)
    set_shd(c, 'F0F4F8'); c.width = Cm(7.0)
    c = t.cell(0,1); c.text = ''
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{w['definition']}   "); r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xBB,0xBB,0xBB); r.font.italic = True
    r = p.add_run(f"({'＿'*12})"); r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xCC,0xCC,0xCC)
    c.width = Cm(7.5)
    c = t.cell(1,0); c.text = ''
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    if w['root']:
        r = p.add_run(f"🔬 词根【{w['root']}】= {w['root_meaning']}"); r.font.size = Pt(7.5); r.font.italic = True; r.font.color.rgb = RGBColor(0x2C,0x3E,0x50)
    else:
        r = p.add_run("📖 基础词汇，无特定词根"); r.font.size = Pt(7.5); r.font.italic = True; r.font.color.rgb = RGBColor(0x2C,0x3E,0x50)
    c = t.cell(1,1); c.text = ''
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    dots = '●'*w['stars']+'○'*(5-w['stars'])
    r = p.add_run(f"写作指数: {dots}  "); r.font.size = Pt(7); r.font.color.rgb = RGBColor(0x88,0x88,0x88)
    if w['stars'] >= 4: r = p.add_run('⭐ 高频词'); r.font.size = Pt(7); r.font.bold = True; r.font.color.rgb = RGBColor(0xC0,0x39,0x2B)
    for rowr in t.rows:
        for cell in rowr.cells:
            set_bdr(cell,top={'val':'single','sz':'4','color':'DDD'},bottom={'val':'single','sz':'4','color':'DDD'},
                    start={'val':'single','sz':'4','color':'DDD'},end={'val':'single','sz':'4','color':'DDD'})
    sp = doc.add_paragraph(); sp.paragraph_format.space_before = Pt(0); sp.paragraph_format.space_after = Pt(0)

# ═══ REVIEW TITLE ═══
def review_title(doc, label, subtitle):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(label); r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    r = p.add_run(f'  {subtitle}'); r.font.size = Pt(7); r.font.color.rgb = RGBColor(0x99,0x99,0x99)

def name_date_table(doc, total):
    t = doc.add_table(rows=1, cols=3); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,(txt,w,clr) in enumerate([('姓名: ___________',Cm(6.0),'2C3E50'),
        ('日期: ___________',Cm(5.0),'2C3E50'),(f'正确: ___ / {total}',Cm(4.0),'C0392B')]):
        c=t.cell(0,i); c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1); p.paragraph_format.line_spacing=Pt(8)
        r=p.add_run(txt); r.font.size=Pt(8); r.font.bold=True; r.font.color.rgb=RGBColor(*hex_rgb(clr))
        set_shd(c,'F5F5F5')
        set_bdr(c,top={'val':'single','sz':'4','color':'CCC'},bottom={'val':'single','sz':'4','color':'CCC'},
                start={'val':'single','sz':'4','color':'CCC'},end={'val':'single','sz':'4','color':'CCC'})
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)

# ═══ SIX-COL REVIEW (col_widths=[0.35, 3.48, 5.17, 0.35, 3.48, 5.17]) ═══
CW = [Cm(0.2), Cm(3.56), Cm(5.24), Cm(0.2), Cm(3.56), Cm(5.24)]

def sixcol(doc, word_list, mode):
    total = len(word_list); rn = math.ceil(total/2)
    t = doc.add_table(rows=rn+1, cols=6)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Full width
    tbl = t._tbl; tblPr = tbl.tblPr
    if tblPr is None: tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'),'11000'); tblW.set(qn('w:type'),'dxa'); tblPr.append(tblW)
    h = {'en2cn':['#','英文','中文（填空）','#','英文','中文（填空）'],
         'cn2en':['#','中文释义','英文（填空）','#','中文释义','英文（填空）'],
         'phonetic':['#','🔊音标','英+中（填空）','#','🔊音标','英+中（填空）'],
         'mixed':['#','原文','填空','#','原文','填空']}.get(mode, ['#','原文','填空','#','原文','填空'])
    for ci in range(6):
        c=t.cell(0,ci); c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
        r=p.add_run(h[ci]); r.font.size=Pt(8); r.font.bold=True; r.font.color.rgb=RGBColor(0x33,0x33,0x33)
        set_shd(c,'EAEAEA')
        set_bdr(c,top={'val':'single','sz':'4','color':'AAA'},bottom={'val':'single','sz':'6','color':'666'},
                start={'val':'single','sz':'4','color':'AAA'},end={'val':'single','sz':'4','color':'AAA'})
        c.width=CW[ci]
    trPr0=t.rows[0]._tr.get_or_add_trPr()
    trH0=OxmlElement('w:trHeight'); trH0.set(qn('w:val'),'400'); trH0.set(qn('w:hRule'),'atLeast'); trPr0.append(trH0)
    for ri in range(rn):
        for col_offset, idx in [(0, ri), (3, ri+rn)]:
            if idx>=total: continue
            w=word_list[idx]; sm='★ ' if w['stars']>=4 else ''
            c=t.cell(ri+1, col_offset); c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
            r=p.add_run(str(idx+1)); r.font.size=Pt(7); r.font.color.rgb=RGBColor(0x99,0x99,0x99)
            tcPr=c._tc.get_or_add_tcPr(); vAlign=OxmlElement('w:vAlign'); vAlign.set(qn('w:val'),'center'); tcPr.append(vAlign)
            c=t.cell(ri+1, col_offset+1); c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
            if mode=='en2cn':
                r=p.add_run(f"{sm}{w['word']}"); r.font.size=Pt(10); r.font.bold=True
                r.font.color.rgb=RGBColor(0x1A,0x1A,0x2E) if w['stars']>=4 else RGBColor(0x33,0x33,0x33)
            elif mode=='cn2en':
                r=p.add_run(w['definition'][:30]); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x44,0x44,0x44)
            elif mode=='phonetic':
                r=p.add_run(f"🔊{w['phonetic']}"); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x2C,0x3E,0x50)
            elif mode=='mixed':
                is_en = (idx%2==0)
                if is_en:
                    r=p.add_run(f"{sm}{w['word']}"); r.font.size=Pt(10); r.font.bold=True
                    r.font.color.rgb=RGBColor(0x1A,0x1A,0x2E) if w['stars']>=4 else RGBColor(0x33,0x33,0x33)
                else:
                    r=p.add_run(w['definition'][:28]); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x44,0x44,0x44)
            tcPr=c._tc.get_or_add_tcPr(); vAlign=OxmlElement('w:vAlign'); vAlign.set(qn('w:val'),'center'); tcPr.append(vAlign)
            c=t.cell(ri+1, col_offset+2); c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
            if mode=='en2cn': r=p.add_run(f"({'＿'*14})")
            elif mode=='cn2en': r=p.add_run(f"({'＿'*14})")
            elif mode=='phonetic':
                r=p.add_run(f"英[{'＿'*8}]"); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0xCC,0xCC,0xCC)
                r.add_break()
                r=p.add_run(f"中[{'＿'*8}]"); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0xCC,0xCC,0xCC)
            elif mode=='mixed':
                if idx%2==0: r=p.add_run(f"中文[{'＿'*10}]")
                else: r=p.add_run(f"英文[{'＿'*10}]")
            r.font.size=Pt(9); r.font.color.rgb=RGBColor(0xCC,0xCC,0xCC)
            tcPr=c._tc.get_or_add_tcPr(); vAlign=OxmlElement('w:vAlign'); vAlign.set(qn('w:val'),'center'); tcPr.append(vAlign)
            for cix in range(3):
                set_bdr(t.cell(ri+1, col_offset+cix),
                    top={'val':'single','sz':'4','color':'EEE'},bottom={'val':'single','sz':'4','color':'EEE'},
                    start={'val':'single','sz':'4','color':'EEE'},end={'val':'single','sz':'4','color':'EEE'})
        trPr=t.rows[ri+1]._tr.get_or_add_trPr()
        trH=OxmlElement('w:trHeight'); trH.set(qn('w:val'),'600'); trH.set(qn('w:hRule'),'atLeast'); trPr.append(trH)
    for row in t.rows:
        for ci in range(6): row.cells[ci].width=CW[ci]

# ═══ ENCOURAGEMENT ═══
ENCOURAGE = [
    (5,  [
        "🎉 你已完成总词汇量的5%，加油！好的开始是成功的一半！",
        "🚀 5%完成！每一步都在靠近目标！",
        "🌱 种子已经种下，5%的词汇正在生根发芽！",
    ]),
    (10, [
        "💪 很厉害！你又进一步，记住了总词汇量的10%！",
        "📈 10%达成！你已经看到了进步的趋势！",
        "🎯 十分之一完成，九分之九还在路上！",
    ]),
    (15, [
        "✨ 百尺竿头，更进一步，马上要完成15%了！",
        "🔥 15%！记住的单词正在悄然改变你的大脑！",
        "🏃 跑过了15%的路程，节奏已经起来了！",
    ]),
    (20, [
        "🚀 已经完成20%了！词根记忆法正在发力！",
        "⭐ 五分之一的词汇已被征服！每一分努力都算数！",
        "💎 20%！累计的词汇量开始产生化学反应！",
    ]),
    (25, [
        "🌟 四分之一的词汇已经被你征服！继续前进！",
        "🎊 25%！四分之一里程碑，值得庆祝！",
        "📚 一本词典的四分之一，你已经装进脑子里了！",
    ]),
    (30, [
        "🔥 30%！记住的每个词都会在考试中帮你得分！",
        "⚡ 接近三分之一了，你的词汇量已经碾压大多数人了！",
        "🎪 30%达成！越到后面词根串记越轻松！",
    ]),
    (35, [
        "💎 超过三分之一了！你比昨天的自己更强大！",
        "🌈 35%！越来越多的单词变得面熟了！",
        "🏗 词汇大厦已经建了三成五，地基已经扎实！",
    ]),
    (40, [
        "⚡ 40%的词汇量，你已经超过大多数人了！",
        "🎯 四成完成！节奏稳了，继续冲！",
        "💰 每记一个词，都是在为未来存钱！40%了！",
    ]),
    (45, [
        "🎯 45%！距离半程只差一步之遥！",
        "🏔 45%！登顶前的最后冲刺！",
        "👀 45%完成，你已经开始用英语思维了！",
    ]),
    (50, [
        "🎊 哇，加油啊，已经过半了！50%的词汇已拿下！",
        "🥇 半程冠军！你已经比开始时强了一倍！",
        "🏁 50%！一半已过，剩下的都是下坡路！",
        "🎆 里程碑！所有单词你已认识一半！",
    ]),
    (55, [
        "🌈 55%！下半程开始，冲刺吧！",
        "⛰ 55%！回头看，起点已经很远了！",
        "📊 55%完成率，你的坚持令人敬佩！",
    ]),
    (60, [
        "🏆 60%！大多数词汇你都已经掌握了！",
        "🎯 六成完成！量变正在引起质变！",
        "🚀 60%！从这往后越来越顺！",
    ]),
    (65, [
        "⭐ 65%！三度已过二，胜利在望！",
        "💪 接近七成！你的词汇量已经让大多数人仰望！",
        "🌄 65%！山顶越来越近了！",
    ]),
    (70, [
        "🚩 70%！你已经进入高阶词汇区了！",
        "👑 七成！你的英语水平已经上了个大台阶！",
        "🔥 70%！离精通只差30%！",
    ]),
    (75, [
        "🎪 四分之三！剩下的都是小菜一碟！",
        "🏗 75%！词汇大厦即将封顶！",
        "🎯 四分之三完成，你已经是词汇高手了！",
    ]),
    (80, [
        "👑 80%！你已经是一位词汇达人了！",
        "💎 八成！剩下的20%用联想记忆轻松拿下！",
        "🌟 80%！你已经突破了英语词汇的关键门槛！",
    ]),
    (85, [
        "💥 85%！离胜利只差最后一程！",
        "🏃 冲刺阶段！最后15%需要的只是坚持！",
        "🎆 85%！回头看一路走来，所有的努力都值得！",
    ]),
    (90, [
        "🎆 90%！最后的冲刺，冲鸭！",
        "🏁 九成完成！最后10%决定胜负！",
        "🥇 90%！你已经站在了词汇金字塔的顶端！",
    ]),
    (95, [
        "🏁 95%！胜利就在眼前！",
        "🚀 最后5%，冲刺！每一个词都在创造纪录！",
        "🌟 95%！你是千分之五十的坚持者！",
    ]),
    (100,[
        "🥇 🎉🎉🎉 100%！恭喜你完成了全部词汇的学习！太棒了！",
        "🎊🏆🎊 100%完美通关！你已经是英语词汇大师！",
        "💎 100%！所有的汗水在这一刻绽放！恭喜你！",
    ]),
]

GENERIC_MSGS = [
    "📖 第{}组完成！词汇大厦稳步添砖加瓦，继续加油！",
    "💪 又一组拿下！积少成多，你比昨天更强了！",
    "🎯 第{}组完成！每50词都在拉近你和目标之间的距离！",
    "🌟 坚持就是胜利！第{}组记住了，下一组也不在话下！",
    "📈 第{}组完成！词汇量正在以肉眼可见的速度增长！",
    "⚡ 又消灭了一组！你的词汇库越来越大！",
    "🏃 第{}组顺利通关！保持节奏，胜利就在前方！",
    "🎯 第{}组收入囊中！已经上瘾了这种通关的感觉吧？",
    "💎 每组50词，第{}组完成！聚沙成塔的力量！",
    "🚀 第{}组拿下！离目标又近了50词！",
]

def get_encouragement(g_idx, total_g, cum_words):
    """Always return a message. Milestone messages at key %s, generic otherwise."""
    pct = cum_words * 100 / TOTAL_WORDS
    # Check milestones first
    for milestone, msgs in ENCOURAGE:
        if abs(pct - milestone) < 2.0:  # within 2% of milestone
            return random.choice(msgs)
        if pct >= milestone and milestone == 100:
            return random.choice(msgs)
    # Fall back to generic
    return random.choice(GENERIC_MSGS).format(g_idx + 1)

# ═══════════════════════════════════════════
# BUILD
# ═══════════════════════════════════════════
print("Creating document...")
doc = Document()
style = doc.styles['Normal']; style.font.name = '微软雅黑'; style.font.size = Pt(9)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
for sec in doc.sections:
    sec.top_margin = Cm(0); sec.bottom_margin = Cm(0)
    sec.left_margin = Cm(0); sec.right_margin = Cm(0)

# Cover
build_cover(doc, len(root_words), len(basic_words), len(root_groups), len(basic_groups))
# Content section
sec2 = doc.add_section()
sec2.top_margin = Cm(0.8); sec2.bottom_margin = Cm(0.8)
sec2.left_margin = Cm(0.8); sec2.right_margin = Cm(0.8)
add_page_number_footer(doc)

# Instructions
print("Instructions...")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('📋 使用说明'); r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
for t in [
    f'📖 Part 1: 词根词汇（{len(root_words)}词 / {len(root_groups)}组，完整6轮艾宾浩斯）',
    f'📝 Part 2: 基础词汇（{len(basic_words)}词 / {len(basic_groups)}组，4轮复习无学习环节）',
    '🎲 各组内乱序编排，非字母顺序',
    '✍ 写作指数基于真实英文语料库频次数据（1-5星）',
]:
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(t); r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x44,0x44,0x44)

# ═══════════════════════════════════════
# PART 1: ROOT WORDS (full 6 rounds + encouragement)
# ═══════════════════════════════════════
print(f"=== PART 1: {len(root_groups)} root groups ===")
cumulative_so_far = 0

for gi, g in enumerate(root_groups):
    wl = g['words']; tw = len(wl)
    t0 = time.time()
    print(f"  Root {gi+1}/{len(root_groups)} ({tw}w)...", end='')
    
    # DAY 1
    doc.add_page_break()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'DAY 1  学习页  · Part 1 · 第{gi+1}/{len(root_groups)}组')
    r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'词根: {g["label"]}  ·  {tw}词汇'); r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x99,0x99,0x99)
    for w in wl: card_learn(doc, w)
    
    # DAY 2
    doc.add_page_break()
    review_title(doc, f'DAY 2 · Part 1 · 第{gi+1}组', f'英→中 · {g["label"]} · {tw}词')
    name_date_table(doc, tw)
    sixcol(doc, wl, 'en2cn')
    
    # DAY 4
    doc.add_page_break()
    review_title(doc, f'DAY 4 · Part 1 · 第{gi+1}组', f'中→英 · {g["label"]} · {tw}词')
    name_date_table(doc, tw)
    sixcol(doc, wl, 'cn2en')
    
    # DAY 7
    doc.add_page_break()
    review_title(doc, f'DAY 7 · Part 1 · 第{gi+1}组', f'🔊 音标 · {g["label"]} · {tw}词')
    name_date_table(doc, tw)
    sixcol(doc, wl, 'phonetic')
    
    # DAY 15
    sl = [w for w in wl if w['stars'] >= 4]
    if sl:
        doc.add_page_break()
        review_title(doc, f'DAY 15 · Part 1 · 第{gi+1}组', f'★ 高频词 · ⭐{len(sl)}词')
        name_date_table(doc, len(sl))
        sixcol(doc, sl, 'cn2en')
    
    # DAY 30 (mixed) + encouragement
    doc.add_page_break()
    review_title(doc, f'DAY 30 · Part 1 · 第{gi+1}组', f'🎲 混合终测 · {g["label"]} · {tw}词')
    name_date_table(doc, tw)
    sixcol(doc, wl, 'mixed')
    
    # Encouragement at bottom
    cumulative_so_far += tw
    msg = get_encouragement(gi, len(root_groups), cumulative_so_far)
    if msg:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(6)
        r = p.add_run(msg); r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = RGBColor(0xC0,0x39,0x2B)
    
    dt = time.time()-t0
    print(f" {dt:.0f}s")

# ═══════════════════════════════════════
# PART 2: BASIC WORDS (4 rounds, no Day1/Day30)
# ═══════════════════════════════════════
print(f"\n=== PART 2: {len(basic_groups)} basic groups ===")

for gi, g in enumerate(basic_groups):
    wl = g['words']; tw = len(wl)
    t0 = time.time()
    print(f"  Basic {gi+1}/{len(basic_groups)} ({tw}w)...", end='')
    
    # DAY 2: en2cn
    doc.add_page_break()
    review_title(doc, f'DAY 2 · Part 2 · 第{gi+1}/{len(basic_groups)}组', f'英→中 · 基础词汇 · {tw}词')
    name_date_table(doc, tw)
    sixcol(doc, wl, 'en2cn')
    
    # DAY 4: cn2en
    doc.add_page_break()
    review_title(doc, f'DAY 4 · Part 2 · 第{gi+1}/{len(basic_groups)}组', f'中→英 · 基础词汇 · {tw}词')
    name_date_table(doc, tw)
    sixcol(doc, wl, 'cn2en')
    
    # DAY 7: phonetic
    doc.add_page_break()
    review_title(doc, f'DAY 7 · Part 2 · 第{gi+1}/{len(basic_groups)}组', f'🔊 音标 · 基础词汇 · {tw}词')
    name_date_table(doc, tw)
    sixcol(doc, wl, 'phonetic')
    
    # DAY 15: stars
    sl = [w for w in wl if w['stars'] >= 4]
    if sl:
        doc.add_page_break()
        review_title(doc, f'DAY 15 · Part 2 · 第{gi+1}/{len(basic_groups)}组', f'★ 高频词 · ⭐{len(sl)}词')
        name_date_table(doc, len(sl))
        sixcol(doc, sl, 'cn2en')
    # Encouragement after last review round of this basic group
    msg_b = random.choice([
        "📝 基础词汇第{}组全部通关！坚持就是胜利！",
        "💪 基础词汇又搞定一组！稳扎稳打！",
        "🎯 第{}组基础词汇完成！基础牢固才能走得更远！",
        "📖 基础词汇第{}组拿下！你已经超过大多数人了！",
        "🌟 第{}组基础词汇通关！每一步都算数！",
        "🔥 基础词汇又一组拿下！积少成多！",
    ]).format(gi+1)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(msg_b); r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = RGBColor(0xC0,0x39,0x2B)
    
    dt = time.time()-t0
    print(f" {dt:.0f}s")

# ═══ SAVE ═══
print("Saving...")
doc.save(OUT)
fsize = os.path.getsize(OUT)/1024
print(f"\n✅ {OUT} ({fsize:.0f}KB)")
print(f"📊 {TOTAL_WORDS}词")
print(f"📖 Part 1: 词根词汇 {len(root_words)}词 / {len(root_groups)}组")
print(f"📝 Part 2: 基础词汇 {len(basic_words)}词 / {len(basic_groups)}组")
print(f"🎯 {len(root_groups)}组 × 6轮 + {len(basic_groups)}组 × 4轮")