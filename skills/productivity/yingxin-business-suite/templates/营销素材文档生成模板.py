"""
营销素材文档生成模板 — 通用版
基于tax-planning-fin-analysis-industry技能

用途：生成带结构化内容的营销素材Word文档（如高级会计师数据分析、行业调研报告等）
依赖：pip install python-docx

使用方式：
1. 修改 build_content() 函数，返回一个章节列表
2. 每个章节结构: {'title': '章节标题', 'level': 1|2|3, 'type': 'text'|'table'|'script_table', 'content': ...}
3. 运行: python3 本文件.py
4. 输出：Word(.docx) 到桌面
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

# ═══════════════════════════════════════════
# ★★★ 修改区 START ★★★
# ═══════════════════════════════════════════

DOC_TITLE = '文档标题'           # 封面标题
DOC_SUBTITLE = '文档副标题'      # 封面副标题
DOC_FILENAME = '文档文件名'      # 不含扩展名，如'高级会计师名单_营销素材库'

def build_content():
    """
    返回章节列表，每章格式：
    {
        'title': '一、章节标题',
        'level': 1,             # 1=大标题, 2=中标题, 3=小标题
        'type': 'text',         # text/table/script_table
        'content': ...          # 根据type不同格式不同
    }

    text类型: content = '纯文本内容'
    table类型: content = {
        'headers': ['列1', '列2'],
        'rows': [['行1列1', '行1列2'], ['行2列1', '行2列2']]
    }
    script_table类型（分镜表）: content = {
        'headers': ['时间', '画面', '口播文案', '字幕'],
        'rows': [['0:00', '画面描述', '口播文字', '字幕文字'], ...]
    }
    """
    return [
        {
            'title': '一、章节示例',
            'level': 1,
            'type': 'text',
            'content': '这是一个文本段落示例。替换为你的实际内容。'
        },
        {
            'title': '1.1 数据表格示例',
            'level': 2,
            'type': 'table',
            'content': {
                'headers': ['维度', '数据'],
                'rows': [
                    ['行1', '数据A'],
                    ['行2', '数据B'],
                ]
            }
        },
        {
            'title': '脚本示例',
            'level': 2,
            'type': 'script_table',
            'content': {
                'headers': ['时间', '画面', '口播文案', '字幕（大号）'],
                'rows': [
                    ['0:00', '镜头怼脸', '"这是第一句口播。"', '第一句字幕'],
                    ['0:05', '切画面', '"这是第二句口播。"', '第二句字幕'],
                ]
            }
        },
    ]

# ═══════════════════════════════════════════
# ★★★ 修改区 END ★★★
# ═══════════════════════════════════════════

# ── 文档生成引擎 ──

def build_doc(title, subtitle, chapters):
    doc = Document()

    # 全局样式
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.paragraph_format.space_after = Pt(6)

    def set_font(run, size=10.5, bold=False, color=None):
        run.font.name = '微软雅黑'
        run.font.size = Pt(size)
        run.bold = bold
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if color:
            run.font.color.rgb = color

    def add_heading_para(text, level=1):
        p = doc.add_paragraph()
        run = p.add_run(text)
        if level == 1:
            set_font(run, 16, True, RGBColor(0x1F, 0x49, 0x7D))
            p.paragraph_format.space_before = Pt(18)
        elif level == 2:
            set_font(run, 13, True, RGBColor(0x1F, 0x49, 0x7D))
            p.paragraph_format.space_before = Pt(12)
        else:
            set_font(run, 11, True, RGBColor(0x40, 0x40, 0x40))
            p.paragraph_format.space_before = Pt(8)
        return p

    def add_body(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_font(run)
        return p

    def add_table(headers, rows, font_size=9):
        t = doc.add_table(rows=1+len(rows), cols=len(headers))
        t.style = 'Light Grid Accent 1'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(font_size)
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = t.rows[ri+1].cells[ci]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.size = Pt(font_size)
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        return t

    # 封面
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_font(run, 22, True, RGBColor(0x1F, 0x49, 0x7D))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    set_font(run, 10, False, RGBColor(0x66, 0x66, 0x66))

    doc.add_paragraph()

    # 正文
    for ch in chapters:
        if ch['type'] == 'text':
            if ch.get('level'):
                add_heading_para(ch['title'], ch['level'])
            add_body(ch['content'])
        elif ch['type'] == 'table':
            if ch.get('title'):
                add_heading_para(ch['title'], ch.get('level', 2))
            c = ch['content']
            add_table(c['headers'], c['rows'])
            doc.add_paragraph()
        elif ch['type'] == 'script_table':
            if ch.get('title'):
                add_heading_para(ch['title'], ch.get('level', 2))
            c = ch['content']
            add_table(c['headers'], c['rows'], font_size=8)
            doc.add_paragraph()

    # 结尾线
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— 完 —')
    set_font(run, 9, False, RGBColor(0x99, 0x99, 0x99))

    return doc


# ═══════════════════════════════════════════
# 生成与保存
# ═══════════════════════════════════════════

if __name__ == '__main__':
    chapters = build_content()
    doc = build_doc(DOC_TITLE, DOC_SUBTITLE, chapters)

    output_dir = '/mnt/d/360MoveData/Users/Admin/Desktop/'
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f'{DOC_FILENAME}.docx')
    doc.save(output_path)

    print(f'✅ 文档已保存: {output_path}')
    print(f'   文件大小: {os.path.getsize(output_path):,} 字节')
    print(f'   章节数: {len(chapters)}')
